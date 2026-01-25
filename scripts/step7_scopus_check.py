#!/usr/bin/env python3
"""
step7_scopus_check.py

Eligible-benchmark vs Scopus retrieval comparison with iteration tracking.

Compares Scopus retrieval (Step 2) against benchmark lists and eligible sets.

Benchmark list (full):
- Step5 wide eligibility file (step5_eligibility_wide.csv) merged with benchmark metadata
- Plotted as: step7_retrieval_analysis_all_2x2.png
- Archived as: history/iterXXX_benchmark_all_2x2.png

Eligible sets (overall-only bars):
- Baseline eligible (from step6 excel; fallback step5 passed): step7_retrieval_analysis_eligible_2x2.png

Benchmark item coverage summary (across iterations):
- step7_benchmark_iteration_matrix.csv
- step7_benchmark_iteration_heatmap.png
- Archived per iteration: history/iterXXX_full_benchmark_status.csv

Suggestions:
- step7_searchstring_suggestions.txt includes:
  - per-block routing
  - final summary (deduped) with title_hits / abstract_hits / score

Appendix:
- step7_appendix_searchstring_evolution.docx includes iteration summaries and embedded figures,
  but does NOT include suggested search string improvements.
"""

from __future__ import annotations

import os
import re
import json
import time
import html
import glob
import hashlib
import shutil
from typing import Any, Optional, List, Tuple, Dict
from collections import Counter, defaultdict

import numpy as np
import requests
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.colors import ListedColormap
from tqdm import tqdm
import subprocess


try:
    from docx import Document
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("'python-docx' not installed. Word files will not be generated. Run: pip install python-docx")


PALETTE_RETRIEVAL = {True: "#93c47d", False: "#e06666"}
CATS_SOURCE = ["Peer Reviewed", "Grey Literature"]

BASELINE_ELIGIBILITY_XLSX_BASENAME = "eligibility_report.xlsx"
BASELINE_ELIGIBILITY_SHEETNAME = "All Papers"
STEP1_QUERIES_JSON_BASENAME = "step1_queries.json"
STEP1_PLOT_BASENAME = "step1_hits_plot.png"
STEP1_SUMMARY_BASENAME = "step1_summary.json"

CROSSREF_BASE = "https://api.crossref.org/works/"
CROSSREF_TIMEOUT = 7

STOPWORDS = {
    "the", "and", "for", "with", "from", "into", "over", "under", "between", "within",
    "towards", "toward", "among", "using", "use", "used", "based", "study", "studies",
    "analysis", "effects", "effect", "impact", "impacts", "results", "evidence",
    "case", "cases", "approach", "methods", "method", "data", "paper", "review",
    "climate", "change", "adaptation", "adaptive", "agriculture", "farm", "farming",
    "farmer", "farmers", "rural", "smallholder", "smallholders",
    "this", "that", "these", "those", "their", "there", "here", "where", "when",
    "what", "which", "who", "whom", "whose", "been", "being", "were", "was", "are",
    "also", "very", "more", "most", "less", "many", "some", "such", "than", "then",
    "onto", "across", "via", "per", "without",
}

GENERIC_BAD_SUGGESTIONS = {
    "journal", "international", "understanding", "measures", "determinants", "strategies",
    "central", "towards", "based", "using", "study", "studies", "analysis", "results",
    "evidence", "review", "effects", "effect", "impacts", "impact", "data", "paper",
}

ELIGIBILITY_BOOST_TERMS = {
    "evaluation", "evaluations", "impact", "impacts", "effect", "effects", "effectiveness", "efficacy",
    "outcome", "outcomes", "performance", "assessment", "assessments", "evidence", "results",
    "randomized", "randomised", "trial", "trials", "quasi", "baseline",
    "counterfactual", "matching", "propensity", "regression", "instrumental",
    "difference", "differences", "rct", "experiment", "experimental",
    "intervention", "interventions", "program", "programme", "programs", "policy", "policies",
    "insurance", "index", "irrigation", "extension", "advisory", "forecast", "forecasts",
    "climatesmart", "climate-smart", "agroforestry", "conservation", "mulching", "fertilizer", "fertiliser",
    "seed", "seeds", "varieties", "drought", "droughts", "heat", "flood", "floods",
    "resilience", "resilient", "risk", "risks", "coping", "mitigation",
    "yield", "yields", "income", "incomes", "profit", "profits", "livelihood", "livelihoods",
    "nutrition", "food", "security", "livestock", "crop", "crops", "pastoral", "pastoralist",
    "smallholder", "smallholders", "farmer", "farmers", "household", "households",
}

ROUTE_TOKENS = {
    "M__measurement_eval": {
        "evaluation", "evaluations", "impact", "impacts", "effect", "effects", "effectiveness",
        "assessment", "assessments", "indicator", "indicators", "metric", "metrics", "monitor", "monitoring",
        "survey", "surveys", "randomized", "randomised", "trial", "trials", "rct", "experiment",
        "counterfactual", "matching", "propensity", "baseline", "difference", "differences", "regression",
        "quasi", "method", "methods", "mel", "m&e", "evaluation", "evaluat",
    },
    "C_context_climate__hazards": {
        "drought", "droughts", "flood", "floods", "heat", "heatwave", "heatwaves", "rainfall",
        "aridity", "storm", "storms", "cyclone", "cyclones", "hurricane", "typhoon", "salinity",
        "intrusion", "shock", "shocks", "hazard", "hazards",
    },
    "C_context_climate__climate_framing": {
        "climate", "variability", "resilient", "resilience", "risk", "risks", "hazard", "hazards",
        "exposure", "shock", "shocks", "climatesmart", "csa",
    },
    "C_context_agriculture__ag_systems": {
        "agriculture", "agricultural", "farm", "farming", "crop", "crops", "livestock",
        "irrigation", "rainfed", "rangeland", "fisheries", "aquaculture", "forestry",
        "agroforestry", "silvopastoral", "pastoral", "pastoralist",
    },
    "C_concept__adaptation_actions": {
        "adoption", "uptake", "participatory", "governance", "learning", "practice", "practices",
        "behaviour", "behavior", "decision", "decisions", "strategy", "strategies", "adjustment",
        "adjust", "change", "changes", "coping", "insurance", "extension", "advisory",
    },
    "C_concept__capacities_knowledge": {
        "capacity", "capacities", "knowledge", "information", "awareness", "skills",
        "training", "literacy",
    },
    "C_concept__livelihood_productivity": {
        "income", "incomes", "yield", "yields", "productivity", "profit", "profits", "livelihood",
        "livelihoods",
    },
    "C_concept__resilience_outcomes": {
        "resilience", "resilient", "wellbeing", "well-being", "vulnerability", "exposure",
        "risk", "risks", "outcome", "outcomes",
    },
    "C_concept__maladaptation": {
        "maladaptation", "maladaptive", "maladapt",
    },
    "P__smallholder_small_scale": {
        "smallholder", "smallholders", "subsistence", "resourcepoor", "resource-poor",
        "marginal", "smallscale", "small-scale",
    },
    "P__households_poverty": {
        "household", "households", "poverty", "poor", "lowincome", "low-income", "family",
    },
    "P__livestock_pastoral": {
        "pastoral", "pastoralist", "herder", "livestock", "dairy",
    },
    "P__fisheries_aquaculture": {
        "fishery", "fisheries", "fisher", "fisherfolk", "aquaculture", "shrimp",
    },
    "P__agroforestry_tree": {
        "agroforestry", "silvopastoral", "tree", "trees", "forest", "fruit",
    },
    "P__marginalized_groups": {
        "women", "female", "youth", "indigenous", "tribal", "ethnic", "migrant",
        "landless", "tenant", "sharecropper",
    },
}

# Functions

def _rewrite_jsonl(path: str, rows: List[dict]) -> None:
    with open(path, "w", encoding="utf-8") as f:
        for obj in rows:
            f.write(json.dumps(obj, ensure_ascii=False) + "\n")

def _atomic_docx_save(doc: "Document", dst_path: str) -> None:
    import os, tempfile
    d = os.path.dirname(dst_path) or "."
    fd, tmp = tempfile.mkstemp(prefix=".tmp_", suffix=".docx", dir=d)
    os.close(fd)
    try:
        doc.save(tmp)
        os.replace(tmp, dst_path)  # atomic rename
    finally:
        try:
            if os.path.exists(tmp):
                os.remove(tmp)
        except Exception:
            pass

def _now_local_str() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S %Z", time.localtime())


def _sha1(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8")).hexdigest()


def _safe_str(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float) and pd.isna(val):
        return ""
    return html.unescape(str(val)).strip()


def clean_text(text: Any) -> str:
    if not isinstance(text, str):
        text = _safe_str(text)
    return re.sub(r"[^a-z0-9]", "", text.lower())


def find_col(df: pd.DataFrame, keywords: List[str]) -> Optional[str]:
    kw_lower = [k.lower() for k in keywords]
    for col in df.columns:
        if str(col).strip().lower() in kw_lower:
            return col
    for col in df.columns:
        col_l = str(col).lower()
        for k in keywords:
            k_l = k.lower()
            if len(k_l) > 2 and k_l in col_l:
                return col
    return None


def _clean_doi(val: Any) -> str:
    s = _safe_str(val).lower().strip()
    for p in ("https://doi.org/", "http://doi.org/", "doi:"):
        if s.startswith(p):
            s = s[len(p):]
    return s.strip().rstrip(" .),;]}>").strip()


def _decision_is_include(val: Any) -> bool:
    v = str(val or "").strip().lower()
    return bool(re.search(r"\b(include|included|pass|yes)\b", v))


def _map_source_type(val: Any) -> str:
    v = str(val or "").lower()
    if any(k in v for k in ["grey", "report", "thesis", "working paper", "brief"]):
        return "Grey Literature"
    return "Peer Reviewed"


def _simple_status(val: Any) -> str:
    v = str(val or "").lower()
    if "include" in v or "pass" in v or re.fullmatch(r"\s*yes\s*", v):
        return "Passed"
    if "unclear" in v or "pending" in v or "no_abstract" in v:
        return "Unclear"
    return "Failed"


def _read_text_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception:
        return ""

def _load_step1_summary(out_root: str) -> Tuple[int, dict, str]:
    """
    Returns:
      - step1_total_all (int) from step1_summary.json["total_all"]
      - raw summary dict (or {})
      - source string
    """
    p = os.path.join(out_root, "step1", STEP1_SUMMARY_BASENAME)
    if not os.path.exists(p):
        return 0, {}, "step1_summary_missing"

    try:
        with open(p, "r", encoding="utf-8") as f:
            d = json.load(f) or {}
    except Exception:
        return 0, {}, "step1_summary_unreadable"

    total_all = int(d.get("total_all", 0) or 0)
    return total_all, d, f"step1_summary_from:{os.path.basename(p)}"


def _load_jsonl(path: str) -> List[dict]:
    rows = []
    if not os.path.exists(path):
        return rows
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except Exception:
                continue
    return rows


def _write_jsonl_append(path: str, obj: dict) -> None:
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(obj, ensure_ascii=False) + "\n")


def _resolve_benchmark_csv(out_root: str, configured: Optional[str]) -> str:
    if configured:
        p = str(configured).strip()
        if p and os.path.exists(p):
            return p
        p2 = os.path.join(out_root, p)
        if os.path.exists(p2):
            return p2

    fallback = os.path.join(out_root, "step3", "step3_benchmark_list.csv")
    if os.path.exists(fallback):
        return fallback
    return ""


def _ensure_dir(p: str) -> None:
    os.makedirs(p, exist_ok=True)


def _archive_copy(src_path: str, dst_path: str) -> None:
    try:
        if src_path and os.path.exists(src_path):
            shutil.copy2(src_path, dst_path)
    except Exception:
        pass


def _file_fingerprint(path: Optional[str]) -> str:
    if not path or not os.path.exists(path):
        return "missing"
    try:
        st = os.stat(path)
        return f"{os.path.basename(path)}|{st.st_size}|{int(st.st_mtime)}"
    except Exception:
        return f"{os.path.basename(path)}|unstatable"


def fetch_crossref_metadata(doi: str) -> dict:
    clean_doi = _clean_doi(doi)
    if not clean_doi:
        return {}

    url = f"{CROSSREF_BASE}{clean_doi}"
    try:
        r = requests.get(url, timeout=CROSSREF_TIMEOUT)
        if r.status_code != 200:
            return {}
        data = r.json().get("message", {}) or {}

        authors_list = data.get("author", []) or []
        author_parts = []
        for a in authors_list:
            family = a.get("family")
            given = a.get("given")
            if family:
                name = family
                if given:
                    name += f", {given[0]}."
                author_parts.append(name)
        if len(author_parts) > 3:
            author_str = ", ".join(author_parts[:3]) + " et al."
        else:
            author_str = ", ".join(author_parts)

        container = data.get("container-title", []) or []
        journal = container[0] if container else ""

        issued = (data.get("issued", {}) or {}).get("date-parts", [[None]])[0][0]
        year = str(issued) if issued else ""

        return {"fetched_authors": author_str, "fetched_journal": journal, "fetched_year": year}
    except Exception:
        return {}


def build_citation(row: dict) -> str:
    author = _safe_str(row.get("fetched_authors")) or _safe_str(row.get("author_names")) or _safe_str(row.get("authors")) or "Unknown Author"

    year_raw = _safe_str(row.get("fetched_year")) or _safe_str(row.get("year_scopus")) or _safe_str(row.get("year"))
    y_match = re.search(r"\d{4}", year_raw)
    year_str = y_match.group(0) if y_match else "n.d."

    title = (_safe_str(row.get("title")) or _safe_str(row.get("title_scopus"))).rstrip(". ")

    journal = _safe_str(row.get("fetched_journal")) or _safe_str(row.get("publicationName"))

    doi = _safe_str(row.get("doi")) or _safe_str(row.get("doi_scopus"))
    if doi and not doi.lower().startswith("http"):
        doi = f"https://doi.org/{_clean_doi(doi)}"

    parts = [author, f"({year_str})"]
    if title:
        parts.append(f"'{title}'.")
    if journal:
        parts.append(f"{journal}.")
    else:
        st = _safe_str(row.get("source_type"))
        parts.append("[Grey Literature]." if "Grey" in st else "[Journal unavailable].")
    if doi:
        parts.append(doi)

    return " ".join([p for p in parts if p]).strip()


def export_to_word_citations(df: pd.DataFrame, title: str, filename: str) -> None:
    """
    Exports a journal-ready reference list to DOCX using the full `citation` string.
    - Uses df['citation'] when available.
    - If citation is missing/blank, regenerates via build_citation(row).
    - Applies hanging indent (common for reference lists).
    """
    if not HAS_DOCX:
        return

    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    doc.add_heading(title, level=1)

    if df is None or df.empty:
        doc.add_paragraph("No records.")
        _atomic_docx_save(doc, filename)
        print(f"Saved Word document: {filename}")
        return

    # Ensure citation exists; if not, generate it
    if "citation" not in df.columns:
        df = df.copy()
        df["citation"] = df.apply(lambda r: build_citation(r.to_dict()), axis=1)

    # Write reference list
    for _, row in df.iterrows():
        c = _safe_str(row.get("citation"))
        if not c:
            # regenerate if blank
            c = build_citation(row.to_dict())
        c = _safe_str(c)
        if not c:
            continue

        p = doc.add_paragraph(c)

        # Hanging indent: 0.5" left indent with -0.5" first line
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(6)

    _atomic_docx_save(doc, filename)
    print(f"Saved Word document: {filename}")

def _doc_max_image_width(doc: "Document", pad_in: float = 0.25):
    """
    Returns maximum image width that fits inside page margins.
    pad_in is a little breathing room so Word never wraps/crops.
    """
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    w = usable - Inches(pad_in)
    # Safety floor so we never pass a non-positive width
    return w if w > Inches(1.0) else Inches(1.0)


def _build_appendix_from_evolution(
    appendix_path: str,
    evo_rows: List[dict],
    progress_chart_path: str,
    heatmap_path: str,
) -> None:
    if not HAS_DOCX:
        return

    doc = Document()
    max_w = _doc_max_image_width(doc, pad_in=0.25)

    doc.add_heading("Appendix: Search String Evolution Log", level=1)

    doc.add_paragraph("Running summary (updated each iteration):")
    latest_scopus_total = int((evo_rows[-1].get("scopus_total_returned", 0) if evo_rows else 0) or 0)
    doc.add_paragraph(f"Total studies returned by Scopus search string: {latest_scopus_total}", style="List Bullet")
    if os.path.exists(progress_chart_path):
        try:
            doc.add_picture(progress_chart_path, width=max_w)
        except Exception:
            doc.add_paragraph("[Could not embed progress chart]")

    if heatmap_path and os.path.exists(heatmap_path):
        doc.add_paragraph("Benchmark item retrieval by iteration (eligible baseline only):")
        try:
            doc.add_picture(heatmap_path, width=max_w)
        except Exception as e:
            doc.add_paragraph(f"[Could not embed heatmap: {type(e).__name__}: {e}]")


    doc.add_paragraph("")

    for r in evo_rows:
        ts = r.get("timestamp_local", "")
        it = r.get("iteration", "")
        ss = r.get("search_string", "")
        src = r.get("search_source", "")
        sh = r.get("search_hash", "")
        elig = r.get("eligibility_sources", {}) or {}
        cov = r.get("coverage", {}) or {}
        cov_final = cov.get("final_target", {}) or {}
        missed_n = cov.get("final_missed_n", 0)

        figs = (r.get("figures", {}) or {})
        fig_comp = figs.get("coverage_comparison", "")
        fig_base = figs.get("baseline_eligible_plot", "")
        fig_benchmark = figs.get("benchmark_all_plot", "")
        fig_step1 = figs.get("step1_hits_plot", "") 

        doc.add_page_break()
        doc.add_heading(f"Iteration {it} — {ts}", level=2)

        doc.add_paragraph("Summary:")
        doc.add_paragraph(f"Search hash: {sh} (source: {src})", style="List Bullet")
        sc_total = int(r.get("scopus_total_returned", 0) or 0)
        doc.add_paragraph(f"Total Scopus returned: {sc_total}", style="List Bullet")
        doc.add_paragraph(f"Final target: {elig.get('final_target','')}", style="List Bullet")
        doc.add_paragraph(
            f"Final coverage: {cov_final.get('n_in_overall',0)}/{cov_final.get('n_total_overall',0)} "
            f"({float(cov_final.get('pct_overall',0.0)):.1f}%)",
            style="List Bullet",
        )
        doc.add_paragraph(f"Final missed count: {missed_n}", style="List Bullet")

        doc.add_paragraph("Search string used:")
        doc.add_paragraph(ss or "[missing]")

        # Step 1 plot (counts by query) — insert right after search string
        if fig_step1 and os.path.exists(fig_step1):
            doc.add_paragraph("Figure: Step 1 query hit counts")
            try:
                doc.add_picture(fig_step1, width=max_w)
            except Exception:
                doc.add_paragraph(f"[Could not embed figure: {os.path.basename(fig_step1)}]")
            doc.add_paragraph("")

        for cap, path in [
            ("Figure: Eligible benchmark coverage — Overall only", fig_comp),
        ]:
            if path and os.path.exists(path):
                doc.add_paragraph(cap)
                try:
                    doc.add_picture(path, width=max_w)
                except Exception:
                    doc.add_paragraph(f"[Could not embed figure: {os.path.basename(path)}]")
                doc.add_paragraph("")


    _atomic_docx_save(doc, appendix_path)
    print(f"Rebuilt appendix: {appendix_path}")

def _build_appendix_pdf_from_evolution(
    pdf_path: str,
    evo_rows: List[dict],
    progress_chart_path: str,
    heatmap_path: str,
    timestamp: str,
) -> None:
    # reportlab is already in your env
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
    )
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import utils
    import os
    import html as _html

    os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)

    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    normal = styles["BodyText"]

    # a simple mono-ish style for search strings
    mono = ParagraphStyle(
        "Mono",
        parent=normal,
        fontName="Courier",
        fontSize=9,
        leading=11,
    )

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Appendix: Search String Evolution Log",
    )

    max_w = doc.width  # available width between margins

    def _img(path: str, max_width=max_w, max_height=8.5 * inch):
        # Keep aspect ratio, fit inside max_width/max_height
        if not path or not os.path.exists(path):
            return None
        img = utils.ImageReader(path)
        iw, ih = img.getSize()
        if not iw or not ih:
            return None
        scale = min(max_width / iw, max_height / ih, 1.0)
        return RLImage(path, width=iw * scale, height=ih * scale)

    story = []
    story.append(Paragraph("Appendix: Search String Evolution Log", h1))
    story.append(Paragraph(f"Generated: {_html.escape(timestamp)}", normal))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Running summary (updated each iteration):", normal))
    latest_scopus_total = int((evo_rows[-1].get("scopus_total_returned", 0) if evo_rows else 0) or 0)
    story.append(Paragraph(f"Total studies returned by Scopus search string: {latest_scopus_total}", normal))
    im = _img(progress_chart_path, max_height=5.5 * inch)
    if im:
        story.append(Spacer(1, 0.1 * inch))
        story.append(im)
    else:
        story.append(Paragraph("[Could not embed progress chart]", normal))

    if heatmap_path:
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Benchmark item retrieval by iteration (eligible baseline only):", normal))
        im2 = _img(heatmap_path, max_height=6.5 * inch)
        if im2:
            story.append(Spacer(1, 0.1 * inch))
            story.append(im2)
        else:
            story.append(Paragraph("[Could not embed heatmap]", normal))

    for r in (evo_rows or []):
        ts = str(r.get("timestamp_local", "") or "")
        it = str(r.get("iteration", "") or "")
        ss = str(r.get("search_string", "") or "")
        src = str(r.get("search_source", "") or "")
        sh = str(r.get("search_hash", "") or "")

        elig = r.get("eligibility_sources", {}) or {}
        cov = r.get("coverage", {}) or {}
        cov_final = cov.get("final_target", {}) or {}
        missed_n = int(cov.get("final_missed_n", 0) or 0)

        figs = (r.get("figures", {}) or {})
        fig_step1 = figs.get("step1_hits_plot", "")
        fig_comp = figs.get("coverage_comparison", "")

        story.append(PageBreak())
        story.append(Paragraph(f"Iteration {it} — {_html.escape(ts)}", h2))

        story.append(Paragraph("Summary:", normal))
        story.append(Paragraph(f"• Search hash: {_html.escape(sh)} (source: {_html.escape(src)})", normal))
        story.append(Paragraph(f"• Final target: {_html.escape(str(elig.get('final_target','') or ''))}", normal))

        n_in = int(cov_final.get("n_in_overall", 0) or 0)
        n_total = int(cov_final.get("n_total_overall", 0) or 0)
        pct = float(cov_final.get("pct_overall", 0.0) or 0.0)
        story.append(Paragraph(f"• Final coverage: {n_in}/{n_total} ({pct:.1f}%)", normal))
        story.append(Paragraph(f"• Final missed count: {missed_n}", normal))

        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Search string used:", normal))

        # Escape for Paragraph and preserve line breaks a bit
        ss_esc = _html.escape(ss).replace("\n", "<br/>")
        story.append(Paragraph(ss_esc or "[missing]", mono))

        def _add_fig(caption: str, path: str):
            if not path or not os.path.exists(path):
                return
            story.append(Spacer(1, 0.15 * inch))
            story.append(Paragraph(_html.escape(caption), normal))
            imx = _img(path, max_height=6.8 * inch)
            if imx:
                story.append(Spacer(1, 0.1 * inch))
                story.append(imx)

        _add_fig("Figure: Step 1 query hit counts", fig_step1)
        _add_fig("Figure: Eligible benchmark coverage — Overall only", fig_comp)

    doc.build(story)
    print(f"Saved: {pdf_path}")


def _eligibility_key_from_row(doi_val: Any, title_val: Any) -> str:
    doi = _clean_doi(doi_val)
    if doi:
        return f"doi:{doi}"
    t = _safe_str(title_val)
    if t:
        return f"title:{clean_text(t)}"
    return ""


def _load_baseline_eligible_keys(step6_dir: str, fallback_step5_df: pd.DataFrame) -> Tuple[set, str]:
    xlsx_path = os.path.join(step6_dir, BASELINE_ELIGIBILITY_XLSX_BASENAME)
    if os.path.exists(xlsx_path):
        try:
            df = pd.read_excel(xlsx_path, sheet_name=BASELINE_ELIGIBILITY_SHEETNAME)
        except Exception:
            df = pd.read_excel(xlsx_path)

        doi_col = find_col(df, ["doi", "DOI"])
        title_col = find_col(df, ["title", "Title"])
        dec_col = find_col(df, ["final_decision", "Final Decision", "decision"])

        if dec_col:
            elig = df[df[dec_col].apply(_decision_is_include)].copy()
            keys = set()
            for _, r in elig.iterrows():
                k = _eligibility_key_from_row(r.get(doi_col, ""), r.get(title_col, ""))
                if k:
                    keys.add(k)
            return keys, f"baseline_from_step6_excel:{os.path.basename(xlsx_path)}"

    s5_title_col = find_col(fallback_step5_df, ["title", "study"])
    s5_doi_col = find_col(fallback_step5_df, ["doi", "DOI"])
    s5_status_col = find_col(fallback_step5_df, ["final_decision", "decision"])
    if not s5_status_col:
        return set(), "baseline_fallback:missing_status_col"

    passed = fallback_step5_df[fallback_step5_df[s5_status_col].apply(_decision_is_include)].copy()
    keys = set()
    for _, r in passed.iterrows():
        k = _eligibility_key_from_row(r.get(s5_doi_col, ""), r.get(s5_title_col, ""))
        if k:
            keys.add(k)
    return keys, "baseline_fallback:step5_passed"


def _load_abstract_lookup(out_root: str) -> Tuple[Dict[str, str], Dict[str, str], str]:
    step4_path = os.path.join(out_root, "step4", "step4_abstracts.csv")
    if not os.path.exists(step4_path):
        return {}, {}, "abstracts_missing"

    try:
        df = pd.read_csv(step4_path)
    except Exception:
        return {}, {}, f"abstracts_unreadable:{os.path.basename(step4_path)}"

    title_col = find_col(df, ["title", "Title", "study"])
    doi_col = find_col(df, ["doi", "DOI"])
    abs_col = find_col(df, ["abstract", "Abstract", "description", "summary"])

    if not abs_col:
        return {}, {}, f"abstracts_no_abstract_col:{os.path.basename(step4_path)}"

    doi_to_abs: Dict[str, str] = {}
    title_to_abs: Dict[str, str] = {}

    for _, r in df.iterrows():
        abs_txt = _safe_str(r.get(abs_col))
        if not abs_txt:
            continue

        doi = _clean_doi(r.get(doi_col)) if doi_col else ""
        if doi:
            doi_to_abs[doi] = abs_txt

        title = _safe_str(r.get(title_col)) if title_col else ""
        if title:
            title_to_abs[clean_text(title)] = abs_txt

    return doi_to_abs, title_to_abs, f"abstracts_from:{os.path.basename(step4_path)}"


def _load_step1_blocks(out_root: str) -> Tuple[Dict[str, str], str]:
    p = os.path.join(out_root, "step1", STEP1_QUERIES_JSON_BASENAME)
    if not os.path.exists(p):
        return {}, "step1_queries_missing"

    try:
        with open(p, "r", encoding="utf-8") as f:
            rows = json.load(f)
    except Exception:
        return {}, "step1_queries_unreadable"

    blocks: Dict[str, str] = {}
    for r in rows:
        if not isinstance(r, dict):
            continue
        name = r.get("name")
        query = r.get("query")
        if isinstance(name, str) and isinstance(query, str) and name.strip():
            blocks[name.strip()] = query.strip()

    return blocks, f"step1_queries_from:{os.path.basename(p)}"


def _get_search_string(config: dict, out_root: str) -> Tuple[str, str]:
    if config.get("search_string"):
        return str(config["search_string"]).strip(), "config.search_string"

    step1_txt = os.path.join(out_root, "step1", "step1_total_query.txt")
    s = _read_text_file(step1_txt)
    if s:
        return s, f"file:{step1_txt}"

    return "", "missing"


def _tokenize(text: str) -> List[str]:
    t = re.sub(r"[^a-z0-9\s\-]", " ", (text or "").lower())
    toks = []
    for x in t.split():
        if len(x) < 4:
            continue
        toks.append(x.strip("-"))
    return toks


def _extract_author_tokens(df: pd.DataFrame) -> set:
    toks = set()
    for col in ["fetched_authors", "author_names"]:
        if col not in df.columns:
            continue
        for v in df[col].dropna().astype(str).tolist():
            for t in re.findall(r"[a-z]{4,}", v.lower()):
                toks.add(t)
    return toks


def _anchors_from_search_string(search_string: str) -> set:
    if not search_string:
        return set()
    anchors = set(re.findall(r"[a-z]{4,}", search_string.lower()))
    anchors -= STOPWORDS
    anchors -= GENERIC_BAD_SUGGESTIONS
    return anchors


def _route_term_to_block(term: str, step1_blocks: Dict[str, str]) -> str:
    t = (term or "").lower().strip()
    if not t:
        return "UNASSIGNED"

    for block, vocab in ROUTE_TOKENS.items():
        if t in vocab and (block in step1_blocks):
            return block

    for block_name, q in step1_blocks.items():
        if t in (q or "").lower():
            return block_name

    return "UNASSIGNED"

def _search_terms_present(search_string: str) -> set:
    """
    Extract "terms" that appear in the search string, including wildcard stems.
    Example: drought* -> 'drought' (stored as a wildcard stem).
    """
    s = (search_string or "").lower()
    # keep only a-z0-9 and * so we can see wildcarded terms
    s = re.sub(r"[^a-z0-9\*\s]", " ", s)

    out = set()
    for raw in s.split():
        if len(raw) < 3:
            continue
        out.add(raw)
    return out


def _token_covered_by_search(tok: str, ss_terms: set) -> bool:
    """
    True if tok is already covered by the search string, accounting for:
      - exact matches
      - singular/plural
      - wildcard stems like drought* covering droughts
      - common stem variants (e.g., vulnerabilit* covering vulnerability)
    """
    if not tok:
        return False

    t = tok.lower().strip()
    if not t:
        return False

    # 1) exact token present
    if t in ss_terms:
        return True

    # 2) singular/plural present
    if t.endswith("s") and len(t) > 4:
        if t[:-1] in ss_terms:
            return True
    else:
        if (t + "s") in ss_terms:
            return True

    # 3) wildcard stem coverage
    # if search has "stem*" and tok starts with stem -> covered
    for stem_len in (4, 5, 6, 7, 8, 9, 10):
        if len(t) >= stem_len:
            stem = t[:stem_len]
            if (stem + "*") in ss_terms:
                return True

    # 4) handle cases where query uses a slightly different stem:
    # behaviour vs behav* ; vulnerability vs vulnerabilit*
    # Check if any wildcard stem in search is a prefix of tok or vice versa.
    for w in ss_terms:
        if not w.endswith("*"):
            continue
        stem = w[:-1]
        if len(stem) < 4:
            continue
        if t.startswith(stem) or stem.startswith(t[: max(4, min(len(t), len(stem)))]):
            return True

    return False


def _suggest_improvements_by_block(
    search_string: str,
    missed_records: List[Dict[str, str]],
    author_tokens: set,
    step1_blocks: Dict[str, str],
    top_k: int = 18,
) -> Tuple[str, List[str]]:
    ss = (search_string or "").lower()
    ss_terms = _search_terms_present(search_string)
    anchors = _anchors_from_search_string(search_string)

    counter = Counter()
    evidence: Dict[str, Dict[str, int]] = {}

    for rec in missed_records:
        title = (rec.get("title") or "").lower()
        abstract = (rec.get("abstract") or "").lower()

        record_text = f"{title} {abstract}".strip()
        anchor_hit = any(a in record_text for a in anchors) if anchors else True
        anchor_mult = 2.0 if anchor_hit else 1.0

        title_toks = _tokenize(title)
        abs_toks = _tokenize(abstract)

        def _add(tok: str, base_w: float, where: str):
            if not tok:
                return
            if not re.fullmatch(r"[a-z]+", tok):
                return
            if tok in STOPWORDS or tok in GENERIC_BAD_SUGGESTIONS:
                return
            if tok in author_tokens:
                return
            if _token_covered_by_search(tok, ss_terms):
                return

            boost = 1.0
            if tok in ELIGIBILITY_BOOST_TERMS:
                boost *= 2.25
            if anchors and any((tok.startswith(a[:5]) or a.startswith(tok[:5])) for a in anchors if len(a) >= 5 and len(tok) >= 5):
                boost *= 1.35

            w = base_w * anchor_mult * boost
            counter[tok] += w

            if tok not in evidence:
                evidence[tok] = {"title_hits": 0, "abstract_hits": 0}
            if where == "title":
                evidence[tok]["title_hits"] += 1
            else:
                evidence[tok]["abstract_hits"] += 1

        for t in title_toks:
            _add(t, base_w=2.0, where="title")
        for t in abs_toks:
            _add(t, base_w=1.0, where="abstract")

    if not counter:
        return "No eligible, non-noise missed-title/abstract keywords found to suggest additions.", []

    picks = counter.most_common(max(60, top_k * 3))

    by_block: Dict[str, List[dict]] = defaultdict(list)
    unassigned: List[dict] = []

    for tok, score in picks:
        ev = evidence.get(tok, {"title_hits": 0, "abstract_hits": 0})
        routed = _route_term_to_block(tok, step1_blocks)
        entry = {
            "term": tok,
            "score": float(score),
            "title_hits": int(ev.get("title_hits", 0)),
            "abstract_hits": int(ev.get("abstract_hits", 0)),
            "block": routed,
        }
        if routed == "UNASSIGNED":
            unassigned.append(entry)
        else:
            by_block[routed].append(entry)

    per_block_max = 8
    for b in list(by_block.keys()):
        by_block[b] = sorted(by_block[b], key=lambda x: x["score"], reverse=True)[:per_block_max]

    ordered_blocks = [
        "M__measurement_eval",
        "C_concept__adaptation_actions",
        "C_concept__resilience_outcomes",
        "C_concept__capacities_knowledge",
        "C_concept__livelihood_productivity",
        "C_context_climate__hazards",
        "C_context_climate__climate_framing",
        "C_context_agriculture__ag_systems",
        "C_concept__maladaptation",
        "P__smallholder_small_scale",
        "P__households_poverty",
        "P__livestock_pastoral",
        "P__fisheries_aquaculture",
        "P__agroforestry_tree",
        "P__marginalized_groups",
    ]
    for b in step1_blocks.keys():
        if b not in ordered_blocks and not b.endswith("__ALL"):
            ordered_blocks.append(b)

    flat: List[dict] = []
    for b in ordered_blocks:
        flat.extend(by_block.get(b, []))
        if len(flat) >= top_k:
            break
    flat = flat[:top_k]

    capped_by_block: Dict[str, List[dict]] = defaultdict(list)
    capped_unassigned: List[dict] = []
    for e in flat:
        if e["block"] == "UNASSIGNED":
            capped_unassigned.append(e)
        else:
            capped_by_block[e["block"]].append(e)

    seen = set()
    deduped_terms: List[str] = []
    deduped_meta: Dict[str, Dict[str, float]] = {}
    for e in flat:
        t = e["term"]
        if t in seen:
            continue
        seen.add(t)
        deduped_terms.append(t)
        deduped_meta[t] = {
            "score": float(e["score"]),
            "title_hits": int(e["title_hits"]),
            "abstract_hits": int(e["abstract_hits"]),
        }

    lines: List[str] = []
    lines.append("Candidate terms from missed eligible titles + abstracts (filtered + eligibility-weighted)")
    lines.append("and routed to Step1 query blocks.")
    lines.append("")

    def _fmt_block_snippet(terms: List[str]) -> str:
        return " OR ".join([f"{t}*" for t in terms])

    any_block_output = False
    for b in ordered_blocks:
        entries = capped_by_block.get(b, [])
        if not entries:
            continue
        any_block_output = True

        terms = [e["term"] for e in entries]
        lines.append(f"BLOCK: {b}")
        if b in step1_blocks:
            lines.append("  Where to add: inside this block's TITLE-ABS-KEY(...) OR list in step1_queries.json")
        else:
            lines.append("  Where to add: add to the closest concept block")
        lines.append("  Suggested terms:")
        lines.append("    " + ", ".join(terms))
        lines.append("  Paste-ready OR-snippet:")
        lines.append("    " + _fmt_block_snippet(terms))
        lines.append("  Rationale (weighted score; where they occur):")
        for e in entries:
            lines.append(f"    - {e['term']}: {e['score']:.2f} (title_hits={e['title_hits']}, abstract_hits={e['abstract_hits']})")
        lines.append("")

    if not any_block_output:
        lines.append("No routed terms found.")
        lines.append("")

    if capped_unassigned:
        terms = [e["term"] for e in capped_unassigned]
        lines.append("UNASSIGNED (could not confidently map to a block; review manually):")
        lines.append("  " + ", ".join(terms))
        lines.append("")

    if deduped_terms:
        lines.append("FINAL SUMMARY (deduped across all blocks; with title/abstract hit counts):")
        for t in deduped_terms:
            m = deduped_meta.get(t, {})
            lines.append(
                f"  - {t}: title_hits={int(m.get('title_hits',0))}, abstract_hits={int(m.get('abstract_hits',0))}, score={float(m.get('score',0.0)):.2f}"
            )
        lines.append("")
        lines.append("FINAL OR-SNIPPET (wildcarded):")
        lines.append("  " + " OR ".join([f"{t}*" for t in deduped_terms]))
        lines.append("")

    return "\n".join(lines).strip(), deduped_terms


def _stamp_figure(fig, stamp: str) -> None:
    fig.text(0.99, 0.006, stamp, ha="right", va="bottom", fontsize=8, color="black")


def _plot_eligible_overall(df: pd.DataFrame, out_path: str, title: str, timestamp: str) -> None:
    n_total = int(len(df))
    n_in = int(df["in_scopus_retrieval"].sum()) if n_total else 0
    n_missed = n_total - n_in
    pct = (100.0 * n_in / n_total) if n_total else 0.0

    fig, ax = plt.subplots(figsize=(8.8, 5.2))
    ax.bar(["Eligible"], [n_in], label="In Scopus", color=PALETTE_RETRIEVAL[True], alpha=0.9)
    ax.bar(["Eligible"], [n_missed], bottom=[n_in], label="Missed", color=PALETTE_RETRIEVAL[False], alpha=0.9)

    ymax = max(1, int((n_total * 1.35)))
    ax.set_ylim(0, ymax)
    ax.set_ylabel("Count")
    ax.set_title(title, fontweight="bold")

    if n_in > 0:
        ax.text(0, n_in / 2, f"{n_in}", ha="center", va="center", color="white", fontweight="bold")
    if n_missed > 0:
        ax.text(0, n_in + n_missed / 2, f"{n_missed}", ha="center", va="center", color="white", fontweight="bold")
    ax.text(0, n_total + max(0.5, n_total * 0.03), f"{pct:.1f}%", ha="center", va="bottom", fontweight="bold")

    ax.legend(frameon=False, loc="upper right")
    _stamp_figure(fig, f"Generated: {timestamp}")

    plt.tight_layout()
    plt.subplots_adjust(bottom=0.16)
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Saved chart: {out_path}")


def _plot_retrieval_2x2(df: pd.DataFrame, out_path: str, title: str, timestamp: str) -> None:
    plot_configs = [
        {"title": "All Studies", "filter": lambda d: d},
        {"title": "Passed Studies", "filter": lambda d: d[d["simple_status"] == "Passed"]},
        {"title": "Unclear Studies", "filter": lambda d: d[d["simple_status"] == "Unclear"]},
        {"title": "Failed Studies", "filter": lambda d: d[d["simple_status"] == "Failed"]},
    ]

    fig, axes = plt.subplots(2, 2, figsize=(14, 11))
    axes = axes.flatten()

    for i, cfg in enumerate(plot_configs):
        ax = axes[i]
        subset = cfg["filter"](df)
        counts = subset.groupby(["source_type", "in_scopus_retrieval"]).size().unstack(fill_value=0)
        counts = counts.reindex(CATS_SOURCE).fillna(0)
        if True not in counts.columns:
            counts[True] = 0
        if False not in counts.columns:
            counts[False] = 0

        p1 = ax.bar(counts.index, counts[True], label="In Scopus", color=PALETTE_RETRIEVAL[True], alpha=0.9)
        p2 = ax.bar(counts.index, counts[False], bottom=counts[True], label="Missed", color=PALETTE_RETRIEVAL[False], alpha=0.9)

        max_height = (counts[True] + counts[False]).max()
        if max_height > 0:
            ax.set_ylim(0, max_height * 1.20)

        ax.set_title(f"{cfg['title']} (n={len(subset)})", fontweight="bold")

        for bar_group in (p1, p2):
            for bar in bar_group:
                h = bar.get_height()
                if h > 0:
                    ax.text(
                        bar.get_x() + bar.get_width() / 2.0,
                        bar.get_y() + h / 2.0,
                        f"{int(h)}",
                        ha="center",
                        va="center",
                        color="white",
                        fontweight="bold",
                    )

        totals = counts[True] + counts[False]
        for idx, cat in enumerate(counts.index):
            total = totals.loc[cat]
            found = counts.loc[cat, True]
            if total > 0:
                pct = (found / total) * 100
                ax.text(
                    idx,
                    total + (max_height * 0.03 if max_height else 0.5),
                    f"{pct:.1f}%",
                    ha="center",
                    va="bottom",
                    color="black",
                    fontweight="bold",
                    fontsize=10,
                )

    fig.suptitle(title, fontsize=16, fontweight="bold")
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", bbox_to_anchor=(0.5, 0.94), ncol=2, frameon=False)

    _stamp_figure(fig, f"Generated: {timestamp}")

    plt.tight_layout()
    plt.subplots_adjust(top=0.88, bottom=0.12)
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Saved chart: {out_path}")


def _plot_coverage_overall_only(base_stats: dict, out_path: str, timestamp: str) -> None:
    base_pct = float(base_stats.get("pct_overall", 0.0))

    fig, ax = plt.subplots(figsize=(8.8, 5.2))

    ax.bar([0], [base_pct], 0.6, label="Eligible (baseline)", edgecolor="black")
    ax.set_ylim(0, 115)
    ax.set_ylabel("Coverage in Scopus (%)")
    ax.set_xticks([0])
    ax.set_xticklabels(["Overall"])

    ax.text(0, base_pct + 2, f"{base_pct:.1f}%", ha="center", va="bottom", fontsize=10, fontweight="bold")

    fig.suptitle("Eligible Benchmark Coverage in Scopus", fontsize=14, fontweight="bold")
    sub = f"Eligible: {base_stats['n_in_overall']}/{base_stats['n_total_overall']}"
    ax.text(0.5, 1.02, sub, transform=ax.transAxes, ha="center", va="bottom", fontsize=10)

    ax.legend(frameon=False, loc="upper right")
    _stamp_figure(fig, f"Generated: {timestamp}")

    plt.tight_layout(rect=[0, 0.08, 1, 0.92])
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Saved chart: {out_path}")

from matplotlib.ticker import MaxNLocator

def _plot_progress_over_iterations(evo_rows: List[dict], out_path: str, timestamp: str) -> None:
    if not evo_rows:
        return

    points = []
    for r in evo_rows:
        it = int(r.get("iteration") or 0)
        cov_final = (r.get("coverage", {}) or {}).get("final_target", {}) or {}

        pct = float(cov_final.get("pct_overall", 0.0) or 0.0)
        sc_total = int(r.get("scopus_total_returned", 0) or 0)

        n_in = int(cov_final.get("n_in_overall", 0) or 0)
        n_total = int(cov_final.get("n_total_overall", 0) or 0)
        points.append((it, pct, sc_total, n_in, n_total))

    points.sort(key=lambda x: x[0])
    xs = [p[0] for p in points]
    ys = [p[1] for p in points]

    fig, ax = plt.subplots(figsize=(9.5, 4.5))
    ax.plot(xs, ys, marker="o")

    ax.set_ylim(0, 120)
    ax.set_xlabel("Iteration")
    ax.set_ylabel("Final Target Coverage (%)")
    ax.set_title("Scopus Coverage Progress Over Iterations", fontweight="bold")

    # integer iterations only
    ax.xaxis.set_major_locator(MaxNLocator(integer=True))
    ax.set_xticks(xs)

    # give headroom for labels above points
    ax.margins(y=0.14)

    for it, pct, sc_total, n_in, n_total in points:
        ax.annotate(
            f"Scopus: {sc_total:,}",
            xy=(it, pct),
            xytext=(0, 8),
            textcoords="offset points",
            ha="center",
            va="bottom",
            fontsize=8,
            color="blue",
        )

        ax.annotate(
            f"Benchmark: {n_in}/{n_total}",
            xy=(it, pct),
            xytext=(0, -6),
            textcoords="offset points",
            ha="center",
            va="top",
            fontsize=8,
            color="black",
        )

        # 2) percent above that (black, larger/bold)
        ax.annotate(
            f"{pct:.1f}%",
            xy=(it, pct),
            xytext=(0, 20),
            textcoords="offset points",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
            color="black",
        )

    _stamp_figure(fig, f"Generated: {timestamp}")
    plt.tight_layout()
    plt.subplots_adjust(bottom=0.18)
    plt.savefig(out_path, dpi=300)
    plt.close()

def _merge_scopus_metadata(step_df: pd.DataFrame, scopus_df: pd.DataFrame) -> Tuple[pd.DataFrame, set, set]:
    sc_title_col = find_col(scopus_df, ["Title", "Article Title"])
    sc_doi_col = find_col(scopus_df, ["DOI", "doi"])
    if not sc_title_col:
        raise RuntimeError("Scopus file missing a title column (expected Title / Article Title).")

    scopus_df = scopus_df.copy()
    scopus_df["clean_title"] = scopus_df[sc_title_col].apply(clean_text)
    scopus_titles = set(scopus_df["clean_title"].dropna().unique())
    scopus_titles.discard("")

    scopus_dois = set()
    if sc_doi_col:
        scopus_df["clean_doi"] = scopus_df[sc_doi_col].apply(clean_text)
        scopus_dois = set(scopus_df["clean_doi"].dropna().unique())
        scopus_dois.discard("")

    sc_author_col = find_col(scopus_df, ["Authors", "Author", "author_names", "Creator"])
    sc_journal_col = find_col(scopus_df, ["Source title", "publicationName", "Journal"])
    sc_date_col = find_col(scopus_df, ["Year", "coverDate", "Date"])
    sc_cit_col = find_col(scopus_df, ["citation_raw", "Citation Raw", "citation", "Citation"])

    needed_cols = [sc_title_col]
    if sc_author_col:
        needed_cols.append(sc_author_col)
    if sc_journal_col:
        needed_cols.append(sc_journal_col)
    if sc_date_col:
        needed_cols.append(sc_date_col)
    if sc_doi_col:
        needed_cols.append(sc_doi_col)
    if sc_cit_col:
        needed_cols.append(sc_cit_col)

    scopus_meta = scopus_df[needed_cols].dropna(subset=[sc_title_col]).copy()
    rename_map = {sc_title_col: "title_scopus"}
    if sc_author_col:
        rename_map[sc_author_col] = "author_names"
    if sc_journal_col:
        rename_map[sc_journal_col] = "publicationName"
    if sc_date_col:
        rename_map[sc_date_col] = "year_scopus"
    if sc_doi_col:
        rename_map[sc_doi_col] = "doi_scopus"
    if sc_cit_col:
        rename_map[sc_cit_col] = "citation_raw_scopus"
    scopus_meta = scopus_meta.rename(columns=rename_map)
    scopus_meta["match_title"] = scopus_meta["title_scopus"].apply(clean_text)

    out = step_df.merge(scopus_meta, how="left", on="match_title", suffixes=("", "_scopus_merge"))
    return out, scopus_titles, scopus_dois


def _compute_in_scopus(row: pd.Series, scopus_titles: set, scopus_dois: set) -> bool:
    md = _safe_str(row.get("match_doi"))
    mt = _safe_str(row.get("match_title"))
    return (bool(md) and md in scopus_dois) or (bool(mt) and mt in scopus_titles)


def _compute_coverage_stats(df: pd.DataFrame) -> dict:
    def _pct(n_in: int, n_total: int) -> float:
        return (100.0 * n_in / n_total) if n_total > 0 else 0.0

    out = {}
    out["n_total_overall"] = int(len(df))
    out["n_in_overall"] = int(df["in_scopus_retrieval"].sum())
    out["pct_overall"] = _pct(out["n_in_overall"], out["n_total_overall"])

    peer = df[df["source_type"] == "Peer Reviewed"]
    grey = df[df["source_type"] == "Grey Literature"]

    out["n_total_peer"] = int(len(peer))
    out["n_in_peer"] = int(peer["in_scopus_retrieval"].sum())
    out["pct_peer"] = _pct(out["n_in_peer"], out["n_total_peer"])

    out["n_total_grey"] = int(len(grey))
    out["n_in_grey"] = int(grey["in_scopus_retrieval"].sum())
    out["pct_grey"] = _pct(out["n_in_grey"], out["n_total_grey"])
    return out


def _first_author_lastname(author_str: str) -> str:
    s = _safe_str(author_str)
    if not s:
        return "Unknown"
    chunk = re.split(r"\s*;\s*|\s+and\s+|,?\s+et\s+al\.?", s, flags=re.IGNORECASE)[0].strip()
    if not chunk:
        return "Unknown"
    if "," in chunk:
        return chunk.split(",")[0].strip() or "Unknown"
    parts = [p for p in chunk.split() if p.strip()]
    return (parts[-1].strip() if parts else "Unknown") or "Unknown"


def _author_year_label_from_row(r: pd.Series) -> str:
    author_src = (
        _safe_str(r.get("fetched_authors"))
        or _safe_str(r.get("author_names"))
        or _safe_str(r.get("authors"))
    )

    year_raw = (
        _safe_str(r.get("fetched_year"))
        or _safe_str(r.get("year_scopus"))
        or _safe_str(r.get("year"))
    )

    if (not author_src) or (not re.search(r"\d{4}", year_raw or "")):
        cit = _safe_str(r.get("citation"))
        if cit:
            m = re.match(r"^\s*(.+?)\s*\((\d{4}|n\.d\.)\)", cit)
            if m:
                author_src = author_src or m.group(1).strip()
                year_raw = year_raw or m.group(2).strip()

    last = _first_author_lastname(author_src)
    y_match = re.search(r"\d{4}", year_raw or "")
    year = y_match.group(0) if y_match else ("n.d." if "n.d" in (year_raw or "").lower() else "n.d.")
    return f"{last} ({year})"

def _benchmark_key_from_row(r: pd.Series) -> str:
    d = _clean_doi(r.get("doi"))
    if d:
        return f"doi:{d}"
    t = _safe_str(r.get("title"))
    return f"title:{clean_text(t)}" if t else ""

def _build_benchmark_iteration_summary(
    step7_dir: str,
    history_dir: str,
    evo_rows: List[dict],
    eligible_csv: str,
) -> Tuple[str, str]:
    out_csv = os.path.join(step7_dir, "step7_benchmark_iteration_matrix.csv")
    out_png = os.path.join(step7_dir, "step7_benchmark_iteration_heatmap.png")

    if not eligible_csv or not os.path.exists(eligible_csv):
        return out_csv, out_png

    try:
        elig_df = pd.read_csv(eligible_csv)
    except Exception:
        return out_csv, out_png

    if "title" not in elig_df.columns:
        return out_csv, out_png
    if "doi" not in elig_df.columns:
        elig_df["doi"] = ""
    if "citation" not in elig_df.columns:
        elig_df["citation"] = ""
    if "in_scopus_retrieval" not in elig_df.columns:
        elig_df["in_scopus_retrieval"] = False

    elig_df = elig_df.copy()
    elig_df["bench_key"] = elig_df.apply(_benchmark_key_from_row, axis=1)
    elig_df = elig_df[elig_df["bench_key"].astype(str) != ""].copy()
    if elig_df.empty:
        return out_csv, out_png

    elig_df["author_year"] = elig_df.apply(_author_year_label_from_row, axis=1)

    display = elig_df.drop_duplicates("bench_key").set_index("bench_key")

    keys_unique: List[str] = []
    seen = set()
    for k in elig_df["bench_key"].astype(str).tolist():
        if k and k not in seen:
            seen.add(k)
            keys_unique.append(k)

    iter_to_pct: Dict[int, float] = {}
    for r in (evo_rows or []):
        it = int(r.get("iteration", 0) or 0)
        cov_final = (r.get("coverage", {}) or {}).get("final_target", {}) or {}
        iter_to_pct[it] = float(cov_final.get("pct_overall", 0.0))

    hist_files = sorted(glob.glob(os.path.join(history_dir, "iter*_full_benchmark_status.csv")))
    per_iter: List[Tuple[int, str]] = []
    for fp in hist_files:
        m = re.search(r"iter(\d{3})_full_benchmark_status\.csv$", os.path.basename(fp))
        if m:
            per_iter.append((int(m.group(1)), fp))
    if not per_iter:
        return out_csv, out_png

    iter_maps: Dict[int, Dict[str, bool]] = {}
    for it, fp in per_iter:
        try:
            df = pd.read_csv(fp)
        except Exception:
            continue
        if "in_scopus_retrieval" not in df.columns:
            continue
        if "title" not in df.columns:
            continue
        if "doi" not in df.columns:
            df["doi"] = ""

        df = df.copy()
        df["bench_key"] = df.apply(_benchmark_key_from_row, axis=1)
        df = df[df["bench_key"].astype(str) != ""].copy()
        if df.empty:
            continue

        mp = df.set_index("bench_key")["in_scopus_retrieval"].astype(bool).to_dict()
        iter_maps[it] = mp

    cols_sorted = sorted(iter_maps.keys())
    if not cols_sorted:
        return out_csv, out_png

    latest_it = max(cols_sorted)
    latest_map = iter_maps.get(latest_it, {})

    def _sort_key(k: str) -> Tuple[int, str, str]:
        passed = 1 if latest_map.get(k, False) else 0
        ay = _safe_str(display.loc[k].get("author_year")) if k in display.index else ""
        title = _safe_str(display.loc[k].get("title")) if k in display.index else ""
        return (-passed, ay.lower(), title.lower())

    keys_sorted = sorted(keys_unique, key=_sort_key)

    col_labels: List[str] = []
    for it in cols_sorted:
        pct = iter_to_pct.get(it, None)
        col_labels.append(f"iter{it:03d} ({pct:.1f}%)" if pct is not None else f"iter{it:03d}")

    matrix = pd.DataFrame(index=keys_sorted)
    matrix["#"] = list(range(1, len(keys_sorted) + 1))
    matrix["Author (Year)"] = [
        (_safe_str(display.loc[k].get("author_year")) if k in display.index else "")
        for k in keys_sorted
    ]
    matrix["Title"] = [
        (_safe_str(display.loc[k].get("title")) if k in display.index else "")
        for k in keys_sorted
    ]
    matrix["DOI"] = [
        (_safe_str(display.loc[k].get("doi")) if k in display.index else "")
        for k in keys_sorted
    ]

    for it, lab in zip(cols_sorted, col_labels):
        mp = iter_maps.get(it, {})
        matrix[lab] = ["Pass" if mp.get(k, False) else "Miss" for k in keys_sorted]

    matrix.to_csv(out_csv, index=False, encoding="utf-8")
    print(f"Saved: {out_csv}")

    heat_vals = np.zeros((len(keys_sorted), len(cols_sorted)), dtype=float)
    for j, it in enumerate(cols_sorted):
        mp = iter_maps.get(it, {})
        for i, k in enumerate(keys_sorted):
            heat_vals[i, j] = 1.0 if mp.get(k, False) else 0.0

    n_rows = len(keys_sorted)
    n_cols = len(cols_sorted)

    # --- FIGURE SIZE: force landscape (wider than tall) ---
    # Width grows with number of iteration columns (labels on top).
    # Height grows with number of rows, but is capped more aggressively so it doesn't become "portrait".
    fig_w = max(8.5, min(22.0, 2.0 + 1.25 * n_cols))
    fig_h = max(5.5, min(14.0, 3.0 + 0.12 * n_rows))

    # Enforce landscape ratio (at least 1.35x wider than tall)
    if fig_w < fig_h * 1.35:
        fig_w = min(22.0, fig_h * 1.35)

    cmap = ListedColormap([PALETTE_RETRIEVAL[False], PALETTE_RETRIEVAL[True]])

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.imshow(heat_vals, aspect="auto", cmap=cmap, vmin=0, vmax=1)

    ax.set_xticks(np.arange(n_cols))
    ax.set_xticklabels(col_labels, rotation=45, ha="left", fontsize=9)

    # put x labels on top
    ax.xaxis.tick_top()
    ax.tick_params(axis="x", which="both", top=True, labeltop=True, bottom=False, labelbottom=False, pad=6)


    y_labels = [
        f"{int(matrix.iloc[i]['#']):03d}  {matrix.iloc[i]['Author (Year)']}".strip()
        for i in range(n_rows)
    ]
    show_every = max(1, n_rows // 80)
    yticks = np.arange(0, n_rows, show_every)
    ax.set_yticks(yticks)
    ax.set_yticklabels([y_labels[i] for i in yticks], fontsize=7)

    # wrapped title so it doesn't get cut off
    import textwrap
    title_txt = "Eligible study retrieval by iteration (Pass=In Scopus; Miss=Not retrieved)"
    title_wrapped = textwrap.fill(title_txt, width=52 if n_cols == 1 else 70)
    # use figure-level title so it never gets clipped by tight_layout
    fig.suptitle(title_wrapped, fontweight="bold", y=0.985)

    # reserve headroom for suptitle + top ticks
    plt.tight_layout(rect=[0.02, 0.02, 0.98, 0.955])

    plt.savefig(out_png, dpi=300)
    plt.close()
    print(f"Saved: {out_png}")

    return out_csv, out_png


def run(config: dict) -> None:
    print("--- Running Step 7: Eligible Benchmark vs. Scopus Retrieval ---")

    out_root = config.get("out_dir") or "scripts/outputs"
    step7_dir = os.path.join(out_root, "step7")
    _ensure_dir(step7_dir)
    step6_dir = os.path.join(out_root, "step6")

    history_dir = os.path.join(step7_dir, "history")
    _ensure_dir(history_dir)

    scopus_file = os.path.join(out_root, "step2", "step2_total_records.csv")
    step5_file = os.path.join(out_root, "step5", "step5_eligibility_wide.csv")
    bm_file = _resolve_benchmark_csv(out_root, config.get("benchmark_csv"))

    missing = []
    for p in [scopus_file, step5_file, bm_file]:
        if not p or not os.path.exists(p):
            missing.append(p)
    if missing:
        raise RuntimeError("Missing one or more input files:\n" + "\n".join([f"- {m}" for m in missing]))

    timestamp = _now_local_str()

    search_string, search_source = _get_search_string(config, out_root)
    search_hash = _sha1(search_string or "")[:12]

    doi_to_abs, title_to_abs, abs_source = _load_abstract_lookup(out_root)
    step1_blocks, step1_blocks_source = _load_step1_blocks(out_root)

    evolution_jsonl = os.path.join(step7_dir, "step7_searchstring_evolution.jsonl")
    evolution_csv = os.path.join(step7_dir, "step7_searchstring_evolution.csv")
    appendix_docx = os.path.join(step7_dir, "step7_appendix_searchstring_evolution.docx")
    appendix_pdf = os.path.join(step7_dir, "step7_appendix_searchstring_evolution.pdf")
    suggestions_txt = os.path.join(step7_dir, "step7_searchstring_suggestions.txt")
    progress_png = os.path.join(step7_dir, "step7_progress_over_iterations.png")

    previous = _load_jsonl(evolution_jsonl)

    baseline_xlsx = os.path.join(step6_dir, BASELINE_ELIGIBILITY_XLSX_BASENAME)
    step1_queries_json = os.path.join(out_root, "step1", STEP1_QUERIES_JSON_BASENAME)
    step1_plot_src = os.path.join(out_root, "step1", STEP1_PLOT_BASENAME)  # <-- move/insert here
    step1_summary_path = os.path.join(out_root, "step1", STEP1_SUMMARY_BASENAME)

    run_fingerprint_src = "|".join([
        f"search_hash:{search_hash}",
        _file_fingerprint(scopus_file),
        _file_fingerprint(step5_file),
        _file_fingerprint(bm_file),
        _file_fingerprint(baseline_xlsx),
        _file_fingerprint(step1_queries_json),
        _file_fingerprint(step1_plot_src),
        _file_fingerprint(step1_summary_path),
    ])

    run_fingerprint = _sha1(run_fingerprint_src)[:12]

    deduped = False
    if previous and previous[-1].get("run_fingerprint") == run_fingerprint:
        deduped = True
        iteration = int(previous[-1].get("iteration", len(previous)))
        print(f"Detected identical rerun (run_fingerprint={run_fingerprint}). Not creating a new iteration.")
    else:
        iteration = len(previous) + 1

    iter_tag = f"iter{iteration:03d}"
    step1_plot_src = os.path.join(out_root, "step1", STEP1_PLOT_BASENAME)
    step1_plot_arch = os.path.join(history_dir, f"{iter_tag}_step1_hits_plot.png")
    step1_summary_arch = os.path.join(history_dir, f"{iter_tag}_step1_summary.json")

    scopus_df = pd.read_csv(scopus_file)     # still needed for matching
    step5_df = pd.read_csv(step5_file)

    step1_total_all, step1_summary_dict, step1_summary_source = _load_step1_summary(out_root)

    # Use Step1 as authoritative "total returned" for plotting/logging
    if step1_total_all > 0:
        scopus_total_returned = int(step1_total_all)
        scopus_total_source = step1_summary_source
    else:
        # fallback only if step1 summary missing/bad
        scopus_total_returned = int(len(scopus_df))
        scopus_total_source = "fallback:len(step2_total_records.csv)"

    # Ensure notes exists (Step5 -> Step7)
    if "notes" not in step5_df.columns:
        step5_df["notes"] = ""
    step5_df["notes"] = step5_df["notes"].fillna("").astype(str)

    bm_df = pd.read_csv(bm_file)

    s5_title_col = find_col(step5_df, ["title", "study"])
    s5_doi_col = find_col(step5_df, ["doi", "DOI"])
    s5_status_col = find_col(step5_df, ["final_decision", "decision"])
    if not s5_title_col or not s5_status_col:
        raise RuntimeError("Step 5 missing required columns (need title and final_decision).")

    step5_df = step5_df.copy()
    step5_df["title"] = step5_df[s5_title_col].apply(_safe_str)
    step5_df["doi"] = step5_df[s5_doi_col].apply(_clean_doi) if s5_doi_col else ""
    step5_df["match_title"] = step5_df["title"].apply(clean_text)
    step5_df["match_doi"] = step5_df["doi"].apply(clean_text) if s5_doi_col else ""
    step5_df["simple_status"] = step5_df[s5_status_col].apply(_simple_status)

    # Reorder: notes immediately after the final decision column
    if s5_status_col and s5_status_col in step5_df.columns and "notes" in step5_df.columns:
        cols = list(step5_df.columns)
        cols.remove("notes")
        fi = cols.index(s5_status_col)
        cols.insert(fi + 1, "notes")
        step5_df = step5_df[cols]

    bm_title_col = find_col(bm_df, ["Study", "Title", "title"])
    bm_type_col = find_col(bm_df, ["Type", "Source Type", "source_type"])
    if bm_title_col and bm_type_col:
        bm_df = bm_df.copy()
        bm_df["match_title_bm"] = bm_df[bm_title_col].apply(clean_text)
        type_lookup = bm_df.drop_duplicates("match_title_bm").set_index("match_title_bm")[bm_type_col].to_dict()
        step5_df["restored_type"] = step5_df["match_title"].map(type_lookup).fillna("Unknown")
    else:
        step5_df["restored_type"] = "Unknown"
    step5_df["source_type"] = step5_df["restored_type"].apply(_map_source_type)

    step5_df, scopus_titles, scopus_dois = _merge_scopus_metadata(step5_df, scopus_df)
    
    # Prefer citation_raw coming from Scopus export (Step2)
    if "citation_raw" not in step5_df.columns:
        step5_df["citation_raw"] = ""

    step5_df["citation_raw"] = step5_df["citation_raw"].fillna("").astype(str)

    if "citation_raw_scopus" in step5_df.columns:
        step5_df["citation_raw_scopus"] = step5_df["citation_raw_scopus"].fillna("").astype(str)
        # fill only when missing
        step5_df.loc[step5_df["citation_raw"].str.strip() == "", "citation_raw"] = (
            step5_df.loc[step5_df["citation_raw"].str.strip() == "", "citation_raw_scopus"]
        )

    # Optional: remove hard line breaks that can make CSV *look* truncated in viewers
    step5_df["citation_raw"] = step5_df["citation_raw"].str.replace(r"[\r\n]+", " ", regex=True).str.strip()


    if "author_names" not in step5_df.columns:
        step5_df["author_names"] = ""
    step5_df["in_scopus_retrieval"] = step5_df.apply(lambda r: _compute_in_scopus(r, scopus_titles, scopus_dois), axis=1)

    baseline_keys, baseline_source = _load_baseline_eligible_keys(step6_dir, step5_df)

    def _row_key(r: pd.Series) -> str:
        d = _clean_doi(r.get("doi"))
        if d:
            return f"doi:{d}"
        return f"title:{clean_text(_safe_str(r.get('title')))}"

    step5_df["eligible_baseline"] = step5_df.apply(lambda r: _row_key(r) in baseline_keys, axis=1)

    print("--- repairing missing metadata for Eligible benchmark records (Crossref) ---")
    has_scopus_author = step5_df["author_names"].notna() & (step5_df["author_names"] != "")
    has_doi = step5_df["doi"].astype(str).str.strip() != ""
    mask_needs_repair = (
        (step5_df["eligible_baseline"])
        & has_doi
        & ((step5_df["in_scopus_retrieval"] == False) | (~has_scopus_author))
    )

    to_repair = step5_df[mask_needs_repair].index.tolist()
    print(f"Fetching metadata for {len(to_repair)} records...")
    for idx in tqdm(to_repair, total=len(to_repair), unit="rec"):
        doi = step5_df.at[idx, "doi"]
        meta = fetch_crossref_metadata(doi)
        if meta:
            step5_df.at[idx, "fetched_authors"] = meta.get("fetched_authors", "")
            step5_df.at[idx, "fetched_journal"] = meta.get("fetched_journal", "")
            step5_df.at[idx, "fetched_year"] = meta.get("fetched_year", "")
        time.sleep(0.1)

    # Use citation_raw if present; fallback to build_citation
    step5_df["citation"] = step5_df["citation_raw"].fillna("").astype(str).str.strip()
    mask_missing_cit = step5_df["citation"].str.strip() == ""
    if mask_missing_cit.any():
        step5_df.loc[mask_missing_cit, "citation"] = step5_df.loc[mask_missing_cit].apply(
            lambda r: build_citation(r.to_dict()), axis=1
        )

    full_csv = os.path.join(step7_dir, "step7_full_benchmark_status.csv")
    step5_df.to_csv(full_csv, index=False, encoding="utf-8")
    print(f"Saved: {full_csv}")

    full_status_arch = os.path.join(history_dir, f"{iter_tag}_full_benchmark_status.csv")
    if not deduped:
        _archive_copy(full_csv, full_status_arch)

    benchmark_plot = os.path.join(step7_dir, "step7_retrieval_analysis_all_2x2.png")
    _plot_retrieval_2x2(step5_df, benchmark_plot, "Scopus Retrieval Analysis — Full Benchmark List (All Decisions)", timestamp)

    base_elig_df = step5_df[step5_df["eligible_baseline"]].copy()
    base_missed_df = base_elig_df[base_elig_df["in_scopus_retrieval"] == False].copy()

    out_base_all_csv = os.path.join(step7_dir, "step7_all_eligible_studies_baseline.csv")
    out_base_missed_csv = os.path.join(step7_dir, "step7_missed_eligible_studies_baseline.csv")
    base_elig_df.to_csv(out_base_all_csv, index=False, encoding="utf-8")
    base_missed_df.to_csv(out_base_missed_csv, index=False, encoding="utf-8")
    print(f"Saved: {out_base_all_csv}")
    print(f"Saved: {out_base_missed_csv}")

    if HAS_DOCX:
        export_to_word_citations(
            base_missed_df,
            "Eligible Benchmark (Baseline): Missed in Scopus",
            os.path.join(step7_dir, "step7_missed_eligible_studies_baseline.docx"),
        )
        export_to_word_citations(
            base_elig_df,
            "Eligible Benchmark (Baseline): Full List",
            os.path.join(step7_dir, "step7_all_eligible_studies_baseline.docx"),
        )

    base_overall_plot = ""
    if not base_elig_df.empty:
        base_overall_plot = os.path.join(step7_dir, "step7_retrieval_analysis_eligible_2x2.png")
        _plot_eligible_overall(base_elig_df, base_overall_plot, "Eligible Benchmark (Baseline): In Scopus vs Missed (Overall)", timestamp)

    base_stats = _compute_coverage_stats(base_elig_df) if not base_elig_df.empty else {
        "n_total_overall": 0, "n_in_overall": 0, "pct_overall": 0.0,
        "n_total_peer": 0, "n_in_peer": 0, "pct_peer": 0.0,
        "n_total_grey": 0, "n_in_grey": 0, "pct_grey": 0.0,
    }

    comparison_plot = os.path.join(step7_dir, "step7a_retrieval_analysis_comparison.png")
    _plot_coverage_overall_only(base_stats, comparison_plot, timestamp)

    final_target_df = base_elig_df
    final_target_label = "baseline"

    final_missed_df = final_target_df[final_target_df["in_scopus_retrieval"] == False].copy()
    final_missed_csv = os.path.join(step7_dir, "step7_benchmark_scopus_missed.csv")
    final_missed_df.to_csv(final_missed_csv, index=False, encoding="utf-8")
    print(f"Saved: {final_missed_csv}")

    if HAS_DOCX and final_target_label != "baseline":
        export_to_word_citations(
            final_missed_df,
            f"Final Eligible Benchmark ({final_target_label}): Missed in Scopus",
            os.path.join(step7_dir, "step7_benchmark_scopus_missed.docx"),
        )

    missed_records: List[Dict[str, str]] = []
    for _, r in final_missed_df.iterrows():
        title = _safe_str(r.get("title"))
        doi = _clean_doi(r.get("doi"))
        abs_txt = ""
        if doi and doi in doi_to_abs:
            abs_txt = doi_to_abs.get(doi, "")
        if not abs_txt and title:
            abs_txt = title_to_abs.get(clean_text(title), "")
        missed_records.append({"title": title, "abstract": abs_txt})

    author_tokens = _extract_author_tokens(final_missed_df)
    suggestion_text, deduped_terms = _suggest_improvements_by_block(
        search_string,
        missed_records,
        author_tokens=author_tokens,
        step1_blocks=step1_blocks,
        top_k=18,
    )

    with open(suggestions_txt, "w", encoding="utf-8") as f:
        f.write(f"(Abstract source: {abs_source})\n")
        f.write(f"(Step1 blocks source: {step1_blocks_source})\n")
        f.write(suggestion_text + "\n")
    print(f"Saved: {suggestions_txt}")

    final_stats = _compute_coverage_stats(final_target_df) if not final_target_df.empty else {
        "n_total_overall": 0, "n_in_overall": 0, "pct_overall": 0.0,
        "n_total_peer": 0, "n_in_peer": 0, "pct_peer": 0.0,
        "n_total_grey": 0, "n_in_grey": 0, "pct_grey": 0.0,
    }

    coverage_csv = os.path.join(step7_dir, "step7_benchmark_scopus_coverage.csv")
    coverage_row = {
        "timestamp_local": timestamp,
        "iteration": iteration,
        "run_fingerprint": run_fingerprint,
        "search_source": search_source,
        "search_hash": search_hash,
        "baseline_eligibility_source": baseline_source,
        "final_target": final_target_label,
        "scopus_total_returned": scopus_total_returned,
        "baseline_n_total": base_stats["n_total_overall"],
        "baseline_n_in_scopus": base_stats["n_in_overall"],
        "baseline_pct": round(base_stats["pct_overall"], 3),
        "final_n_total": final_stats["n_total_overall"],
        "final_n_in_scopus": final_stats["n_in_overall"],
        "final_pct": round(final_stats["pct_overall"], 3),
        "final_missed_n": int(len(final_missed_df)),
        "abstract_source": abs_source,
        "step1_blocks_source": step1_blocks_source,
        "suggested_terms_deduped": ", ".join(deduped_terms),
        "scopus_total_source": scopus_total_source,
        "step1_total_all": int(step1_total_all or 0),
        "step1_query_signature": _safe_str(step1_summary_dict.get("query_signature")),
        "step1_summary_timestamp_utc": _safe_str(step1_summary_dict.get("timestamp_utc")),
    }
    pd.DataFrame([coverage_row]).to_csv(coverage_csv, index=False, encoding="utf-8")
    print(f"Saved: {coverage_csv}")

    fig_benchmark_arch = os.path.join(history_dir, f"{iter_tag}_benchmark_all_2x2.png")
    fig_comp_arch = os.path.join(history_dir, f"{iter_tag}_coverage_overall.png")
    fig_base_arch = os.path.join(history_dir, f"{iter_tag}_baseline_eligible_overall.png")

    if not deduped:
        _archive_copy(benchmark_plot, fig_benchmark_arch)
        _archive_copy(comparison_plot, fig_comp_arch)
        _archive_copy(step1_plot_src, step1_plot_arch)
        if base_overall_plot:
            _archive_copy(base_overall_plot, fig_base_arch)
        _archive_copy(step1_summary_path, step1_summary_arch)


        evolution_rec = {
            "timestamp_local": timestamp,
            "iteration": iteration,
            "run_fingerprint": run_fingerprint,
            "search_source": search_source,
            "search_string": search_string,
            "search_hash": search_hash,
            "scopus_total_returned": scopus_total_returned,
            "scopus_total_source": scopus_total_source,
            "step1_summary": {
                "total_all": int(step1_total_all or 0),
                "query_signature": _safe_str(step1_summary_dict.get("query_signature")),
                "timestamp_utc": _safe_str(step1_summary_dict.get("timestamp_utc")),
                "archived_path": step1_summary_arch if os.path.exists(step1_summary_arch) else "",
            },
            "inputs": {
                "benchmark_csv": bm_file,
                "scopus_csv": scopus_file,
                "step5_csv": step5_file,
                "abstract_source": abs_source,
                "step1_blocks_source": step1_blocks_source,
            },
            "eligibility_sources": {"baseline": baseline_source, "final_target": final_target_label},
            "coverage": {"baseline": base_stats, "final_target": final_stats, "final_missed_n": int(len(final_missed_df))},
            "missed_titles_sample": [m.get("title", "") for m in missed_records[:10]],
            "suggestion_text": suggestion_text,
            "figures": {
                "benchmark_all_plot": fig_benchmark_arch if os.path.exists(fig_benchmark_arch) else "",
                "coverage_comparison": fig_comp_arch if os.path.exists(fig_comp_arch) else "",
                "baseline_eligible_plot": fig_base_arch if os.path.exists(fig_base_arch) else "",
                "step1_hits_plot": step1_plot_arch if os.path.exists(step1_plot_arch) else "",  # <-- add
            },

        }
        _write_jsonl_append(evolution_jsonl, evolution_rec)
        print(f"Appended evolution log: {evolution_jsonl}")

    evo_rows = _load_jsonl(evolution_jsonl)

    # After evo_rows = _load_jsonl(evolution_jsonl) OR right before reloading for plotting
    if deduped and os.path.exists(evolution_jsonl):
        evo_rows_tmp = _load_jsonl(evolution_jsonl)
        if evo_rows_tmp:
            evo_rows_tmp[-1]["scopus_total_returned"] = scopus_total_returned
            evo_rows_tmp[-1]["scopus_total_source"] = scopus_total_source

            # optionally keep step1 summary fields in sync too:
            evo_rows_tmp[-1]["step1_summary"] = {
                "total_all": int(step1_total_all or 0),
                "query_signature": _safe_str(step1_summary_dict.get("query_signature")),
                "timestamp_utc": _safe_str(step1_summary_dict.get("timestamp_utc")),
                "archived_path": "",
            }

            evo_rows_tmp[-1]["coverage"]["final_target"] = final_stats
            evo_rows_tmp[-1]["coverage"]["baseline"] = base_stats
            evo_rows_tmp[-1]["coverage"]["final_missed_n"] = int(len(final_missed_df))
            _rewrite_jsonl(evolution_jsonl, evo_rows_tmp)

            # IMPORTANT: plot from updated rows
            evo_rows = evo_rows_tmp

    flat = []
    for r in evo_rows:
        cov_final = (r.get("coverage", {}) or {}).get("final_target", {}) or {}
        flat.append({
            "iteration": r.get("iteration"),
            "timestamp_local": r.get("timestamp_local"),
            "search_hash": r.get("search_hash"),
            "final_target": (r.get("eligibility_sources", {}) or {}).get("final_target"),
            "final_n_total": cov_final.get("n_total_overall", 0),
            "final_n_in_scopus": cov_final.get("n_in_overall", 0),
            "final_pct": cov_final.get("pct_overall", 0.0),
        })
    pd.DataFrame(flat).to_csv(evolution_csv, index=False, encoding="utf-8")
    print(f"Updated evolution CSV: {evolution_csv}")

    _plot_progress_over_iterations(evo_rows, progress_png, timestamp)
    print(f"Saved: {progress_png}")

    if HAS_DOCX:

        _, heatmap_png = _build_benchmark_iteration_summary(
            step7_dir,
            history_dir,
            evo_rows,
            os.path.join(step7_dir, "step7_all_eligible_studies_baseline.csv"),
        )
        _build_appendix_from_evolution(appendix_docx, evo_rows, progress_png, heatmap_png)

        # Build PDF directly (cross-platform, no Word/LibreOffice)
        _build_appendix_pdf_from_evolution(appendix_pdf, evo_rows, progress_png, heatmap_png, timestamp)
        if not deduped:
            appendix_pdf_arch = os.path.join(history_dir, f"{iter_tag}_appendix_searchstring_evolution.pdf")
            _archive_copy(appendix_pdf, appendix_pdf_arch)


    print("--- SUMMARY ---")
    print(f"Benchmark CSV used:        {bm_file}")
    print(f"Baseline eligible source:  {baseline_source}")
    print(f"Final target:             {final_target_label}")
    print(f"Final coverage:           {final_stats['n_in_overall']}/{final_stats['n_total_overall']} ({final_stats['pct_overall']:.1f}%)")
    print(f"Final missed:             {len(final_missed_df)}")
    print(f"Abstract source:          {abs_source}")
    print(f"Step1 blocks source:      {step1_blocks_source}")
    print(f"Appendix:                 {appendix_docx if HAS_DOCX else '[docx not available]'}")
    if deduped:
        print("NOTE: This run was deduped (no new iteration appended).")


if __name__ == "__main__":
    run({"out_dir": "scripts/outputs"})
