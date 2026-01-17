#!/usr/bin/env python3
"""
step3_benchmark.py

Step 3 (benchmark prep + DOI enrichment):
  - Load benchmark CSV (config.benchmark_csv)
  - Extract a clean title from the raw "Study" text
  - Parse DOIs from text (best-effort)
  - If DOI missing, enrich via Crossref/OpenAlex/SemanticScholar (cached)
  - STRICT MODE: Any DOI flagged as "Needs Review" is discarded automatically.

MINIMAL UPDATE:
  - Add safe `year` extraction (from citation_raw) and carry `year` forward to step3 output.
  - Use year (when available) to improve DOI enrichment scoring, and backfill year from enrichment only if missing.

Outputs:
  - outputs/step3/step3_benchmark_list.csv
  - outputs/step3/step3_benchmark_list.enriched_only.csv
  - outputs/step3/step3_benchmark_list.summary.json
  - outputs/step3/step3_benchmark_doi_enrichment_cache.json
  - outputs/step3/step3_benchmark_list.docx
  - outputs/step3/step3_benchmark_list.bib
  - outputs/step3/step3_benchmark_list.ris
"""

from __future__ import annotations

import hashlib
import os
import re
import time
from difflib import SequenceMatcher
from typing import Optional

import pandas as pd
import requests
from tqdm import tqdm

from docx import Document
from docx.shared import Pt

import config as cfg
from utils import load_json, save_json, doi_from_text, normalize_doi, request_with_retries

# --- CONFIGURATION & THRESHOLDS ---
YEAR_SPLIT_REGEX = re.compile(r"[\s\.,]\(\d{4}[a-z]?\)[\.,]?\s*")

# --- NEW: safe year extraction ---
YEAR_PARENS_RE = re.compile(r"\((\d{4})[a-z]?\)")
YEAR_ANY_RE = re.compile(r"\b(19|20)\d{2}\b")


def _extract_year_any(text: str) -> Optional[int]:
    if not isinstance(text, str):
        return None
    m = YEAR_PARENS_RE.search(text)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    m2 = YEAR_ANY_RE.search(text)
    if m2:
        try:
            return int(m2.group(0))
        except Exception:
            return None
    return None


CROSSREF_WORKS_URL = "https://api.crossref.org/works"
OPENALEX_WORKS_URL = "https://api.openalex.org/works"
SEMANTIC_SCHOLAR_SEARCH_URL = "https://api.semanticscholar.org/graph/v1/paper/search"

DOI_AUTO_ACCEPT_SCORE = 0.90     # Auto-accept if >= 90%
DOI_MIN_CANDIDATE_SCORE = 0.60   # Candidate threshold


def _find_col(df: pd.DataFrame, candidates: list[str]) -> Optional[str]:
    norm = {c.strip().lower(): c for c in df.columns}
    for k in candidates:
        if k.strip().lower() in norm:
            return norm[k.strip().lower()]
    for c in df.columns:
        cl = c.strip().lower()
        for k in candidates:
            kl = k.strip().lower()
            if kl and kl in cl:
                return c
    return None


def _extract_clean_title(full_text: str) -> str:
    if not isinstance(full_text, str):
        return ""
    parts = YEAR_SPLIT_REGEX.split(full_text, maxsplit=1)
    cleaned = parts[1].strip() if len(parts) > 1 and len(parts[1].strip()) > 10 else full_text.strip()
    cleaned = re.sub(r"^[\W_]+", "", cleaned)
    return cleaned.strip()


def _normalize_doi_url(x: str) -> str:
    """
    Returns canonical DOI URL, stripping common trailing punctuation that breaks APIs.
    """
    d = normalize_doi(x)
    d = d.strip().rstrip(" .),;]}>")
    return f"https://doi.org/{d}" if d else ""


def _doi_key(x: str) -> str:
    """
    Returns DOI key (no https prefix) suitable for record_key when DOI exists.
    """
    d = normalize_doi(x)
    return d.strip().rstrip(" .),;]}>") if d else ""


def _record_key(clean_title: str, doi_url: Optional[str]) -> str:
    """
    Stable key:
      - DOI key if DOI exists
      - else a safe, compact title key
    """
    if isinstance(doi_url, str) and doi_url.strip():
        dk = _doi_key(doi_url)
        if dk:
            return dk

    t = (clean_title or "").strip().lower()
    t = re.sub(r"[^a-z0-9]+", "", t)
    return t[:200] if t else ""


def _title_norm(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"[^\w\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _title_score(a: str, b: str) -> float:
    a2, b2 = _title_norm(a), _title_norm(b)
    if not a2 or not b2:
        return 0.0
    base = SequenceMatcher(None, a2, b2).ratio()
    if len(a2) > 15 and len(b2) > 15:
        if a2 in b2 or b2 in a2:
            return 1.0
    return base


def _http_get_json(
    session: requests.Session,
    url: str,
    *,
    headers: dict | None = None,
    params: dict | None = None,
) -> Optional[dict]:
    try:
        data, _ = request_with_retries(
            session=session,
            method="GET",
            url=url,
            headers=headers or {},
            params=params or {},
            data=None,
            tries=6,
        )
        return data
    except Exception:
        return None


def _crossref_candidates(session: requests.Session, title: str, mailto: str | None = None, rows: int = 5) -> list[dict]:
    params = {"query.bibliographic": title, "rows": rows}
    if mailto:
        params["mailto"] = mailto
    data = _http_get_json(session, CROSSREF_WORKS_URL, params=params)
    items = ((data or {}).get("message") or {}).get("items") or []
    out: list[dict] = []
    for it in items:
        t = it.get("title")[0] if it.get("title") else None
        year = it.get("issued", {}).get("date-parts", [[None]])[0][0]
        doi = it.get("DOI")
        if doi:
            out.append({"source": "crossref", "doi": _normalize_doi_url(doi), "title": t, "year": year})
    return out


def _openalex_candidates(session: requests.Session, title: str, mailto: str | None = None, per_page: int = 5) -> list[dict]:
    params = {"search": title, "per-page": per_page}
    if mailto:
        params["mailto"] = mailto
    data = _http_get_json(session, OPENALEX_WORKS_URL, params=params)
    results = (data or {}).get("results") or []
    out: list[dict] = []
    for it in results:
        doi = it.get("doi")
        if doi:
            out.append(
                {
                    "source": "openalex",
                    "doi": _normalize_doi_url(doi),
                    "title": it.get("title"),
                    "year": it.get("publication_year"),
                }
            )
    return out


def _semantic_candidates(session: requests.Session, title: str, contact_email: str | None = None, limit: int = 5) -> list[dict]:
    params = {"query": title, "limit": limit, "fields": "title,year,externalIds"}
    headers = {"Accept": "application/json", "User-Agent": f"step3-enrichment ({contact_email or 'no-email'})"}
    data = _http_get_json(session, SEMANTIC_SCHOLAR_SEARCH_URL, headers=headers, params=params)
    results = (data or {}).get("data") or []
    out: list[dict] = []
    for it in results:
        doi = (it.get("externalIds") or {}).get("DOI")
        if doi:
            out.append(
                {
                    "source": "semantic_scholar",
                    "doi": _normalize_doi_url(doi),
                    "title": it.get("title"),
                    "year": it.get("year"),
                }
            )
    return out


def _best_doi_for_title(
    session: requests.Session,
    cleaned_title: str,
    year: int | None,
    contact_email: str | None,
    mailto: str | None,
) -> dict:
    cands: list[dict] = []
    cands += _crossref_candidates(session, cleaned_title, mailto=mailto)
    cands += _openalex_candidates(session, cleaned_title, mailto=mailto)
    cands += _semantic_candidates(session, cleaned_title, contact_email=contact_email)

    doi_counts: dict[str, int] = {}
    for c in cands:
        d = c.get("doi")
        if d:
            doi_counts[d] = doi_counts.get(d, 0) + 1

    scored: list[dict] = []
    for c in cands:
        if not c.get("doi") or not c.get("title"):
            continue
        base_score = _title_score(cleaned_title, str(c.get("title") or ""))
        if doi_counts.get(c["doi"], 0) > 1:
            base_score += 0.05
        if year and c.get("year"):
            try:
                if abs(int(year) - int(c["year"])) > 1:
                    base_score -= 0.20
            except Exception:
                pass
        scored.append({**c, "score": min(1.0, float(base_score))})

    scored.sort(key=lambda x: x["score"], reverse=True)
    best = scored[0] if scored and scored[0]["score"] >= DOI_MIN_CANDIDATE_SCORE else None

    if not best:
        return {
            "doi": None,
            "doi_source": None,
            "doi_score": 0.0,
            "matched_title": None,
            "needs_review": False,
            "year": None,            # NEW
            "year_source": None,     # NEW
        }

    return {
        "doi": best["doi"],
        "doi_source": best["source"],
        "doi_score": float(best["score"]),
        "matched_title": best["title"],
        "needs_review": bool(best["score"] < DOI_AUTO_ACCEPT_SCORE),
        "year": best.get("year"),           # NEW
        "year_source": best.get("source"),  # NEW
    }


def _cache_key_for_title(title: str) -> str:
    return hashlib.sha256(_title_norm(title).encode("utf-8")).hexdigest()


def _bib_escape(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace("\\", "\\\\")
    s = s.replace("{", "\\{").replace("}", "\\}")
    s = s.replace("\n", " ").replace("\r", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _ensure_unique_keys(keys: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    out: list[str] = []
    for k in keys:
        k0 = (k or "").strip()
        if not k0:
            k0 = "record"
        if k0 not in seen:
            seen[k0] = 1
            out.append(k0)
        else:
            seen[k0] += 1
            out.append(f"{k0}_{seen[k0]}")
    return out


def _ris_type_from_row(row: dict) -> str:
    t = str(row.get("type") or "").strip().lower()
    if "grey" in t or "report" in t or t == "grey":
        return "RPRT"
    if "peer" in t or "journal" in t:
        return "JOUR"
    return "GEN"


def _write_ris(df: pd.DataFrame, out_path: str) -> None:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    keys = _ensure_unique_keys(df["record_key"].fillna("").astype(str).tolist())

    lines: list[str] = []
    for i, (_, r) in enumerate(df.iterrows()):
        title = str(r.get("title") or "").strip()
        doi_url = str(r.get("doi") or "").strip()
        doi_key = _doi_key(doi_url) if doi_url else ""
        rec_key = keys[i]
        ty = _ris_type_from_row(r.to_dict())

        lines.append(f"TY  - {ty}")
        if rec_key:
            lines.append(f"ID  - {rec_key}")
        if title:
            lines.append(f"TI  - {title}")
        if doi_key:
            lines.append(f"DO  - {doi_key}")
        if doi_url:
            lines.append(f"UR  - {doi_url}")

        note_parts = []
        tval = str(r.get("type") or "").strip()
        if tval:
            note_parts.append(f"type={tval}")
        iv = str(r.get("identified_via") or "").strip()
        if iv:
            note_parts.append(f"identified_via={iv}")
        if note_parts:
            lines.append(f"N1  - {'; '.join(note_parts)}")

        lines.append("ER  -")
        lines.append("")

    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(lines))


def _write_bib(df: pd.DataFrame, out_path: str) -> None:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    keys = _ensure_unique_keys(df["record_key"].fillna("").astype(str).tolist())

    entries: list[str] = []
    for i, (_, r) in enumerate(df.iterrows()):
        title = _bib_escape(r.get("title") or "")
        doi_url = str(r.get("doi") or "").strip()
        doi_key = _bib_escape(_doi_key(doi_url)) if doi_url else ""
        url = _bib_escape(doi_url) if doi_url else ""
        rec_key = _bib_escape(keys[i])

        entry_type = "misc"
        t = str(r.get("type") or "").strip().lower()
        if "peer" in t or "journal" in t:
            entry_type = "article"
        elif "grey" in t or "report" in t:
            entry_type = "techreport"

        fields: list[str] = []
        if title:
            fields.append(f"  title = {{{title}}}")
        if doi_key:
            fields.append(f"  doi = {{{doi_key}}}")
        if url:
            fields.append(f"  url = {{{url}}}")

        note_parts = []
        tval = str(r.get("type") or "").strip()
        if tval:
            note_parts.append(f"type={tval}")
        iv = str(r.get("identified_via") or "").strip()
        if iv:
            note_parts.append(f"identified_via={iv}")
        if note_parts:
            fields.append(f"  note = {{{_bib_escape('; '.join(note_parts))}}}")

        body = ",\n".join(fields)
        entries.append(f"@{entry_type}{{{rec_key},\n{body}\n}}")

    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n\n".join(entries) + "\n")


def _write_docx(df: pd.DataFrame, out_path: str, title: str) -> None:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Records: {len(df)}").font.size = Pt(10)

    if df.empty:
        doc.add_paragraph("No records.")
        doc.save(out_path)
        return

    for _, r in df.iterrows():
        t = str(r.get("title") or "").strip()
        doi_url = str(r.get("doi") or "").strip()
        pieces = []
        if t:
            pieces.append(t.rstrip(". "))
        if doi_url:
            pieces.append(doi_url)
        line = ". ".join(pieces).strip()
        if not line:
            continue
        p = doc.add_paragraph(line, style="List Number")
        p.paragraph_format.space_after = Pt(6)

    doc.save(out_path)


def step3_build_benchmark_list(config: dict) -> dict:
    out_dir = config.get("out_dir") or getattr(cfg, "out_dir", "outputs")
    benchmark_csv = config.get("benchmark_csv") or getattr(cfg, "benchmark_csv", None)
    if not benchmark_csv:
        raise SystemExit("Missing benchmark_csv configuration.")

    step3_dir = os.path.join(out_dir, "step3")
    out_csv = os.path.join(step3_dir, "step3_benchmark_list.csv")
    out_enriched_only_csv = os.path.join(step3_dir, "step3_benchmark_list.enriched_only.csv")
    out_summary = os.path.join(step3_dir, "step3_benchmark_list.summary.json")
    cache_path = os.path.join(step3_dir, "step3_benchmark_doi_enrichment_cache.json")

    out_docx = os.path.join(step3_dir, "step3_benchmark_list.docx")
    out_bib = os.path.join(step3_dir, "step3_benchmark_list.bib")
    out_ris = os.path.join(step3_dir, "step3_benchmark_list.ris")

    os.makedirs(step3_dir, exist_ok=True)

    print(f"[step3] Reading benchmark: {benchmark_csv}")
    bench = pd.read_csv(benchmark_csv)

    study_col = _find_col(bench, ["Study", "study", "Title", "title"]) or bench.columns[0]
    type_col = _find_col(bench, ["Type", "type", "Source Type", "source type"])
    idvia_col = _find_col(bench, ["Identified Via", "identified_via", "identified via"])

    raw_studies = bench[study_col].fillna("").astype(str)
    cleaned_titles = raw_studies.apply(_extract_clean_title)

    parsed_doi = raw_studies.apply(doi_from_text)
    initial_doi_mask = parsed_doi.notna()
    initial_doi_count = int(initial_doi_mask.sum())

    # NEW: parse year from citation raw (Study text)
    parsed_year = raw_studies.apply(_extract_year_any)

    out = pd.DataFrame(
        {
            "citation_raw": raw_studies,
            "title": cleaned_titles,
            "type": bench[type_col] if type_col else "",
            "doi": parsed_doi,
            "identified_via": bench[idvia_col] if idvia_col else "",
            "year": parsed_year,                   # NEW
            "year_source": "parsed_from_citation", # NEW (default)
        }
    )

    contact_email = os.getenv("CONTACT_EMAIL")
    cache = load_json(cache_path, default={})

    print(f"[step3] Processing {len(out)} rows (auto-accept >= {DOI_AUTO_ACCEPT_SCORE*100:.0f}%)")
    results_meta: list[dict] = []

    with requests.Session() as session:
        pbar = tqdm(total=len(out), desc="[step3] Enriching DOIs", unit="row")

        for i, row in out.iterrows():
            existing = row.get("doi")
            if isinstance(existing, str) and existing.strip():
                out.at[i, "doi"] = _normalize_doi_url(existing)
                results_meta.append({"source": "given", "score": 1.0, "review": False, "match": None})
                pbar.update(1)
                continue

            clean_t = row["title"]
            if not clean_t or len(clean_t) < 5:
                results_meta.append({"source": None, "score": 0.0, "review": False, "match": None})
                pbar.update(1)
                continue

            # NEW: pass year (if available) to improve matching
            row_year = row.get("year")
            try:
                row_year_int = int(row_year) if str(row_year).strip() else None
            except Exception:
                row_year_int = None

            ck = _cache_key_for_title(clean_t)
            if ck not in cache:
                got = _best_doi_for_title(
                    session,
                    clean_t,
                    year=row_year_int,  # NEW
                    contact_email=contact_email,
                    mailto=contact_email,
                )
                cache[ck] = got
                if i % 25 == 0:
                    save_json(cache_path, cache)

            got = cache[ck]
            if got.get("doi"):
                out.at[i, "doi"] = _normalize_doi_url(got["doi"])

            # NEW: backfill year only if missing
            cur_y = out.at[i, "year"]
            missing_year = (cur_y is None) or (str(cur_y).strip() == "") or (str(cur_y).lower() == "nan")
            if missing_year and got.get("year"):
                out.at[i, "year"] = got.get("year")
                out.at[i, "year_source"] = got.get("year_source") or got.get("doi_source") or "enriched"

            results_meta.append(
                {
                    "source": got.get("doi_source"),
                    "score": float(got.get("doi_score") or 0.0),
                    "review": bool(got.get("needs_review")),
                    "match": got.get("matched_title"),
                }
            )
            pbar.update(1)
            time.sleep(0.05)

        pbar.close()

    save_json(cache_path, cache)

    out["doi_source"] = [x["source"] for x in results_meta]
    out["doi_score"] = [x["score"] for x in results_meta]
    out["doi_needs_review"] = [x["review"] for x in results_meta]
    out["doi_matched_title"] = [x["match"] for x in results_meta]

    # Strict mode: reject low-confidence DOI enrichments
    reject_mask = out["doi_needs_review"] == True
    rejected_n = int(reject_mask.sum())
    if rejected_n:
        print(f"[step3] Strict mode: rejecting {rejected_n} low-confidence DOI enrichments")
        out.loc[reject_mask, "doi"] = None
        out.loc[reject_mask, "doi_source"] = "rejected_low_score"

    out["record_key"] = out.apply(lambda r: _record_key(r.get("title", ""), r.get("doi")), axis=1)

    # DOI flow flags
    out["_has_doi_final"] = out["doi"].notna()
    out["_has_doi_initial"] = initial_doi_mask.values
    out["_doi_given"] = out["doi_source"] == "given"
    out["_doi_enriched"] = out["_has_doi_final"] & (out["doi_source"] != "given")
    out["_doi_rejected"] = out["doi_source"] == "rejected_low_score"

    # NEW: include year in exported schema
    final_cols = [
        "citation_raw",
        "record_key",
        "title",
        "type",
        "doi",
        "identified_via",
        "year",         # NEW
        "year_source",  # NEW
        "doi_source",
        "doi_score",
        "doi_needs_review",
        "doi_matched_title",
    ]

    out[final_cols].to_csv(out_csv, index=False)

    enriched_only = out[(out["doi_source"].fillna("") != "given") & (out["doi"].notna())].copy()
    enriched_only[final_cols].to_csv(out_enriched_only_csv, index=False)

    type_counts = out["type"].value_counts(dropna=False).to_dict()
    identified_via_counts = out["identified_via"].value_counts(dropna=False).to_dict()

    doi_by_type_deep = (
        out.groupby("type", dropna=False)
        .agg(
            total=("title", "count"),
            initial_with_doi=("_has_doi_initial", "sum"),
            doi_given=("_doi_given", "sum"),
            doi_enriched=("_doi_enriched", "sum"),
            doi_rejected=("_doi_rejected", "sum"),
            final_with_doi=("_has_doi_final", "sum"),
            final_missing=("doi", lambda x: x.isna().sum()),
        )
        .reset_index()
        .to_dict(orient="records")
    )

    doi_by_identified_via_deep = (
        out.groupby("identified_via", dropna=False)
        .agg(
            total=("title", "count"),
            initial_with_doi=("_has_doi_initial", "sum"),
            doi_given=("_doi_given", "sum"),
            doi_enriched=("_doi_enriched", "sum"),
            doi_rejected=("_doi_rejected", "sum"),
            final_with_doi=("_has_doi_final", "sum"),
            final_missing=("doi", lambda x: x.isna().sum()),
        )
        .reset_index()
        .to_dict(orient="records")
    )

    identified_via_doi_coverage = (
        out.assign(has_final_doi=out["doi"].notna())
        .groupby("identified_via", dropna=False)["has_final_doi"]
        .agg(total="count", with_doi="sum")
        .reset_index()
        .assign(pct_with_doi=lambda d: (d["with_doi"] / d["total"]).round(3))
        .to_dict(orient="records")
    )

    summary = {
        "total": len(out),
        "initial_with_doi": initial_doi_count,
        "doi_enriched": int(len(enriched_only)),
        "doi_needs_review_but_rejected": int(rejected_n),
        "missing": int(out["doi"].isna().sum()),
        "by_type": type_counts,
        "by_identified_via": identified_via_counts,
        "identified_via_doi_coverage": identified_via_doi_coverage,
        "doi_flow_by_type": doi_by_type_deep,
        "doi_flow_by_identified_via": doi_by_identified_via_deep,
    }
    save_json(out_summary, summary)

    print(f"[step3] Writing exports: {out_docx}, {out_bib}, {out_ris}")
    _write_docx(out[final_cols], out_docx, "Benchmark List (Step 3)")
    _write_bib(out[final_cols], out_bib)
    _write_ris(out[final_cols], out_ris)

    print(f"[step3] Wrote: {out_csv}")
    return {"status": "ok", "path": out_csv}


def run(config: dict) -> dict:
    return step3_build_benchmark_list(config)


def run_step3(config: dict) -> dict:
    return step3_build_benchmark_list(config)


def main(config: dict | None = None) -> dict:
    return step3_build_benchmark_list(config or {})


if __name__ == "__main__":
    step3_build_benchmark_list({})
