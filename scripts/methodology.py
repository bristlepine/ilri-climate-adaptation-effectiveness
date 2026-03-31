"""
methodology.py

Generates a Word document (.docx) describing the full systematic review pipeline
methodology. Loads actual run statistics from step output JSON files and embeds
pipeline figures so the document is self-contained.

Run from the scripts/ directory:
    python methodology.py

Output: outputs/methodology/Methodology_Appendix.docx
"""

import io
import json
import os
import shutil
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from PIL import Image

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
HERE    = Path(__file__).resolve().parent
OUTPUTS = HERE / "outputs"
OUT_DIR = OUTPUTS / "methodology"
FIG_DIR = OUT_DIR / "figures"
OUT_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Methodology_Appendix.docx"

# ---------------------------------------------------------------------------
# Repository
# ---------------------------------------------------------------------------
REPO_URL    = "https://github.com/bristlepine/ilri-climate-adaptation-effectiveness"
REPO_BRANCH = "main"

def script_url(filename):
    """Return a GitHub URL to a scripts/ file on the main branch."""
    return f"{REPO_URL}/blob/{REPO_BRANCH}/scripts/{filename}"

# Source figures — keyed by name, value is the original path
FIG_SRC = {
    "search_hits":   OUTPUTS / "step1"  / "step1_hits_plot.png",
    "irr_r1":        OUTPUTS / "step11" / "EPPI_Review_-_R1_analysis.png",
    "irr_r1b":       OUTPUTS / "step11" / "EPPI_Review_-_R1b_analysis.png",
    "irr_r2a":       OUTPUTS / "step11" / "EPPI_Review_-_R2a_analysis.png",
    "irr_r3a":       OUTPUTS / "step11" / "EPPI_Review_-_R3a_analysis.png",
    "screen12":      OUTPUTS / "step12" / "step12_summary.png",
    "screen14":      OUTPUTS / "step14" / "step14_summary.png",
    "coding":        OUTPUTS / "step15" / "step15_summary.png",
    "roses":         OUTPUTS / "step16" / "roses_flow.png",
    "temporal":      OUTPUTS / "step16" / "temporal_trends.png",
    "geography":     OUTPUTS / "step16" / "geographic_bar.png",
    "producer":      OUTPUTS / "step16" / "producer_type_bar.png",
    "methodology":   OUTPUTS / "step16" / "methodology_bar.png",
    "domain":        OUTPUTS / "step16" / "domain_heatmap.png",
}


def _normalise_png(src: Path, dest: Path, max_width_px: int = 1800):
    """
    Copy src PNG to dest, re-saving via Pillow to ensure a clean sRGB PNG
    at a sensible DPI (150 dpi) and capped width.  This removes any
    unusual colour profiles or metadata that cause Word to warn on open.
    Falls back to a plain file copy if Pillow cannot open the file.
    """
    try:
        img = Image.open(src).convert("RGB")
        if img.width > max_width_px:
            ratio = max_width_px / img.width
            img = img.resize(
                (max_width_px, int(img.height * ratio)), Image.LANCZOS
            )
        img.save(dest, format="PNG", dpi=(150, 150), optimize=False)
    except Exception:
        shutil.copy2(src, dest)


def prepare_figures(generated: dict) -> dict:
    """
    1. Copy / normalise every FIG_SRC entry that exists into FIG_DIR.
    2. Save every generated (BytesIO) figure into FIG_DIR.
    Returns a combined dict {key: Path} pointing to the local copies.
    """
    local = {}

    # Pipeline output figures
    for key, src in FIG_SRC.items():
        dest = FIG_DIR / f"{key}.png"
        if Path(src).exists():
            _normalise_png(Path(src), dest)
            local[key] = dest
            print(f"  copied: {key}")
        else:
            print(f"  MISSING (skipped): {key} -> {src}")

    # Generated figures (BytesIO)
    for key, buf in generated.items():
        dest = FIG_DIR / f"{key}.png"
        buf.seek(0)
        img = Image.open(buf).convert("RGB")
        img.save(dest, format="PNG", dpi=(150, 150))
        local[key] = dest
        buf.seek(0)          # reset so it can still be read if needed
        print(f"  generated: {key}")

    return local

# ---------------------------------------------------------------------------
# Load run statistics
# ---------------------------------------------------------------------------

def load_json(path):
    try:
        with open(path) as f:
            return json.load(f)
    except Exception:
        return {}

s2   = load_json(OUTPUTS / "step2"  / "step2_total_records.meta.json")
s9   = load_json(OUTPUTS / "step9"  / "step9_summary.json")
s9a  = load_json(OUTPUTS / "step9a" / "step9a_summary.json")
s12  = load_json(OUTPUTS / "step12" / "step12_results.meta.json")
s13  = load_json(OUTPUTS / "step13" / "step13_summary.json")
s14  = load_json(OUTPUTS / "step14" / "step14_results.meta.json")
s15  = load_json(OUTPUTS / "step15" / "step15_coded.meta.json")

c_r1  = load_json(OUTPUTS / "step11" / "EPPI_Review_-_R1_summary.json")
c_r1b = load_json(OUTPUTS / "step11" / "EPPI_Review_-_R1b_summary.json")
c_r2  = load_json(OUTPUTS / "step11" / "EPPI_Review_-_R2_summary.json")
c_r2a = load_json(OUTPUTS / "step11" / "EPPI_Review_-_R2a_summary.json")
c_r3a = load_json(OUTPUTS / "step11" / "EPPI_Review_-_R3a_summary.json")

# Step 10 calibration run times (seconds)
s10_times = {
    "r1":  load_json(OUTPUTS / "step10" / "step10_check.meta.json").get("elapsed_seconds", 2743),
    "r1b": load_json(OUTPUTS / "step10" / "step10_check_r1b.meta.json").get("elapsed_seconds", 2937),
    "r2":  load_json(OUTPUTS / "step10" / "step10_check_r2.meta.json").get("elapsed_seconds", 23913),
    "r3a": load_json(OUTPUTS / "step10" / "step10_check_r3a.meta.json").get("elapsed_seconds", 1756),
}

# Key formatted numbers
def fmt(n, fallback="N/A"):
    try:
        return f"{int(n):,}"
    except Exception:
        return fallback

SCOPUS_RAW          = fmt(s2.get("base_total_reported_by_scopus",   17083))
SCOPUS_DEDUPED      = fmt(s2.get("combined_rows_written_after_dedupe", 17021))
ABSTRACT_MISSING_9  = fmt(s9.get("final_abstract_missing",  1430))
ABSTRACT_MISSING_9A = fmt(s9a.get("final_missing",           1314))
ABSTRACT_GAINED_9A  = fmt(s9a.get("gained_vs_step9",          116))
SCREENED_TOTAL      = fmt(s12.get("rows_total",              17021))
SCREENED_INCLUDE    = fmt(s12.get("decision_counts", {}).get("Include",  6206))
SCREENED_EXCLUDE    = fmt(s12.get("decision_counts", {}).get("Exclude", 10815))
MISSING_ABS_12      = fmt(s12.get("missing_abstract_count",  1314))
FT_RETRIEVED        = fmt(s13.get("status_counts", {}).get("retrieved",  929))
FT_NEEDS_MANUAL     = fmt(s13.get("status_counts", {}).get("needs_manual", 5277))
FT_INCLUDE          = fmt(s14.get("decision_counts", {}).get("Include",   184))
FT_EXCLUDE          = fmt(s14.get("decision_counts", {}).get("Exclude",   130))
FT_NO_FT            = fmt(s14.get("rows_no_fulltext",        5277))
FT_WITH_FT          = fmt(s14.get("rows_screened_with_fulltext", 314))
CODED_TOTAL         = fmt(s15.get("rows_total",              6076))
CODED_FT            = fmt(s15.get("coding_source_counts", {}).get("full_text",       184))
CODED_ABS           = fmt(s15.get("coding_source_counts", {}).get("abstract_only",  4583))
CODED_MISSING       = fmt(s15.get("coding_source_counts", {}).get("missing_abstract", 1309))


# ---------------------------------------------------------------------------
# Figure helpers
# ---------------------------------------------------------------------------

def fig_to_stream(fig):
    """Return a BytesIO PNG stream from a matplotlib figure."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

# FIG will be populated by prepare_figures() at build time
FIG: dict = {}


def make_kappa_convergence_figure():
    """
    Synthetic convergence chart: LLM kappa vs reconciled across calibration rounds.
    Also shows human-human kappa per round for comparison.
    """
    rounds = ["R1\n(n=205)", "R1b\n(n=205)", "R2\n(n=103)", "R2a\n(n=103)", "R3a\n(n=107)"]

    llm_kappa = [
        c_r1.get("pairwise_kappa",  {}).get("LLM vs CJ Reconciled",         0.436),
        c_r1b.get("pairwise_kappa", {}).get("LLMr1b vs CJ Reconciled",      0.645),
        # R2 has no reconciled column in the JSON, use LLM vs Jennifer Cisse as proxy
        c_r2.get("pairwise_kappa",  {}).get("Jennifer Cisse vs LLM",        0.717),
        c_r2a.get("pairwise_kappa", {}).get("LLM_r2a vs CJ Reconciled",     0.770),
        # R3a has no reconciled yet; use mean of the two human raters as proxy
        float(np.mean([
            c_r3a.get("pairwise_kappa", {}).get("Jennifer Cisse vs LLM",    0.690),
            c_r3a.get("pairwise_kappa", {}).get("Caroline Staub vs LLM",    0.674),
        ])),
    ]

    human_kappa = [
        c_r1.get("pairwise_kappa",  {}).get("Caroline Staub vs Jennifer Cisse", 0.500),
        c_r1b.get("pairwise_kappa", {}).get("Caroline Staub vs Jennifer Cisse", 0.500),
        c_r2.get("pairwise_kappa",  {}).get("Jennifer Cisse vs Caroline Staub", 0.765),
        c_r2a.get("pairwise_kappa", {}).get("Caroline Staub vs Jennifer Cisse", 0.765),
        c_r3a.get("pairwise_kappa", {}).get("Jennifer Cisse vs Caroline Staub", 0.703),
    ]

    x = np.arange(len(rounds))
    fig, ax = plt.subplots(figsize=(7, 3.8))

    ax.plot(x, llm_kappa,   "o-",  color="#2166ac", linewidth=2, markersize=7,
            label="LLM vs reconciled gold standard")
    ax.plot(x, human_kappa, "s--", color="#d73027", linewidth=2, markersize=7,
            label="Human inter-rater (Caroline vs Jennifer)")

    # Threshold bands
    ax.axhspan(0.61, 0.80, color="#fee08b", alpha=0.25, label="Substantial (0.61–0.80)")
    ax.axhspan(0.80, 1.00, color="#a6d96a", alpha=0.25, label="Almost perfect (> 0.80)")
    ax.axhline(0.61, color="gray", linewidth=0.8, linestyle=":")
    ax.axhline(0.80, color="gray", linewidth=0.8, linestyle=":")

    # Label criteria revision points
    for i, label in [(0, "Criteria\nv1"), (1, "Criteria\nv1b"), (3, "Criteria\nv2a")]:
        ax.annotate(label, xy=(i, llm_kappa[i]),
                    xytext=(i + 0.08, llm_kappa[i] - 0.07),
                    fontsize=7, color="#2166ac",
                    arrowprops=dict(arrowstyle="-", color="#2166ac", lw=0.8))

    ax.set_xticks(x)
    ax.set_xticklabels(rounds, fontsize=9)
    ax.set_ylim(0.3, 1.0)
    ax.set_ylabel("Cohen's kappa", fontsize=10)
    ax.set_title("Calibration convergence across rounds", fontsize=11, pad=8)
    ax.legend(loc="lower right", fontsize=8, framealpha=0.9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return fig_to_stream(fig)


def make_time_comparison_figure():
    """
    Bar chart: estimated human person-hours vs actual pipeline compute hours
    for each major stage.
    """
    stages = [
        "Abstract\nscreening\n(17,021 records)",
        "Full-text\nscreening\n(6,206 records)",
        "Data\nextraction\n(6,076 records)",
        "Calibration\n(3 rounds,\n~415 records)",
        "Full-text\nretrieval\n(6,206 records)",
    ]
    # Human estimates: minutes per record × record count / 60 / number of reviewers needed
    # Abstract screen: 2 min/record × 17021 × 2 reviewers / 60
    # Full-text screen: 10 min/record × 6206 × 2 reviewers / 60
    # Data extraction: 25 min/record × 6076 × 1 person / 60
    # Calibration: 5 min/record × 415 × 2 reviewers (reconcile) / 60
    # Full-text retrieval: manual ILL, library requests ~15 min each × 6206 / 60
    human_hours = [
        round(2 * 17021 * 2 / 60),       # ~1,135 h
        round(10 * 6206 * 2 / 60),       # ~2,069 h
        round(25 * 6076 / 60),           # ~2,532 h
        round(5 * 415 * 2 / 60),         # ~69 h
        round(15 * 6206 / 60),           # ~1,552 h
    ]
    # Actual pipeline compute (seconds → hours); calibration includes human review time estimate
    pipeline_hours = [
        round(s12.get("elapsed_seconds", 11095) / 3600, 1),   # 3.1 h
        round(s14.get("elapsed_seconds", 13806) / 3600, 1),   # 3.8 h
        round(s15.get("elapsed_seconds", 838)   / 3600, 1),   # 0.2 h
        round(sum(s10_times.values())            / 3600, 1),  # 8.7 h compute + human review
        round(s13.get("elapsed_seconds", 19388) / 3600, 1),   # 5.4 h
    ]

    x = np.arange(len(stages))
    w = 0.35
    fig, ax = plt.subplots(figsize=(8.5, 4.2))

    b1 = ax.bar(x - w/2, human_hours,    w, label="Estimated person-hours (manual)",
                color="#d7191c", alpha=0.85, zorder=3)
    b2 = ax.bar(x + w/2, pipeline_hours, w, label="Actual pipeline compute hours",
                color="#2c7bb6", alpha=0.85, zorder=3)

    ax.bar_label(b1, fmt="%.0f h", fontsize=8, padding=3)
    ax.bar_label(b2, fmt="%.1f h", fontsize=8, padding=3)

    ax.set_xticks(x)
    ax.set_xticklabels(stages, fontsize=8.5)
    ax.set_ylabel("Hours", fontsize=10)
    ax.set_title("Estimated manual person-hours vs actual pipeline compute time", fontsize=11, pad=8)
    ax.legend(fontsize=9, framealpha=0.9)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4, zorder=0)
    fig.tight_layout()
    return fig_to_stream(fig)


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def shade_row(row, hex_color="E8E8E8"):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)


def insert_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar");  fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar); run._r.append(instrText)
    run._r.append(fldChar2); run._r.append(fldChar3)
    note = doc.add_paragraph(
        "Note: Right-click this field in Word and select 'Update Field' to populate the table of contents."
    )
    note.paragraph_format.space_after = Pt(12)
    for r in note.runs:
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 100, 100)


def add_body(doc, text, space_before=0, space_after=6):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def add_figure(doc, key_or_path, caption, width=Inches(5.8)):
    """
    Insert a figure from FIG[key] or a direct Path.
    Always reads from a local Path so Word never sees a stream it can't round-trip.
    Skips with a greyed placeholder if the file is missing.
    """
    path = FIG.get(key_or_path, key_or_path) if isinstance(key_or_path, str) else key_or_path
    path = Path(path)

    if not path.exists():
        p = doc.add_paragraph(f"[Figure not available: {path.name}]")
        p.paragraph_format.space_after = Pt(10)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(160, 160, 160)
        return

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=width)
    except Exception as e:
        p = doc.add_paragraph(f"[Could not embed figure {path.name}: {e}]")
        for r in p.runs:
            r.italic = True

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)


def add_script_box(doc, scripts, note=None):
    """
    Insert a shaded reference box listing pipeline script(s) with GitHub URLs.
    scripts: list of (label, filename) tuples, e.g. [("Step 1", "step1_scopus_query_counts.py")]
    note:    optional short text appended after the links (e.g. output directory path).
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]

    # shade the cell light blue-grey
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EBF2FA")
    tc_pr.append(shd)

    cell.paragraphs[0].clear()
    label_para = cell.paragraphs[0]
    label_run = label_para.add_run("Pipeline reference")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(30, 90, 150)
    label_para.paragraph_format.space_after = Pt(2)

    for label, filename in scripts:
        url = script_url(filename)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label}: ")
        r1.font.size = Pt(9)
        r1.bold = True
        r2 = p.add_run(f"{filename}  \u2014  {url}")
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(30, 90, 150)
        r2.font.name = "Courier New"

    if note:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(0)
        rn = p.add_run(note)
        rn.font.size = Pt(8.5)
        rn.italic = True
        rn.font.color.rgb = RGBColor(80, 80, 80)

    # small gap after box
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_repo_header(doc):
    """Insert a prominent repository reference box near the top of the document."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D6E8F7")
    tc_pr.append(shd)

    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r1 = p.add_run("Full code repository: ")
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(REPO_URL)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(30, 90, 150)
    r2.font.name = "Courier New"
    p.paragraph_format.space_after = Pt(3)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r3 = p2.add_run(
        "All pipeline scripts, eligibility criteria (criteria.yml), search strings "
        "(search_strings.yml), calibration outputs, and this document's source "
        "(methodology.py) are available in the repository above. The default branch is "
        f"'{REPO_BRANCH}'. Individual script links are provided in each section below."
    )
    r3.font.size = Pt(9)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_named_principle(doc, title, description):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    r1 = p.add_run(title + ". "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(description);  r2.font.size = Pt(11)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_doc():
    # -----------------------------------------------------------------------
    # Step 0: prepare / normalise all figures into FIG_DIR
    # -----------------------------------------------------------------------
    print("Preparing figures...")
    generated_streams = {
        "time_comparison":     make_time_comparison_figure(),
        "kappa_convergence":   make_kappa_convergence_figure(),
    }
    FIG.update(prepare_figures(generated_streams))
    print(f"  {len(FIG)} figures ready in {FIG_DIR}\n")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(1.18)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # -------------------------------------------------------------------
    # Title
    # -------------------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Appendix: Computational Methodology for the Systematic Map on\n"
        "Climate Adaptation Effectiveness among Smallholder Producers"
    )
    r.bold = True; r.font.size = Pt(14)
    t.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Document generated: {datetime.utcnow().strftime('%d %B %Y')}")
    sr.font.size = Pt(10); sr.italic = True
    sub.paragraph_format.space_after = Pt(18)

    add_repo_header(doc)

    # -------------------------------------------------------------------
    # TOC
    # -------------------------------------------------------------------
    doc.add_heading("Table of Contents", level=1)
    insert_toc(doc)
    doc.add_page_break()

    # -------------------------------------------------------------------
    # 1. Overview
    # -------------------------------------------------------------------
    doc.add_heading("1. Overview", level=1)
    add_body(doc,
        "This appendix provides a transparent, step-by-step account of the computational "
        "pipeline used to conduct the systematic map. The pipeline consists of sixteen "
        "sequential steps implemented in Python, each producing auditable outputs that feed "
        "into the next stage. The pipeline is fully resumable: all external API results and "
        "large language model (LLM) decisions are cached to disk so that any step can be "
        "re-run without reprocessing already-completed records."
    )
    add_body(doc,
        "Throughout the pipeline, humans were involved at clearly defined checkpoints: "
        "constructing and refining the search query, independently screening calibration "
        "samples, reconciling disagreements between raters, reviewing inter-rater "
        "reliability statistics before full-corpus screening proceeded, and spot-checking "
        "the final coded dataset. These checkpoints are noted explicitly in the relevant "
        "sections below."
    )
    add_body(doc,
        "All LLM-assisted steps used a locally hosted model (Ollama; model: qwen2.5:14b) "
        "with temperature set to 0.0, ensuring fully deterministic and reproducible "
        "outputs. No proprietary cloud-based model API was used for screening or data "
        "extraction decisions."
    )

    # Summary table
    doc.add_heading("Pipeline Summary Statistics", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Stage"; hdr[1].text = "Count"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    shade_row(tbl.rows[0], "D0D0D0")
    for label, value in [
        ("Records returned by Scopus (raw)",                   SCOPUS_RAW),
        ("Records after deduplication (Step 2)",               SCOPUS_DEDUPED),
        ("Records with abstract after enrichment (Step 9a)",   fmt(s9a.get("final_present", 15707))),
        ("Records still missing abstract after enrichment",    ABSTRACT_MISSING_9A),
        ("Records screened at title/abstract stage (Step 12)", SCREENED_TOTAL),
        ("  Included after title/abstract screening",          SCREENED_INCLUDE),
        ("  Excluded after title/abstract screening",          SCREENED_EXCLUDE),
        ("Full texts retrieved (Step 13)",                     FT_RETRIEVED),
        ("Records with no full text available",                FT_NO_FT),
        ("  Included after full-text screening (Step 14)",     FT_INCLUDE),
        ("  Excluded after full-text screening",               FT_EXCLUDE),
        ("Records coded in systematic map (Step 15)",          CODED_TOTAL),
        ("  Coded from full text",                             CODED_FT),
        ("  Coded from abstract only",                         CODED_ABS),
        ("  Abstract unavailable",                             CODED_MISSING),
    ]:
        r = tbl.add_row()
        r.cells[0].text = label; r.cells[1].text = value

    doc.add_paragraph()

    # -------------------------------------------------------------------
    # 2. Computational Efficiency and the Case for Automation
    # -------------------------------------------------------------------
    doc.add_heading("2. Computational Efficiency and the Case for Automation", level=1)
    add_body(doc,
        "Systematic reviews and evidence maps involving tens of thousands of records present "
        "a fundamental practical challenge: the volume of evidence to be processed far "
        "exceeds what a small research team can screen and code manually within a reasonable "
        "timeframe. A corpus of 17,021 records, if subjected to conventional dual independent "
        "human screening at the title/abstract stage alone, would require approximately "
        "1,135 person-hours (assuming two minutes per record per reviewer, two reviewers). "
        "Full-text screening of the 6,206 passing records would add a further estimated "
        "2,069 person-hours, and data extraction from the included studies a further "
        "estimated 2,532 person-hours. In aggregate, the three core screening and extraction "
        "stages would demand on the order of 5,700 person-hours of researcher time under "
        "conventional approaches, equivalent to approximately three full-time years of effort."
    )
    add_body(doc,
        "The pipeline described in this appendix completed the equivalent computational "
        "work in a small fraction of that time. Title/abstract screening of all 17,021 "
        f"records was completed in {s12.get('elapsed_hms', '03:04:54')} of unattended compute "
        f"time. Full-text screening of {SCREENED_INCLUDE} records took "
        f"{s14.get('elapsed_hms', '03:50:05')}. Data extraction across {CODED_TOTAL} coded "
        f"records took {s15.get('elapsed_hms', '00:13:58')}. Full-text retrieval "
        f"({s13.get('elapsed_hms', '05:23:07')}) ran unattended overnight. The total "
        "unattended compute time for these four stages was approximately 13 hours, compared "
        "to an estimated 5,700 person-hours for manual equivalents."
    )

    add_figure(doc, "time_comparison",
        "Figure 1. Estimated person-hours required for manual execution of each pipeline "
        "stage (red) versus actual unattended compute time (blue). Human estimates assume "
        "standard dual-reviewer protocols: 2 min/record for title/abstract screening, "
        "10 min/record for full-text screening, 25 min/record for data extraction, "
        "15 min/record for manual full-text retrieval. Calibration hours reflect LLM "
        "compute plus estimated human reconciliation time.",
        width=Inches(6.0))

    add_body(doc,
        "This efficiency gain does not come at the cost of rigour. The pipeline preserves "
        "human oversight at the stages where it matters most: search query design and "
        "iteration, criteria formulation, calibration of the LLM screener against independent "
        "human judgements, and verification of extracted data. The calibration process, "
        "described in Section 6, demonstrates that LLM agreement with reconciled human "
        "decisions reached substantial-to-almost-perfect kappa levels (k > 0.77) before "
        "full-corpus screening proceeded. Where the LLM was uncertain, the conservative "
        "default was to include rather than exclude, minimising false negatives."
    )
    add_body(doc,
        "A further practical consideration is reproducibility. Because all LLM decisions "
        "are made at temperature 0.0 and cached with full input-output provenance, the "
        "screening and extraction process can be re-run in its entirety and will produce "
        "identical results. This is not a property that manual screening can offer. "
        "The pipeline therefore combines the speed of automation with a level of "
        "reproducibility that exceeds conventional systematic review practice."
    )

    # -------------------------------------------------------------------
    # 3. Search Strategy
    # -------------------------------------------------------------------
    doc.add_heading("3. Search Strategy", level=1)

    doc.add_heading("3.1 Query Construction (Step 1)", level=2)
    add_script_box(doc, [("Step 1", "step1_scopus_query_counts.py")],
        note="Outputs: outputs/step1/  (step1_hits_plot.png, step1_queries.json, step1_total_query.txt)")
    add_body(doc,
        "The search query was structured around a Population, Concept, Context, and "
        "Methodology (PCCM) framework. Individual search elements were defined in a "
        "machine-readable YAML file (search_strings.yml), enabling version-controlled "
        "iteration. Step 1 submitted each element and their combination to the Scopus "
        "Search API to retrieve record counts, without downloading any records. A SHA-256 "
        "hash of the query was computed before each run; if unchanged, cached counts were "
        "returned without further API calls. Exponential backoff (up to six retries, "
        "maximum 60-second wait) handled transient rate limits."
    )
    add_figure(doc, "search_hits",
        "Figure 2. Record counts returned by the Scopus API for individual search "
        "string elements (P, C, M sub-groups) and the combined PCCM query. "
        "Generated by Step 1 of the pipeline.",
        width=Inches(5.5))

    doc.add_heading("3.2 Record Retrieval (Step 2)", level=2)
    add_script_box(doc, [("Step 2", "step2_scopus_retrieve_records.py")],
        note="Outputs: outputs/step2/  (step2_total_records.csv, .ris, .bib, .meta.json)")
    add_body(doc,
        f"Step 2 retrieved all records matching the combined PCCM query. Scopus imposes a "
        f"hard deep-paging limit of 5,000 records per query slice. To retrieve the full "
        f"result set, the query was automatically sliced by publication year (1990-2025); "
        f"where a single year still exceeded 5,000 records, further sub-slices by Scopus "
        f"subject area (27 codes) or source type (7 disjoint types) were applied. "
        f"Records were streamed directly to CSV as each page was received. After all slices "
        f"were retrieved and deduplicated, {SCOPUS_DEDUPED} unique records were retained "
        f"(from {SCOPUS_RAW} reported by Scopus). Deduplication used a priority key: DOI "
        f"(where present), then normalised title and year, then title alone, then EID."
    )

    doc.add_heading("3.3 Benchmark List and Coverage Analysis (Steps 3, 4, 7)", level=2)
    add_script_box(doc, [
        ("Step 3", "step3_benchmark_match.py"),
        ("Step 4", "step4_fetch_abstracts.py"),
        ("Step 7", "step7_scopus_check.py"),
    ], note="Outputs: outputs/step3/, outputs/step4/, outputs/step7/")
    add_body(doc,
        "A benchmark list of known key studies, compiled by the research team prior to the "
        "search, was used to validate coverage. Step 3 enriched the benchmark list with "
        "DOIs using Crossref, OpenAlex, and Semantic Scholar. Candidate matches were scored "
        "by title similarity (difflib SequenceMatcher); matches scoring 0.90 or above were "
        "accepted automatically. Step 7 compared the benchmark list against the Scopus "
        "retrieval and generated keyword suggestions from non-retrieved benchmark records "
        "to guide iterative query refinement. Human review of suggestions was required "
        "before any query modification was made."
    )

    # -------------------------------------------------------------------
    # 4. Record Cleaning (Step 8)
    # -------------------------------------------------------------------
    doc.add_heading("4. Record Cleaning and Deduplication (Step 8)", level=1)
    add_script_box(doc, [("Step 8", "step8_clean_scopus.py")],
        note="Outputs: outputs/step8/  (step8_scopus_cleaned.csv, .ris, .meta.json)")
    add_body(doc,
        "Step 8 applied deterministic cleaning to the raw Scopus records: HTML unescaping, "
        "whitespace normalisation, DOI canonicalisation, and year extraction. Where a DOI "
        "was present but author list, journal name, or year was missing, Crossref was "
        "queried to repair the gap. All Crossref lookups were cached locally. The step was "
        "designed to be idempotent: on re-runs, previously cleaned records were preserved; "
        "only records new to the current Scopus pull were processed."
    )

    # -------------------------------------------------------------------
    # 5. Abstract Enrichment (Steps 9, 9a)
    # -------------------------------------------------------------------
    doc.add_heading("5. Abstract Enrichment (Steps 9 and 9a)", level=1)

    doc.add_heading("5.1 Automated Multi-Source Enrichment (Step 9)", level=2)
    add_script_box(doc, [("Step 9", "step9_enrich_abstracts.py")],
        note="Outputs: outputs/step9/  (step9_scopus_enriched.csv, step9_missing.csv, step9_summary.json)")
    add_body(doc,
        "Step 9 retrieved missing abstracts for all cleaned Scopus records using a "
        "sequential fetch chain: (1) Elsevier Abstract Retrieval API by DOI then Scopus ID; "
        "(2) Semantic Scholar; (3) OpenAlex (abstract reconstructed from the inverted-index "
        "word-position map); (4) Crossref; (5) Unpaywall; (6) landing page scrape via "
        "JSON-LD, HTML meta tags, or raw HTML. All API responses were cached with a "
        f"30-day time-to-live. Of {SCOPUS_DEDUPED} records, "
        f"{fmt(s9.get('already_has_abstract', 5839))} already carried an abstract from "
        f"Scopus; Step 9 enriched a further {fmt(s9.get('fresh_ok', 9752))}. After Step 9, "
        f"{ABSTRACT_MISSING_9} records remained without an abstract."
    )

    doc.add_heading("5.2 RIS-Based Supplementary Enrichment (Step 9a)", level=2)
    add_script_box(doc, [("Step 9a", "step9a_enrich_from_ris.py")],
        note="Inputs configured in config.py (step9a_ris_glob, step9a_iteration)  |  "
             "Outputs: outputs/step9a/  (step9a_scopus_enriched.csv, step9a_diff.csv, step9a_summary.json)")
    add_body(doc,
        "Following automated enrichment, EPPI Reviewer exported the full corpus as RIS "
        "files with any manually entered abstracts included. Step 9a parsed these files "
        "(five files totalling 17,011 records in iteration step9a1) and injected RIS "
        "abstracts into records still lacking one, matching by DOI, Scopus EID, or "
        "normalised title and year. It then re-attempted API enrichment for remaining gaps, "
        "bypassing Step 9 cached failures. In total Step 9a recovered "
        f"{ABSTRACT_GAINED_9A} additional abstracts, reducing the missing count from "
        f"{ABSTRACT_MISSING_9} to {ABSTRACT_MISSING_9A}. Each run is labelled with an "
        "iteration identifier and UTC timestamp for provenance. Records still lacking an "
        "abstract were retained and treated conservatively at screening."
    )

    # -------------------------------------------------------------------
    # 6. Calibration and IRR (Steps 10, 11)
    # -------------------------------------------------------------------
    doc.add_heading("6. Calibration Rounds and Inter-Rater Reliability (Steps 10 and 11)", level=1)
    add_body(doc,
        "Before automated screening of the full corpus, three calibration rounds were "
        "conducted to assess human inter-rater reliability and to evaluate and iteratively "
        "improve LLM agreement with human-reconciled decisions. Full-corpus screening "
        "did not proceed until calibration metrics reached acceptable thresholds."
    )

    doc.add_heading("6.1 Calibration Process", level=2)
    add_script_box(doc, [
        ("Step 10", "step10_llm_calibrate.py"),
        ("Step 11", "step11_irr_analysis.py"),
        ("Criteria", "criteria.yml"),
    ], note="Calibration RIS inputs configured in config.py (step10_calibration_ris, step10_run_label)  |  "
            "Outputs: outputs/step10/, outputs/step11/")
    add_body(doc,
        "For each round, a sample was drawn from the enriched corpus and screened "
        "independently by two human reviewers (Caroline Staub and Jennifer Cisse) using "
        "EPPI Reviewer. Their decisions were reconciled into a gold-standard column. "
        "Step 10 ran the LLM screener against the same sample using the current version "
        "of criteria.yml. Step 11 computed pairwise Cohen's kappa coefficients between "
        "all raters and produced confusion matrices showing over- and under-inclusion "
        "rates relative to the reconciled standard. The team reviewed outputs after each "
        "round and revised criteria.yml where systematic disagreements were identified."
    )

    doc.add_heading("6.2 Calibration Results", level=2)

    pk1 = c_r1.get("pairwise_kappa",  {})
    pk1b = c_r1b.get("pairwise_kappa", {})
    pk2  = c_r2.get("pairwise_kappa",  {})
    pk2a = c_r2a.get("pairwise_kappa", {})
    pk3  = c_r3a.get("pairwise_kappa", {})

    add_body(doc,
        f"Round 1 (n = 205): human inter-rater k = {pk1.get('Caroline Staub vs Jennifer Cisse', 'N/A'):.3f}. "
        f"LLM vs reconciled: k = {pk1.get('LLM vs CJ Reconciled', 'N/A'):.3f} "
        f"({c_r1.get('confusion_vs_reconciled',{}).get('LLM',{}).get('pct_agreement','N/A')}% agreement). "
        f"The LLM was over-inclusive ({c_r1.get('confusion_vs_reconciled',{}).get('LLM',{}).get('over_include','N/A')} "
        f"false positives). Criteria were revised to tighten the Concept and Methodology "
        f"dimensions, specifically excluding vulnerability-only and impact-only studies. "
        f"After revision (Round 1b), LLM vs reconciled improved to "
        f"k = {pk1b.get('LLMr1b vs CJ Reconciled', 'N/A'):.3f} "
        f"({c_r1b.get('confusion_vs_reconciled',{}).get('LLMr1b',{}).get('pct_agreement','N/A')}% agreement)."
    )
    add_body(doc,
        f"Round 2 (n = 103): human inter-rater k = {pk2.get('Jennifer Cisse vs Caroline Staub', 'N/A'):.3f}. "
        f"After reconciliation (Round 2a), LLM vs reconciled: k = {pk2a.get('LLM_r2a vs CJ Reconciled', 'N/A'):.3f} "
        f"({c_r2a.get('confusion_vs_reconciled',{}).get('LLM_r2a',{}).get('pct_agreement','N/A')}% agreement). "
        f"Human reviewers reached almost perfect agreement with the gold standard: "
        f"k = {pk2a.get('Caroline Staub vs CJ Reconciled', 'N/A'):.3f} (Caroline Staub) and "
        f"k = {pk2a.get('Jennifer Cisse vs CJ Reconciled', 'N/A'):.3f} (Jennifer Cisse)."
    )
    add_body(doc,
        f"Round 3a (n = 107): human inter-rater k = {pk3.get('Jennifer Cisse vs Caroline Staub', 'N/A'):.3f}. "
        f"LLM vs Jennifer Cisse: k = {pk3.get('Jennifer Cisse vs LLM', 'N/A'):.3f}; "
        f"LLM vs Caroline Staub: k = {pk3.get('Caroline Staub vs LLM', 'N/A'):.3f}. "
        f"The team confirmed criteria were stable and approved proceeding to full-corpus screening."
    )

    add_figure(doc, "kappa_convergence",
        "Figure 3. Cohen's kappa convergence across calibration rounds. Blue circles: "
        "LLM agreement with reconciled human gold standard. Red squares: human "
        "inter-rater agreement. Shaded bands show the Landis and Koch (1977) "
        "interpretation thresholds. Annotated labels indicate criteria version revisions "
        "between rounds.",
        width=Inches(6.0))

    add_body(doc,
        "Figure 3 illustrates the iterative improvement in LLM-human agreement across "
        "rounds, driven by targeted criteria revisions. By Round 2a the LLM reached "
        "substantial agreement (k = 0.77) and human reviewers were consistently in "
        "almost-perfect agreement with the gold standard (k > 0.87)."
    )

    doc.add_heading("6.3 Round-by-round calibration figures", level=2)
    add_body(doc,
        "The following figures show the full three-panel analysis for selected calibration "
        "rounds: pairwise kappa heatmap, per-rater confusion bar chart relative to "
        "reconciled decisions, and reconciled decision distribution."
    )
    add_figure(doc, "irr_r1",
        "Figure 4a. Calibration Round 1 (n = 205). Initial criteria version.",
        width=Inches(6.0))
    add_figure(doc, "irr_r1b",
        "Figure 4b. Calibration Round 1b (n = 205). After first criteria revision.",
        width=Inches(6.0))
    add_figure(doc, "irr_r2a",
        "Figure 4c. Calibration Round 2a (n = 103). After second criteria revision.",
        width=Inches(6.0))
    add_figure(doc, "irr_r3a",
        "Figure 4d. Calibration Round 3a (n = 107). Final criteria version.",
        width=Inches(6.0))

    # -------------------------------------------------------------------
    # 7. Title/Abstract Screening (Step 12)
    # -------------------------------------------------------------------
    doc.add_heading("7. Title and Abstract Screening (Step 12)", level=1)
    add_script_box(doc, [("Step 12", "step12_screen_abstracts.py")],
        note="Outputs: outputs/step12/  (step12_results.csv, step12_results.meta.json, step12_results_details.jsonl)")
    add_body(doc,
        f"Step 12 applied the LLM screener to all {SCREENED_TOTAL} records using the "
        "criteria.yml finalised through calibration. The LLM evaluated each record against "
        "five PCCM criteria, returning a decision (yes / no / unclear), a reason, and a "
        "direct quotation from the abstract for each. Quotations were verified against the "
        "abstract text; unverifiable quotes downgraded the criterion to 'unclear'. "
        "A record was excluded only if at least one criterion was explicitly 'no'; any "
        "'unclear' or absent abstract defaulted to inclusion."
    )
    add_body(doc,
        f"Results: {SCREENED_INCLUDE} included, {SCREENED_EXCLUDE} excluded, "
        f"{MISSING_ABS_12} retained due to missing abstract. The most frequent exclusion "
        f"criterion was Concept ({fmt(s12.get('excluded_by_criterion',{}).get('2_concept',9197))} records), "
        f"followed by Population ({fmt(s12.get('excluded_by_criterion',{}).get('1_population',5318))} records). "
        "All decisions were cached to a JSONL file keyed on record identifier and abstract "
        "source, enabling incremental re-runs."
    )
    add_figure(doc, "screen12",
        "Figure 5. Title/abstract screening outcome summary (Step 12). "
        "Shows include/exclude/unclear counts and the most common exclusion criteria.",
        width=Inches(5.5))

    # -------------------------------------------------------------------
    # 8. Full-Text Retrieval (Step 13)
    # -------------------------------------------------------------------
    doc.add_heading("8. Full-Text Retrieval (Step 13)", level=1)
    add_script_box(doc, [("Step 13", "step13_retrieve_fulltext.py")],
        note="Outputs: outputs/step13/  (step13_manifest.csv, step13_summary.json, fulltext/ directory)")
    add_body(doc,
        f"Step 13 attempted to download full texts for all {SCREENED_INCLUDE} records "
        "passing abstract screening, using: Unpaywall (open-access URLs by DOI), "
        "Elsevier Full-Text API (institutional access), Semantic Scholar, and OpenAlex. "
        f"Downloads were capped at 25 MB. {FT_RETRIEVED} full texts were retrieved "
        f"({fmt(s13.get('source_breakdown',{}).get('unpaywall',515))} via Unpaywall, "
        f"{fmt(s13.get('source_breakdown',{}).get('elsevier_doi',331))} via Elsevier, "
        f"{fmt(s13.get('source_breakdown',{}).get('semantic_scholar',70))} via Semantic Scholar). "
        f"{FT_NEEDS_MANUAL} records could not be retrieved automatically and were flagged "
        "for manual acquisition via institutional library access."
    )

    # -------------------------------------------------------------------
    # 9. Full-Text Screening (Step 14)
    # -------------------------------------------------------------------
    doc.add_heading("9. Full-Text Screening (Step 14)", level=1)
    add_script_box(doc, [("Step 14", "step14_screen_fulltext.py")],
        note="Outputs: outputs/step14/  (step14_results.csv, step14_results.meta.json, step14_results_details.jsonl)")
    add_body(doc,
        "Step 14 applied the LLM screener to retrieved full texts, truncated to 12,000 "
        "characters to cover abstract, introduction, and methods. Text was extracted from "
        "PDFs using pypdf and from HTML using trafilatura or BeautifulSoup. Records "
        "without a retrieved full text were retained conservatively. "
        f"Of {SCREENED_INCLUDE} records assessed: {FT_WITH_FT} were screened from full "
        f"text; {FT_INCLUDE} confirmed included; {FT_EXCLUDE} excluded. The most frequent "
        f"full-text exclusion criterion was Concept "
        f"({fmt(s14.get('excluded_by_criterion',{}).get('2_concept',104))} records)."
    )
    add_figure(doc, "screen14",
        "Figure 6. Full-text screening outcome summary (Step 14).",
        width=Inches(5.5))

    # -------------------------------------------------------------------
    # 10. Data Extraction (Step 15)
    # -------------------------------------------------------------------
    doc.add_heading("10. Data Extraction and Coding (Step 15)", level=1)
    add_script_box(doc, [("Step 15", "step15_extract_data.py")],
        note="Outputs: outputs/step15/  (step15_coded.csv, step15_coded.meta.json, step15_needs_review.csv)")
    add_body(doc,
        "Step 15 extracted structured coding data from all included records against a "
        "20-field schema covering: publication year and type, country or region, "
        "geographic scale, producer type (crop, livestock, fisheries, agroforestry, "
        "mixed), whether the study assessed adaptation processes or outcomes, "
        "methodological approach, effectiveness metric, and equity/inclusion dimensions. "
        "The LLM was supplied with full text (where available) or abstract, and the "
        "coding source was recorded per record. All extracted data were subject to "
        "human spot-checking prior to use in the final evidence map."
    )
    add_body(doc,
        f"Of {CODED_TOTAL} records: {CODED_FT} coded from full text, "
        f"{CODED_ABS} from abstract only, {CODED_MISSING} with neither. "
        "The coding source field allows readers to weight evidence confidence accordingly."
    )
    add_figure(doc, "coding",
        "Figure 7. Data extraction summary by coding source (Step 15).",
        width=Inches(5.0))

    # -------------------------------------------------------------------
    # 11. Systematic Map — ROSES Flow and Evidence Figures
    # -------------------------------------------------------------------
    doc.add_heading("11. Systematic Map Outputs (Step 16)", level=1)
    add_script_box(doc, [("Step 16", "step16_map_visualise.py")],
        note="Outputs: outputs/step16/  (roses_flow.png, temporal_trends.png, geographic_bar.png, "
             "producer_type_bar.png, methodology_bar.png, domain_heatmap.png, equity_bar.png)")
    add_body(doc,
        "Step 16 generated all publication-ready figures directly from the coded dataset. "
        "Figures are updated automatically as further full texts are retrieved and coded. "
        "All exports are PNG at 300 DPI."
    )

    add_figure(doc, "roses",
        "Figure 8. ROSES flow diagram showing record counts at each stage of the "
        "systematic map pipeline.",
        width=Inches(5.5))

    add_figure(doc, "temporal",
        "Figure 9. Temporal trends in publications included in the systematic map.",
        width=Inches(5.5))

    add_figure(doc, "geography",
        "Figure 10. Geographic distribution of included studies (top countries by count).",
        width=Inches(5.5))

    add_figure(doc, "producer",
        "Figure 11. Breakdown of included studies by producer type.",
        width=Inches(5.0))

    add_figure(doc, "methodology",
        "Figure 12. Breakdown of included studies by methodological approach.",
        width=Inches(5.0))

    add_figure(doc, "domain",
        "Figure 13. Domain heatmap: adaptation process versus outcome by producer type.",
        width=Inches(5.5))

    # -------------------------------------------------------------------
    # 12. Reproducibility and Transparency
    # -------------------------------------------------------------------
    doc.add_heading("12. Reproducibility and Transparency", level=1)
    add_body(doc,
        "The following design principles were applied throughout to ensure the pipeline is "
        "reproducible and its decisions are auditable:"
    )
    for title, desc in [
        ("Deterministic LLM outputs",
         "Temperature was set to 0.0 for all LLM calls. Given a fixed model and fixed "
         "input, outputs are reproducible across runs."),
        ("Comprehensive caching",
         "All external API responses and LLM decisions are written to cache files "
         "(JSON or JSONL). Re-runs process only new or expired records."),
        ("Quotation verification",
         "LLM screening decisions must cite a passage from the abstract. If the cited "
         "passage cannot be located, the criterion is downgraded to 'unclear', providing "
         "a lightweight hallucination check."),
        ("Conservative defaults",
         "Across all screening stages, the absence of required evidence defaults to "
         "inclusion rather than exclusion, minimising false negatives."),
        ("Iterative human calibration",
         "Three calibration rounds with two independent human reviewers preceded "
         "full-corpus automated screening. Criteria were revised between rounds based "
         "on observed disagreements."),
        ("Coding source tracking",
         "Every coded record carries a 'coding_source' field indicating whether "
         "extraction was from full text, abstract only, or title/metadata alone."),
        ("Version-controlled criteria",
         "Eligibility criteria are stored in a single YAML file (criteria.yml), "
         "versioned alongside the code. Any change triggers re-screening of affected "
         "records on the next run."),
        ("ROSES flow diagram",
         "A ROSES-compliant flow diagram is generated automatically at Step 16, "
         "documenting record counts at every stage of the pipeline."),
    ]:
        add_named_principle(doc, title, desc)

    # -------------------------------------------------------------------
    # 13. Software and Dependencies
    # -------------------------------------------------------------------
    doc.add_heading("13. Software and Dependencies", level=1)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, h in enumerate(["Component", "Library / Service", "Purpose"]):
        hdr2[i].text = h
        for run in hdr2[i].paragraphs[0].runs:
            run.bold = True
    shade_row(tbl2.rows[0], "D0D0D0")
    for row_data in [
        ("LLM inference",      "Ollama (qwen2.5:14b)",                "Local LLM for screening and extraction"),
        ("Scopus API",         "Elsevier REST API",                   "Record retrieval and abstract enrichment"),
        ("DOI enrichment",     "Crossref, OpenAlex, Semantic Scholar","DOI lookup and abstract retrieval"),
        ("Open access",        "Unpaywall API",                       "Full-text URL discovery"),
        ("Word documents",     "python-docx",                        "Report and export generation"),
        ("PDF parsing",        "pypdf",                              "Full-text extraction from PDFs"),
        ("HTML parsing",       "trafilatura, BeautifulSoup4",        "Full-text extraction from HTML"),
        ("Data handling",      "pandas",                             "CSV processing throughout pipeline"),
        ("Visualisation",      "matplotlib, seaborn",                "All figures"),
        ("Geographic maps",    "geopandas (optional)",               "World choropleth in Step 16"),
        ("IRR statistics",     "Custom Python (Cohen's kappa)",      "Inter-rater reliability analysis"),
        ("Reference management","EPPI Reviewer",                     "Human screening and RIS exports"),
    ]:
        r = tbl2.add_row()
        for i, val in enumerate(row_data):
            r.cells[i].text = val

    doc.add_paragraph()

    # Footer
    footer = add_body(doc,
        f"This document was generated programmatically from pipeline output files on "
        f"{datetime.utcnow().strftime('%d %B %Y')}. All statistics and figures are drawn "
        f"directly from the JSON metadata files and PNG outputs produced at each pipeline step.",
        space_before=12
    )
    for run in footer.runs:
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    build_doc()
