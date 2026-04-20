"""
Generate ILRI01 Short-Term Consultancy Agreement + Appendices B & C
Built on the Bristlepine template (preserves header/footer/fonts).

Run: python3 scripts/generate_contract.py
Output: documentation/coding/systematic-map/ILRI01_JD_RA_2026.docx
"""

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = Path("documentation/coding/systematic-map/contracts/template_subcontract.docx")
OUT      = Path("documentation/coding/systematic-map/contracts/ILRI01_RA_Contract_2026.docx")

HOURLY_RATE_PKR  = 450    # Rs. per hour
HOURS_MIN        = 0.5    # 30 min per paper (fast)
HOURS_MAX        = 1.0    # 60 min per paper (slow)
HOURS_EST        = 0.8    # ~50 papers / 40 hrs (expected)
CALIB_PAPERS     = 5
BATCH_PAPERS     = 50

def fee_range(n_papers):
    lo  = int(n_papers * HOURS_MIN * HOURLY_RATE_PKR)
    hi  = int(n_papers * HOURS_MAX * HOURLY_RATE_PKR)
    est = int(n_papers * HOURS_EST * HOURLY_RATE_PKR)
    return lo, est, hi

def hrs_range(n_papers):
    lo  = n_papers * HOURS_MIN
    hi  = n_papers * HOURS_MAX
    est = n_papers * HOURS_EST
    # format as int if whole number
    fmt = lambda x: str(int(x)) if x == int(x) else f"{x:.0f}"
    return fmt(lo), fmt(est), fmt(hi)

shutil.copy(TEMPLATE, OUT)
doc = Document(OUT)

# ── Clear body, keep sectPr ───────────────────────────────────────────────
body   = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child is not sectPr:
        body.remove(child)

def _ins(el):
    """Insert element before sectPr."""
    idx = list(body).index(sectPr) if sectPr is not None else len(body)
    body.insert(idx, el)

# ── Core helpers ──────────────────────────────────────────────────────────

def _move(el):
    body.remove(el)
    _ins(el)

def para(text="", style="Normal", bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.LEFT, sa=Pt(5), sb=Pt(0)):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after  = sa
    p.paragraph_format.space_before = sb
    if text:
        _add_runs(p, text, bold=bold, italic=italic)
    _move(p._element)
    return p

def h(text, level=2, sb=Pt(14), sa=Pt(4)):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after  = sa
    p.add_run(text)
    _move(p._element)
    return p

def bul(text, indent=Cm(1.0), sa=Pt(3)):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = sa
    _add_runs(p, f"\u2022  {text}")
    _move(p._element)
    return p

def blank(pt=2):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(pt)
    _move(p._element)

def page_break():
    p   = doc.add_paragraph(style="Normal")
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    _move(p._element)

def set_table_borders(table, color="000000", sz="4", val="single"):
    """Apply a simple solid black border to all cells in a table."""
    sides = ("top", "left", "bottom", "right", "insideH", "insideV")
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove any existing tblBorders element first
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   val)
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def tbl_move(t):
    set_table_borders(t)
    body.remove(t._element)
    _ins(t._element)

# ── Yellow-highlight helper ───────────────────────────────────────────────
# Segments text on [...] tokens and highlights them yellow.

_PLACEHOLDER = re.compile(r"(\[[^\]]+\])")

def _add_runs(p, text, bold=False, italic=False, highlight=False):
    """Add text to paragraph, auto-highlighting any [...] tokens in yellow."""
    parts = _PLACEHOLDER.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.bold   = bold
        run.italic = italic
        is_ph = bool(_PLACEHOLDER.fullmatch(part))
        if is_ph or highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def _cell_runs(cell, label, value, label_bold=True):
    """Add label + value to a cell paragraph, highlighting [...] in value."""
    p = cell.paragraphs[0]
    if p.text:
        p = cell.add_paragraph()
    r1 = p.add_run(label + "  ")
    r1.bold = label_bold
    r1.font.size = Pt(10)
    parts = _PLACEHOLDER.split(value)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        r.font.size = Pt(10)
        if _PLACEHOLDER.fullmatch(part):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(4)

# ── Fee helpers ───────────────────────────────────────────────────────────


# ══════════════════════════════════════════════════════════════════════════
# PAGE 1 — MAIN AGREEMENT
# ══════════════════════════════════════════════════════════════════════════

# Title
p = doc.add_paragraph(style="Heading 1")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
p.add_run("SHORT-TERM CONSULTANCY AGREEMENT")
_move(p._element)

p = doc.add_paragraph(style="Normal")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Contract Reference: ILRI01-JD-RA-2026  \u00b7  Evidence Coding \u2014 Systematic Map (D5.6)")
r.italic = True; r.font.size = Pt(10)
_move(p._element)

# ── Parties table ─────────────────────────────────────────────────────────
pt = doc.add_table(rows=1, cols=2)
pt.style = "TableNormal"
lc = pt.cell(0, 0)
rc = pt.cell(0, 1)

_cell_runs(lc, "Client:", "Bristlepine Resilience Consultants LLC, a limited liability company. Represented by Zarrar Khan, Partner.")
p2 = lc.add_paragraph()
r2 = p2.add_run("Email: info@bristlep.com")
r2.font.size = Pt(10)
p2.paragraph_format.space_after = Pt(4)
p3 = lc.add_paragraph()
r3 = p3.add_run("Collectively referred to as the \u201cParties.\u201d")
r3.italic = True; r3.font.size = Pt(9)

_cell_runs(rc, "Consultant:", "[Full name], [Country of residence], an independent contractor.")
p4 = rc.add_paragraph()
r4 = p4.add_run("Email: ")
r4.font.size = Pt(10)
r5 = p4.add_run("[Consultant email]")
r5.font.size = Pt(10)
r5.font.highlight_color = WD_COLOR_INDEX.YELLOW

tbl_move(pt)
blank()

# ── 1. Nature of Relationship ─────────────────────────────────────────────
h("1.  Nature of Relationship")
para(
    "This Agreement constitutes an independent contractor arrangement. It does not create "
    "an employment relationship, partnership, joint venture, or agency. The Consultant is not "
    "an employee of Bristlepine Resilience Consultants LLC or ILRI/CGIAR and is not entitled "
    "to employment benefits, statutory deductions, or protections arising from an employment "
    "relationship. The Consultant is solely responsible for their own taxes and social-security "
    "obligations in their country of residence.")

# ── 2. Scope of Work ──────────────────────────────────────────────────────
h("2.  Scope of Work")
para(
    "The Consultant will provide evidence-coding services for the ILRI/CGIAR Systematic Map "
    "project (Deliverable D5.6), examining how climate adaptation processes and outcomes are "
    "measured for smallholder agricultural producers in low- and middle-income countries. "
    "Specific responsibilities:")
bul("Code peer-reviewed and grey literature papers against the structured codebook "
    "(16 fields per paper), working independently and without discussion with other coders.")
bul("Participate in the initial 5-paper calibration session to align codebook application "
    "before independent batch coding begins.")
bul("Flag ambiguous cases with notes in the coding sheet; contribute to codebook "
    "refinement discussions between batches.")
bul("Deliver completed coding sheets in CSV format via Google Drive.")
bul("Expected throughput: approximately 40\u201360 papers per week at full effort, "
    "subject to paper complexity.")
para(
    "The Consultant is applying defined criteria systematically. This is not a research "
    "design or analysis role \u2014 rigour and consistency matter more than domain expertise.",
    italic=True)

# ── 3. Compensation and Invoicing ─────────────────────────────────────────
h("3.  Compensation and Invoicing")
h("Hourly Rate and Payment Batches", level=3)
para(
    f"Work is paid on an hourly basis at Rs. {HOURLY_RATE_PKR}/hr (Pakistani Rupees). "
    f"Papers are estimated to take {int(HOURS_MIN*60)}\u2013{int(HOURS_MAX*60)} minutes each, "
    f"with an expected average of ~{int(HOURS_EST*60)} minutes (~{BATCH_PAPERS} papers per 40-hour week). "
    "The Consultant will submit a brief timesheet with each invoice listing actual hours worked "
    "and papers coded. All payments will be made in PKR unless otherwise agreed in writing.")

# Batch payment table
bt = doc.add_table(rows=5, cols=4)
bt.style = "TableNormal"
for i, hdr in enumerate(["Batch", "Papers", "Est. hours", "Est. fee (PKR)"]):
    c = bt.rows[0].cells[i]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)

c_lo, c_est, c_hi = fee_range(CALIB_PAPERS)
b_lo, b_est, b_hi = fee_range(BATCH_PAPERS)
ch_lo, ch_est, ch_hi = hrs_range(CALIB_PAPERS)
bh_lo, bh_est, bh_hi = hrs_range(BATCH_PAPERS)
bdata = [
    ("Calibration (R1a)", str(CALIB_PAPERS),
     f"{ch_lo}\u2013{ch_hi} hrs (~{ch_est} expected)",
     f"Rs. {c_lo:,}\u2013{c_hi:,} (~Rs. {c_est:,})"),
    ("Batch 1",  f"~{BATCH_PAPERS}",
     f"{bh_lo}\u2013{bh_hi} hrs (~{bh_est} expected)",
     f"Rs. {b_lo:,}\u2013{b_hi:,} (~Rs. {b_est:,})"),
    ("Batch 2+", f"~{BATCH_PAPERS}",
     f"{bh_lo}\u2013{bh_hi} hrs (~{bh_est} expected)",
     f"Rs. {b_lo:,}\u2013{b_hi:,} per batch"),
    ("Extension batches", "[TBD]", "[TBD]", "[TBD]"),
]
for ri, row in enumerate(bdata, 1):
    for ci, val in enumerate(row):
        cell = bt.rows[ri].cells[ci]
        parts = _PLACEHOLDER.split(val)
        p = cell.paragraphs[0]
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(10)
            if _PLACEHOLDER.fullmatch(part):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
tbl_move(bt)
blank()

h("Invoicing", level=3)
para(
    "The Consultant shall submit an invoice to Bristlepine upon completion of each batch "
    "using the template in Appendix C. Each invoice must include: batch description, "
    "papers coded, dates worked, hours logged, and total fee. "
    "Payment will be made within 14 calendar days of invoice approval "
    "via bank transfer to the account details in Appendix B.")

h("Conditional Payment", level=3)
para(
    "Payment is conditional on Bristlepine\u2019s acceptance of the batch deliverable. "
    "Bristlepine shall notify the Consultant of acceptance or revision requests within "
    "5 business days of delivery. Reasonable revisions are included at no extra cost.")

# ── 4. Term ───────────────────────────────────────────────────────────────
h("4.  Term")
para(
    "This Agreement commences on 21 April 2026 and has an initial term of three (3) months, "
    "expiring on 21 July 2026. Either Party may extend by mutual written agreement "
    "before the expiry date. If no extension is agreed, the Agreement terminates automatically "
    "at the end of the initial term, with payment for all completed and accepted work to that date.")

# ── 5. Timeline and Delivery ──────────────────────────────────────────────
h("5.  Timeline and Delivery")
para(
    "Batch deadlines will be agreed in writing (email) at the start of each batch. "
    "The Consultant shall notify Bristlepine promptly if a deadline cannot be met. "
    "Repeated missed deadlines without notice are grounds for termination under Clause 6.")

# ── 6. Termination ────────────────────────────────────────────────────────
h("6.  Termination for Convenience")
para(
    "Either Party may terminate this Agreement at any time and for any reason by sending "
    "written notice by email. No notice period is required \u2014 termination is effective "
    "on the date the email is sent. Upon termination:")
bul("Bristlepine will pay in full for all completed and accepted batches.")
bul("For any partially completed batch, Bristlepine will pay a pro-rata amount based on "
    "the number of fully coded and accepted papers delivered.")
bul("The Consultant will deliver all work in progress to Bristlepine within 3 business "
    "days of the termination notice.")

# ── 7. Acceptance of Work ─────────────────────────────────────────────────
h("7.  Acceptance of Work")
para(
    "Deliverables will be reviewed by Bristlepine. Work is considered accepted only when "
    "Bristlepine provides written confirmation by email. If work does not meet the required "
    "standards, Bristlepine may request revisions. The Consultant agrees to provide "
    "reasonable revisions at no additional cost within 3 business days.")

# ── 8. Confidentiality ────────────────────────────────────────────────────
h("8.  Confidentiality")
para(
    "The Consultant agrees to maintain strict confidentiality regarding all proprietary or "
    "confidential information, including paper contents under embargo, internal codebook "
    "drafts, and unpublished project data. This obligation continues indefinitely after "
    "termination of this Agreement.")
para(
    "This clause does not restrict the Consultant from listing this engagement on their "
    "CV or academic profile using the title and dates.")

# ── 9. Data Security ──────────────────────────────────────────────────────
h("9.  Data Security")
para(
    "The Consultant shall take reasonable measures to protect all confidential information "
    "from unauthorised access, use, or disclosure, including using secure storage for any "
    "downloaded papers and deleting project data from personal devices upon request or "
    "termination of this Agreement.")

# ── 10. Intellectual Property ─────────────────────────────────────────────
h("10.  Intellectual Property")
para(
    "All work product, coding outputs, notes, reconciliation records, and related "
    "intellectual property created under this Agreement shall be the property of "
    "Bristlepine Resilience Consultants LLC and ILRI/CGIAR upon delivery. "
    "If any work product is not legally considered work made for hire, the Consultant "
    "hereby assigns all rights, title, and interest to Bristlepine.")
para(
    "Acknowledgement: The Consultant will be named in the Acknowledgements of all published "
    "outputs. Where contribution meets ICMJE authorship criteria, co-authorship will be offered.")

# ── 11. No Subcontracting ─────────────────────────────────────────────────
h("11.  No Subcontracting")
para(
    "The Consultant shall not delegate, subcontract, or assign any portion of the Services "
    "to a third party without prior written approval from Bristlepine.")

# ── 12. No Use of AI Tools ────────────────────────────────────────────────
h("12.  No Use of Artificial Intelligence Tools")
para(
    "The Consultant shall not use artificial intelligence tools, large language models, "
    "automated text generation systems (including but not limited to ChatGPT, Claude, "
    "Copilot, Gemini, or similar technologies), or any other automated coding assistance "
    "to complete any coding task under this Agreement.")
para(
    "All coding must be completed by the Consultant personally through direct reading of "
    "the source papers and the exercise of independent human judgement. This engagement "
    "is explicitly designed as a human-coded process for methodological integrity and "
    "peer-review purposes.")
para(
    "Any use of AI tools for coding purposes will be considered a material breach of this "
    "Agreement and grounds for immediate termination without payment for the affected batch.")

# ── 13. ILRI / Funder Compliance ──────────────────────────────────────────
h("13.  ILRI / Funder Compliance and Quality Standards")
para(
    "This engagement is conducted in support of a project funded by and accountable to the "
    "International Livestock Research Institute (ILRI) / CGIAR. Bristlepine Resilience "
    "Consultants LLC reserves the right to terminate or suspend this Agreement, with "
    "payment for completed and accepted work only, in the event that:")
bul(
    "ILRI or its designated representatives raise substantive objections to the "
    "Consultant\u2019s work quality, methodology, or approach that cannot be resolved "
    "through revision within a reasonable timeframe; or")
bul(
    "ILRI requires a change in process, personnel, or quality standards that makes "
    "continuation of this Agreement under its current terms impracticable.")
para(
    "The Consultant acknowledges that ILRI\u2019s methodological standards and quality "
    "requirements take precedence, and that Bristlepine cannot guarantee continuation "
    "of this engagement beyond any individual batch if ILRI exercises this discretion. "
    "Bristlepine will provide reasonable written notice of any such decision.")

# ── 14. Warranties ────────────────────────────────────────────────────────
h("14.  Warranties")
para("The Consultant warrants that:")
bul("Deliverables will be original, free of defects, and meet specified requirements.")
bul("Deliverables will not infringe on the rights of any third parties.")
bul("The Consultant has the right and authority to enter into this Agreement.")

# ── 15. Indemnification ───────────────────────────────────────────────────
h("15.  Indemnification")
para(
    "The Consultant agrees to indemnify and hold harmless Bristlepine Resilience Consultants LLC "
    "from any claims, damages, or liabilities arising from the Consultant\u2019s breach of this "
    "Agreement, negligence, or wilful misconduct.")

# ── 16. Conflicts of Interest ─────────────────────────────────────────────
h("16.  Conflicts of Interest and Professional Conduct")
para(
    "The Consultant shall perform the Services in a professional and lawful manner and shall "
    "avoid any actual or potential conflict of interest. The Consultant shall promptly disclose "
    "to Bristlepine any situation that may constitute a conflict of interest.")

# ── 17. Method of Notice ──────────────────────────────────────────────────
h("17.  Method of Notice")
para(
    "All notices, approvals, invoices, and instructions under this Agreement shall be sent "
    "by email. Notices are deemed received on the date sent.")
bul("Client:       info@bristlep.com")
bul("Consultant:   [Consultant email]")

# ── 18. Governing Law ─────────────────────────────────────────────────────
h("18.  Governing Law")
para(
    "This Agreement is governed by the laws of the Province of Ontario, Canada, without "
    "regard to conflict-of-law principles. Any disputes will be resolved by good-faith "
    "negotiation before any further escalation.")

# ── 19. Entire Agreement ──────────────────────────────────────────────────
h("19.  Entire Agreement")
para(
    "This Agreement, including all Exhibits and Appendices, constitutes the entire "
    "understanding between the Parties and supersedes all prior discussions. "
    "Any amendments must be agreed in writing by email by both Parties.")

blank()

# ── Signatures ────────────────────────────────────────────────────────────
h("Signatures", level=2)
para(
    "By signing below, both Parties confirm they have read, understood, and agreed to the "
    "terms of this Agreement.")
blank()

sg = doc.add_table(rows=6, cols=2)
sg.style = "TableNormal"
sig_data = [
    ("For Bristlepine Resilience Consultants LLC", "Consultant"),
    ("Name: Zarrar Khan",   "Name: [Full name]"),
    ("Title: Partner",      "Title: [Title]"),
    ("Signature:",          "Signature:"),
    ("Date: [Date]",        "Date: [Date]"),
    ("Email: info@bristlep.com", "Email: [Consultant email]"),
]
for ri, (left, right) in enumerate(sig_data):
    for ci, val in enumerate([left, right]):
        cell = sg.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        parts = _PLACEHOLDER.split(val)
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            run.bold = (ri == 0)
            run.font.size = Pt(10)
            if _PLACEHOLDER.fullmatch(part):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    sg.rows[ri].height = Cm(1.1)
tbl_move(sg)
blank()

# ══════════════════════════════════════════════════════════════════════════
# EXHIBIT A — Scope, Deliverables, and Payment Schedule
# ══════════════════════════════════════════════════════════════════════════
page_break()

h("Exhibit A \u2014 Scope, Deliverables, and Payment Schedule", level=1)

h("Project", level=3)
para(
    "ILRI/CGIAR Systematic Map D5.6 \u2014 Climate adaptation measurement for smallholder "
    "agricultural producers in LMICs. Outputs will be published on Zenodo and CGSpace "
    "and submitted to a peer-reviewed journal.")

h("Role", level=3)
para(
    "Evidence coder \u2014 data extraction phase. The Consultant reads full-text papers and "
    "codes each against 18 structured fields in the project codebook, working independently.")

h("Deliverables and Fees", level=3)

exA = doc.add_table(rows=5, cols=5)
exA.style = "TableNormal"
for i, hdr in enumerate(["#", "Deliverable", "Deadline", "Est. hours", "Fee (PKR)"]):
    c = exA.rows[0].cells[i]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)

exA_data = [
    ("1",
     f"Calibration batch ({CALIB_PAPERS} papers) \u2014 coding_ft_r1a_[initials].csv",
     "[Date]",
     f"{ch_lo}\u2013{ch_hi} hrs",
     f"Rs. {c_lo:,}\u2013{c_hi:,}"),
    ("2",
     f"Batch 1 (~{BATCH_PAPERS} papers) \u2014 coding_[round]_[initials].csv",
     "[Date]",
     f"{bh_lo}\u2013{bh_hi} hrs",
     f"Rs. {b_lo:,}\u2013{b_hi:,}"),
    ("3",
     f"Batch 2 (~{BATCH_PAPERS} papers)",
     "[Date]",
     f"{bh_lo}\u2013{bh_hi} hrs",
     f"Rs. {b_lo:,}\u2013{b_hi:,}"),
    ("4+",
     "Extension batches (if agreed in writing)",
     "[TBD]",
     "[TBD]",
     "[TBD]"),
]
for ri, row in enumerate(exA_data, 1):
    for ci, val in enumerate(row):
        cell = exA.rows[ri].cells[ci]
        p    = cell.paragraphs[0]
        for part in _PLACEHOLDER.split(val):
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(10)
            if _PLACEHOLDER.fullmatch(part):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
tbl_move(exA)
blank()

h("Acceptance Criteria", level=3)
bul("All 16 fields completed for every paper (or marked not_reported with a note).")
bul("Coding sheet delivered in CSV format, uploaded to the shared Google Drive folder.")
bul("Flagged ambiguous cases have explanatory notes in the notes column.")

h("Agreed Hourly Rate", level=3)
p = doc.add_paragraph(style="Normal")
p.paragraph_format.space_after = Pt(5)
r1 = p.add_run(f"Rs. {HOURLY_RATE_PKR}/hr (Pakistani Rupees) \u2014 rate initialled by both Parties: "
               "Client ____  Consultant ____")
_move(p._element)

h("Rate Review", level=3)
para(
    "The hourly rate is an initial estimate and may be reviewed by mutual written agreement "
    "before each new batch. Any rate change applies only to future batches, not completed work.")

blank()

# ── Exhibit A Signatures ───────────────────────────────────────────────────
para("Agreed and accepted:", bold=True)
blank()
exSig = doc.add_table(rows=4, cols=2)
exSig.style = "TableNormal"
exSig_data = [
    ("For Bristlepine Resilience Consultants LLC:", "Consultant:"),
    ("Name: Zarrar Khan",  "Name: [Full name]"),
    ("Signature:",         "Signature:"),
    ("Date: [Date]",       "Date: [Date]"),
]
for ri, (left, right) in enumerate(exSig_data):
    for ci, val in enumerate([left, right]):
        cell = exSig.rows[ri].cells[ci]
        p    = cell.paragraphs[0]
        for part in _PLACEHOLDER.split(val):
            if not part:
                continue
            run = p.add_run(part)
            run.bold = (ri == 0)
            run.font.size = Pt(10)
            if _PLACEHOLDER.fullmatch(part):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    exSig.rows[ri].height = Cm(1.0)
tbl_move(exSig)

# ══════════════════════════════════════════════════════════════════════════
# APPENDIX B — Bank / Payment Information Form
# ══════════════════════════════════════════════════════════════════════════
page_break()

h("Appendix B \u2014 Payment / Bank Information Form", level=1)
para(
    "Please complete this form and return it to info@bristlep.com before your first invoice. "
    "This information will be used solely to process payment under Contract ILRI01-JD-RA-2026. "
    "All information is kept confidential.", italic=True)
blank()

# Section A: Personal details
h("Section A \u2014 Payee Details", level=3)
bank_fields_A = [
    ("Full legal name (as on bank account):", "[Full legal name]"),
    ("National ID / CNIC / Passport number:", "[ID number]"),
    ("Country of residence:",                  "[Country]"),
    ("Mailing address (if required):",         "[Address]"),
    ("Email for payment confirmation:",         "[Email]"),
]

bfA = doc.add_table(rows=len(bank_fields_A), cols=2)
bfA.style = "TableNormal"
for ri, (label, placeholder) in enumerate(bank_fields_A):
    lc = bfA.rows[ri].cells[0]
    rc = bfA.rows[ri].cells[1]
    lrun = lc.paragraphs[0].add_run(label)
    lrun.font.size = Pt(10)
    lrun.bold = True
    rrun = rc.paragraphs[0].add_run(placeholder)
    rrun.font.size = Pt(10)
    rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    bfA.rows[ri].height = Cm(0.9)
tbl_move(bfA)
blank()

# Section B: Bank details
h("Section B \u2014 Bank Account Details", level=3)
bank_fields_B = [
    ("Bank name:",                            "[Bank name]"),
    ("Bank branch / address:",                "[Branch address]"),
    ("Account title:",                        "[Account title]"),
    ("Account number:",                       "[Account number]"),
    ("IBAN (if applicable):",                 "[IBAN]"),
    ("SWIFT / BIC code:",                     "[SWIFT/BIC]"),
    ("Account type (savings / current):",     "[Account type]"),
    ("Currency of account:",                  "[PKR / USD / other]"),
]

bfB = doc.add_table(rows=len(bank_fields_B), cols=2)
bfB.style = "TableNormal"
for ri, (label, placeholder) in enumerate(bank_fields_B):
    lc = bfB.rows[ri].cells[0]
    rc = bfB.rows[ri].cells[1]
    lrun = lc.paragraphs[0].add_run(label)
    lrun.font.size = Pt(10)
    lrun.bold = True
    rrun = rc.paragraphs[0].add_run(placeholder)
    rrun.font.size = Pt(10)
    rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    bfB.rows[ri].height = Cm(0.9)
tbl_move(bfB)
blank()

# Section C: Intermediary bank (optional)
h("Section C \u2014 Intermediary / Correspondent Bank (if required for international wire)", level=3)
bank_fields_C = [
    ("Intermediary bank name:",   "[Bank name or N/A]"),
    ("Intermediary SWIFT / BIC:", "[SWIFT/BIC or N/A]"),
    ("Routing number:",           "[Routing number or N/A]"),
]

bfC = doc.add_table(rows=len(bank_fields_C), cols=2)
bfC.style = "TableNormal"
for ri, (label, placeholder) in enumerate(bank_fields_C):
    lc = bfC.rows[ri].cells[0]
    rc = bfC.rows[ri].cells[1]
    lrun = lc.paragraphs[0].add_run(label)
    lrun.font.size = Pt(10)
    lrun.bold = True
    rrun = rc.paragraphs[0].add_run(placeholder)
    rrun.font.size = Pt(10)
    rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    bfC.rows[ri].height = Cm(0.9)
tbl_move(bfC)
blank()

para(
    "I confirm that the above information is accurate and I authorise Bristlepine Resilience "
    "Consultants LLC to use it for payment processing under this contract.", italic=True)
blank()

bSig = doc.add_table(rows=3, cols=2)
bSig.style = "TableNormal"
bSig_data = [
    ("Consultant name:", "[Full name]"),
    ("Signature:",       ""),
    ("Date:",            "[Date]"),
]
for ri, (label, val) in enumerate(bSig_data):
    for ci, txt in enumerate([label, val] if ci == 0 else ["", val]
                              for ci in range(2)):
        # simpler: just do both cells
        pass
    lc = bSig.rows[ri].cells[0]
    rc = bSig.rows[ri].cells[1]
    lc.paragraphs[0].add_run(label).font.size = Pt(10)
    if val:
        rrun = rc.paragraphs[0].add_run(val)
        rrun.font.size = Pt(10)
        if _PLACEHOLDER.fullmatch(val):
            rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    bSig.rows[ri].height = Cm(1.0)
tbl_move(bSig)

# ══════════════════════════════════════════════════════════════════════════
# APPENDIX C — Invoice Template
# ══════════════════════════════════════════════════════════════════════════
page_break()

h("Appendix C \u2014 Invoice Template", level=1)
para(
    "Complete this form for each batch and send to info@bristlep.com. "
    "Keep a copy for your records.", italic=True)
blank()

# Header block
invH = doc.add_table(rows=1, cols=2)
invH.style = "TableNormal"
lc = invH.cell(0, 0)
rc = invH.cell(0, 1)

lc.paragraphs[0].add_run("FROM:").bold = True
for line, ph in [
    ("Name:",    "[Full name]"),
    ("Address:", "[Address]"),
    ("Email:",   "[Consultant email]"),
    ("Tax ID / NTN (if applicable):", "[Tax ID or N/A]"),
]:
    p = lc.add_paragraph()
    r1 = p.add_run(f"{line}  ")
    r1.font.size = Pt(10)
    r2 = p.add_run(ph)
    r2.font.size = Pt(10)
    if _PLACEHOLDER.fullmatch(ph):
        r2.font.highlight_color = WD_COLOR_INDEX.YELLOW

rc.paragraphs[0].add_run("TO:").bold = True
for line in [
    "Bristlepine Resilience Consultants LLC",
    "Email: info@bristlep.com",
    "Contract Ref: ILRI01-JD-RA-2026",
]:
    p = rc.add_paragraph()
    r = p.add_run(line)
    r.font.size = Pt(10)

tbl_move(invH)
blank()

# Invoice metadata
meta = doc.add_table(rows=3, cols=2)
meta.style = "TableNormal"
meta_data = [
    ("Invoice number:",  "[INV-ILRI01-001]"),
    ("Invoice date:",    "[Date]"),
    ("Payment due:",     "[Date + 14 days]"),
]
for ri, (label, val) in enumerate(meta_data):
    lrun = meta.rows[ri].cells[0].paragraphs[0].add_run(label)
    lrun.bold = True; lrun.font.size = Pt(10)
    rrun = meta.rows[ri].cells[1].paragraphs[0].add_run(val)
    rrun.font.size = Pt(10)
    rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    meta.rows[ri].height = Cm(0.9)
tbl_move(meta)
blank()

# Line items table
h("Services Rendered", level=3)
inv = doc.add_table(rows=5, cols=5)
inv.style = "TableNormal"
for ci, hdr in enumerate(["Batch", "Papers coded", "Hours worked", f"Rate (Rs./hr)", "Amount (PKR)"]):
    c = inv.rows[0].cells[ci]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)

inv_rows = [
    ("[Batch name e.g. Calibration R1a]", "[n]", "[hrs]", f"Rs. {HOURLY_RATE_PKR}", "[Total]"),
    ("", "", "", "", ""),
    ("", "", "", "", ""),
    ("", "", "TOTAL", "", "[Grand total]"),
]
for ri, row in enumerate(inv_rows, 1):
    for ci, val in enumerate(row):
        cell = inv.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        for part in _PLACEHOLDER.split(val):
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.bold = (ri == 4 and ci in (2, 4))
            if _PLACEHOLDER.fullmatch(part):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
tbl_move(inv)
blank()

h("Payment Details", level=3)
para("Please transfer to the account listed in Appendix B on file with Bristlepine.")
para("Reference / narration to include on transfer: [Contract ref + consultant name + invoice number]")
blank()

para(
    "I certify that the hours and deliverables listed above are accurate and that the work "
    "described has been completed to the standard required by the contract.", italic=True)
blank()

cSig = doc.add_table(rows=3, cols=2)
cSig.style = "TableNormal"
cSig_data = [
    ("Consultant name:", "[Full name]"),
    ("Signature:", ""),
    ("Date:", "[Date]"),
]
for ri, (label, val) in enumerate(cSig_data):
    lc = cSig.rows[ri].cells[0]
    rc = cSig.rows[ri].cells[1]
    lc.paragraphs[0].add_run(label).font.size = Pt(10)
    if val:
        rrun = rc.paragraphs[0].add_run(val)
        rrun.font.size = Pt(10)
        if _PLACEHOLDER.fullmatch(val):
            rrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
    cSig.rows[ri].height = Cm(1.0)
tbl_move(cSig)

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Rate: Rs. {HOURLY_RATE_PKR}/hr  |  {int(HOURS_MIN*60)}–{int(HOURS_MAX*60)} min/paper  |  ~{int(HOURS_EST*60)} min expected")
print(f"Calibration ({CALIB_PAPERS} papers): Rs. {c_lo:,}–{c_hi:,}  (~Rs. {c_est:,})")
print(f"Standard batch (~{BATCH_PAPERS} papers): Rs. {b_lo:,}–{b_hi:,}  (~Rs. {b_est:,})")
