#!/usr/bin/env python3
"""
step2_retrieve.py

Step 2 (retrieve TOTAL__ALL records; streaming):
  - Avoid deep paging limits by slicing TOTAL query into PUBYEAR ranges (<= DEEP_PAGING_LIMIT)
  - If a single PUBYEAR slice still exceeds the deep paging limit, sub-slice within-year by:
      1) SUBJAREA() (may overlap -> final dedupe)
      2) SRCTYPE() (disjoint within a SUBJAREA bucket)
  - Stream each slice to CSV, then concatenate -> outputs/step2/step2_total_records.csv
  - Final combined CSV is deduped on record_key to guard against overlap.

Additional outputs:
  - outputs/step2/step2_total_records.docx
  - outputs/step2/step2_total_records.bib
  - outputs/step2/step2_total_records.ris

Designed to work with scripts/run.py "new runner" style:
  - provides run_step1b_retrieve_total(config)
"""

from __future__ import annotations

import os
import json
import time
import re
from typing import Dict, List, Optional, Tuple

import pandas as pd
import requests
from dotenv import load_dotenv
from tqdm import tqdm

from docx import Document
from docx.shared import Pt

import config as cfg
from utils import (
    ScopusAuth,
    scopus_headers,
    request_with_retries,
    read_yaml,
    normalize_doi,
    clean_text,
    record_key,
    utc_now_iso,
)


DEFAULT_USE_POST_FOR_SEARCH = 1
DEFAULT_VIEW = "STANDARD"

DEFAULT_COUNT_PER_PAGE_RETRIEVE = 25
DEFAULT_STEP2_SLEEP_S = 0.15
DEFAULT_DEEP_PAGING_LIMIT = 5000
DEFAULT_PUBYEAR_MIN = 1990
DEFAULT_PUBYEAR_MAX = 2025
DEFAULT_MAX_RESULTS_TOTAL = None


# Scopus subject area codes (commonly used by SUBJAREA in advanced query).
# Note: items may be tagged with multiple subject areas -> overlap is possible.
SCOPUS_SUBJAREAS = [
    "AGRI", "ARTS", "BIOC", "BUSI", "CENG", "CHEM", "COMP", "DECI", "DENT", "EART",
    "ECON", "ENER", "ENGI", "ENVI", "HEAL", "IMMU", "MATE", "MATH", "MEDI", "NEUR",
    "NURS", "PHAR", "PHYS", "PSYC", "SOCI", "VETE", "MULT",
]

# Scopus source type codes (SRCTYPE in advanced query). These are disjoint.
# If any are invalid for your tenant, you'll see it immediately during counting; remove the bad code.
SCOPUS_SRCTYPES = ["j", "p", "k", "b", "d", "t", "r"]  # journal, proceedings, book series, book, trade, etc.


def _sanitize_csv_tail(path: str) -> None:
    """
    If the file doesn't end with a newline, truncate to the last newline.
    Protects against a crash mid-write leaving a partial last row.
    """
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        return

    with open(path, "rb") as f:
        f.seek(0, os.SEEK_END)
        size = f.tell()
        if size == 0:
            return
        f.seek(-1, os.SEEK_END)
        last = f.read(1)

        if last == b"\n":
            return

        chunk = min(size, 1024 * 1024)
        f.seek(-chunk, os.SEEK_END)
        buf = f.read(chunk)
        idx = buf.rfind(b"\n")
        if idx == -1:
            new_size = 0
        else:
            new_size = size - (chunk - (idx + 1))

    with open(path, "ab") as f:
        f.truncate(new_size)


def _count_csv_data_rows(path: str) -> int:
    """
    Fast-ish line count. Assumes one record per line (pandas CSV output).
    Returns number of data rows (excluding header).
    """
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        return 0

    n = 0
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        header = f.readline()
        if not header:
            return 0
        for _ in f:
            n += 1
    return n


def _save_json(path: str, obj) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2, ensure_ascii=False)
    os.replace(tmp, path)


def _build_queries(search_cfg: dict) -> Tuple[Dict[str, str], Dict[str, str], str]:
    field = (search_cfg.get("field", "TITLE-ABS-KEY") or "TITLE-ABS-KEY").strip()
    elements = search_cfg.get("elements", {}) or {}

    subgroup_queries: Dict[str, str] = {}
    element_queries: Dict[str, str] = {}
    element_raw_or: Dict[str, str] = {}

    for element_key, subgroups in elements.items():
        raw_exprs: List[str] = []
        for sub_key, expr in (subgroups or {}).items():
            expr = str(expr).strip()
            qname = f"{element_key}__{sub_key}"
            subgroup_queries[qname] = f"{field}({expr})"
            raw_exprs.append(f"({expr})")

        if raw_exprs:
            or_raw = " OR ".join(raw_exprs)
            element_raw_or[element_key] = f"({or_raw})"
            element_queries[f"{element_key}__ALL"] = f"{field}({element_raw_or[element_key]})"

    combined_raw = " AND ".join([f"({v})" for v in element_raw_or.values()])
    combined_query = f"{field}({combined_raw})" if combined_raw else ""
    return subgroup_queries, element_queries, combined_query


def scopus_count_only(
    auth: ScopusAuth,
    scopus_search_url: str,
    query: str,
    use_post: int,
    view: str,
    count_per_page: int = 1,
) -> Tuple[int, dict]:
    headers = scopus_headers(auth)
    meta = {"query": query, "count_per_page": count_per_page, "view": view, "rate_headers_last": {}}

    with requests.Session() as session:
        params = {"query": query, "start": "0", "count": str(count_per_page), "view": view}
        method = "POST" if use_post else "GET"
        data, rate = request_with_retries(session, method, scopus_search_url, headers, params=params, data=params)
        meta["rate_headers_last"] = rate

    sr = (data or {}).get("search-results", {}) or {}
    total = int(sr.get("opensearch:totalResults", "0"))
    return total, meta


def _year_from_coverdate(s: str) -> str:
    if not isinstance(s, str):
        return ""
    m = re.search(r"\b(19|20)\d{2}\b", s)
    return m.group(0) if m else ""


def _build_scopus_citation_raw(creator: str, year: str, title: str, journal: str, doi_core: str) -> str:
    parts = []
    if creator:
        parts.append(str(creator).strip())
    if year:
        parts.append(f"({year})")
    if title:
        parts.append(f"'{str(title).strip().rstrip('. ')}'.")
    if journal:
        parts.append(f"{str(journal).strip().rstrip('. ')}.")
    if doi_core:
        parts.append(f"https://doi.org/{doi_core}")
    return " ".join([p for p in parts if p]).strip()


def _extract_entry_row(e: dict) -> dict:
    dcid = e.get("dc:identifier")
    scopus_id = dcid.split(":", 1)[1].strip() if isinstance(dcid, str) and dcid.startswith("SCOPUS_ID:") else None

    title = e.get("dc:title")
    cover = e.get("prism:coverDate")
    journal = e.get("prism:publicationName")

    doi_raw = e.get("prism:doi")
    doi_raw = doi_raw.strip() if isinstance(doi_raw, str) else None
    doi_core = normalize_doi(doi_raw) if doi_raw else ""

    creator = e.get("dc:creator")
    year = _year_from_coverdate(cover)

    citation_raw = _build_scopus_citation_raw(
        creator=str(creator).strip() if isinstance(creator, str) else "",
        year=year,
        title=str(title).strip() if isinstance(title, str) else "",
        journal=str(journal).strip() if isinstance(journal, str) else "",
        doi_core=doi_core,
    )

    rk = record_key(doi_core, str(title or ""))

    return {
        "citation_raw": citation_raw,
        "record_key": rk,
        "title": title,
        "coverDate": cover,
        "publicationName": journal,
        "doi": doi_core or None,
        "eid": e.get("eid"),
        "scopus_id": scopus_id,
        "citedby_count": e.get("citedby-count"),
        "prism_url": e.get("prism:url"),
    }


def _with_pubyear_range(base_query: str, y0: int, y1: int) -> str:
    return f"({base_query}) AND PUBYEAR > {y0 - 1} AND PUBYEAR < {y1 + 1}"


def _with_pubyear_missing_or_outside(base_query: str, y0: int, y1: int) -> str:
    return f"({base_query}) AND NOT (PUBYEAR > {y0 - 1} AND PUBYEAR < {y1 + 1})"


def _safe_tag(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_]+", "_", s.strip())
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "SLICE"


def _plan_year_slices(
    auth: ScopusAuth,
    scopus_search_url: str,
    base_query: str,
    y0: int,
    y1: int,
    *,
    deep_paging_limit: int,
    use_post: int,
    view: str,
) -> List[Tuple[int, int, int]]:
    total, _ = scopus_count_only(
        auth,
        scopus_search_url,
        _with_pubyear_range(base_query, y0, y1),
        use_post=use_post,
        view=view,
        count_per_page=1,
    )
    if total <= deep_paging_limit or y0 == y1:
        return [(y0, y1, total)]
    mid = (y0 + y1) // 2
    return (
        _plan_year_slices(
            auth, scopus_search_url, base_query, y0, mid,
            deep_paging_limit=deep_paging_limit, use_post=use_post, view=view
        )
        + _plan_year_slices(
            auth, scopus_search_url, base_query, mid + 1, y1,
            deep_paging_limit=deep_paging_limit, use_post=use_post, view=view
        )
    )


def _plan_within_year_subjarea_slices(
    auth: ScopusAuth,
    scopus_search_url: str,
    base_q: str,
    year: int,
    *,
    deep_paging_limit: int,
    use_post: int,
    view: str,
) -> List[dict]:
    """
    Create SUBJAREA sub-slices for a single year slice that is still > deep paging limit.
    Adds a remainder slice for records not matching any SUBJAREA codes in our list.
    Note: SUBJAREA can overlap; we will dedupe final output on record_key.
    """
    parent_q = _with_pubyear_range(base_q, year, year)
    parent_total, _ = scopus_count_only(auth, scopus_search_url, parent_q, use_post=use_post, view=view, count_per_page=1)

    slices: List[dict] = []
    covered_exprs: List[str] = []
    covered_sum = 0

    for code in SCOPUS_SUBJAREAS:
        expr = f"SUBJAREA({code})"
        q = f"({parent_q}) AND {expr}"
        t, _ = scopus_count_only(auth, scopus_search_url, q, use_post=use_post, view=view, count_per_page=1)
        if t <= 0:
            continue
        slices.append(
            {
                "tag": _safe_tag(f"PUBYEAR_{year}_{year}__SUBJAREA_{code}"),
                "query": q,
                "expected": int(t),
                "kind": "PUBYEAR_SUBJAREA",
                "year": year,
                "subjarea": code,
            }
        )
        covered_exprs.append(expr)
        covered_sum += int(t)

    # remainder: not in any of the subject area codes we enumerated
    if covered_exprs:
        or_part = " OR ".join(covered_exprs)
        remainder_q = f"({parent_q}) AND NOT ({or_part})"
        remainder_t, _ = scopus_count_only(auth, scopus_search_url, remainder_q, use_post=use_post, view=view, count_per_page=1)
        if remainder_t > 0:
            slices.append(
                {
                    "tag": _safe_tag(f"PUBYEAR_{year}_{year}__SUBJAREA_REMAINDER"),
                    "query": remainder_q,
                    "expected": int(remainder_t),
                    "kind": "PUBYEAR_SUBJAREA_REMAINDER",
                    "year": year,
                    "subjarea": None,
                }
            )

    # We don't require covered_sum == parent_total because SUBJAREA can overlap (multi-tagging).
    # Dedupe later avoids double counting.
    return slices


def _plan_within_slice_srctype_slices(
    auth: ScopusAuth,
    scopus_search_url: str,
    parent_slice: dict,
    *,
    deep_paging_limit: int,
    use_post: int,
    view: str,
) -> List[dict]:
    """
    Further sub-slice an oversized slice by SRCTYPE() into disjoint buckets.
    Adds a remainder bucket for anything not matching our SRCTYPE list.
    """
    parent_q = parent_slice["query"]
    parent_total = int(parent_slice["expected"])

    slices: List[dict] = []
    covered_exprs: List[str] = []
    covered_sum = 0

    for code in SCOPUS_SRCTYPES:
        expr = f"SRCTYPE({code})"
        q = f"({parent_q}) AND {expr}"
        t, _ = scopus_count_only(auth, scopus_search_url, q, use_post=use_post, view=view, count_per_page=1)
        if t <= 0:
            continue
        slices.append(
            {
                "tag": _safe_tag(f"{parent_slice['tag']}__SRCTYPE_{code}"),
                "query": q,
                "expected": int(t),
                "kind": "SRCTYPE",
                "parent_tag": parent_slice["tag"],
                "srctype": code,
            }
        )
        covered_exprs.append(expr)
        covered_sum += int(t)

    if covered_exprs:
        or_part = " OR ".join(covered_exprs)
        remainder_q = f"({parent_q}) AND NOT ({or_part})"
        remainder_t, _ = scopus_count_only(auth, scopus_search_url, remainder_q, use_post=use_post, view=view, count_per_page=1)
        if remainder_t > 0:
            slices.append(
                {
                    "tag": _safe_tag(f"{parent_slice['tag']}__SRCTYPE_REMAINDER"),
                    "query": remainder_q,
                    "expected": int(remainder_t),
                    "kind": "SRCTYPE_REMAINDER",
                    "parent_tag": parent_slice["tag"],
                    "srctype": None,
                }
            )

    # SRCTYPE is disjoint, so covered_sum + remainder_t should be close to parent_total.
    # We won't hard-fail; just return what we planned.
    return slices


def _expand_slices(
    auth: ScopusAuth,
    scopus_search_url: str,
    base_q: str,
    year_slices: List[Tuple[int, int, int]],
    *,
    deep_paging_limit: int,
    use_post: int,
    view: str,
) -> List[dict]:
    """
    Expand year slices. If a slice is a single year and still > limit, expand within-year.
    If any within-year slice still > limit, expand that further by SRCTYPE.
    """
    expanded: List[dict] = []

    for y0, y1, t in year_slices:
        if y0 == y1 and t > deep_paging_limit:
            year = y0
            subj_slices = _plan_within_year_subjarea_slices(
                auth,
                scopus_search_url,
                base_q,
                year,
                deep_paging_limit=deep_paging_limit,
                use_post=use_post,
                view=view,
            )

            for ss in subj_slices:
                if int(ss["expected"]) > deep_paging_limit:
                    # further expand by SRCTYPE (disjoint)
                    sr_slices = _plan_within_slice_srctype_slices(
                        auth,
                        scopus_search_url,
                        ss,
                        deep_paging_limit=deep_paging_limit,
                        use_post=use_post,
                        view=view,
                    )
                    expanded.extend(sr_slices)
                else:
                    expanded.append(ss)
        else:
            expanded.append(
                {
                    "tag": _safe_tag(f"PUBYEAR_{y0}_{y1}"),
                    "query": _with_pubyear_range(base_q, y0, y1),
                    "expected": int(t),
                    "kind": "PUBYEAR",
                    "year_start": y0,
                    "year_end": y1,
                }
            )

    return expanded


def _concat_csvs(csv_paths: List[str], out_csv: str) -> int:
    os.makedirs(os.path.dirname(out_csv), exist_ok=True)
    if os.path.exists(out_csv):
        os.remove(out_csv)

    wrote_header = False
    total_rows = 0

    with open(out_csv, "w", encoding="utf-8", newline="") as out_f:
        for p in csv_paths:
            if not os.path.exists(p) or os.path.getsize(p) == 0:
                continue
            with open(p, "r", encoding="utf-8", errors="ignore") as in_f:
                for i, line in enumerate(in_f):
                    if i == 0:
                        if wrote_header:
                            continue
                        wrote_header = True
                        out_f.write(line)
                        continue
                    out_f.write(line)
                    total_rows += 1
    return total_rows


def _dedupe_csv_in_place(path: str, key_col: str = "record_key") -> int:
    """
    Dedupe the combined CSV in place by record_key.
    Returns number of rows after dedupe (data rows, excluding header).
    """
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        return 0
    df = pd.read_csv(path, engine="python", on_bad_lines="skip")
    if df.empty:
        return 0
    if key_col in df.columns:
        df = df.drop_duplicates(subset=[key_col], keep="first")
    else:
        df = df.drop_duplicates(keep="first")
    df.to_csv(path, index=False)
    return int(len(df))


def scopus_retrieve_stream_to_csv_start(
    auth: ScopusAuth,
    scopus_search_url: str,
    *,
    query: str,
    out_csv: str,
    count_per_page: int,
    sleep_s: float,
    max_results: Optional[int],
    use_post: int,
    view: str,
) -> Tuple[int, int, dict]:
    headers = scopus_headers(auth)
    method = "POST" if use_post else "GET"

    os.makedirs(os.path.dirname(out_csv), exist_ok=True)

    # --- RESUME LOGIC ---
    if os.path.exists(out_csv) and os.path.getsize(out_csv) > 0:
        _sanitize_csv_tail(out_csv)
        retrieved = _count_csv_data_rows(out_csv)
        resume_mode = True
    else:
        retrieved = 0
        resume_mode = False
        if os.path.exists(out_csv):
            os.remove(out_csv)

    total_reported = None

    meta = {
        "query": query,
        "view": view,
        "count_per_page": count_per_page,
        "sleep_s": sleep_s,
        "max_results": max_results,
        "rate_headers_last": {},
        "started_utc": utc_now_iso(),
        "resume_mode": bool(resume_mode),
        "resume_existing_rows": int(retrieved),
    }

    pbar = None

    with requests.Session() as session:
        while True:
            if max_results is not None and retrieved >= max_results:
                break

            params = {"query": query, "start": str(retrieved), "count": str(count_per_page), "view": view}

            data, rate = request_with_retries(session, method, scopus_search_url, headers, params=params, data=params)
            meta["rate_headers_last"] = rate

            sr = (data or {}).get("search-results", {}) or {}

            if total_reported is None:
                total_reported = int(sr.get("opensearch:totalResults", "0"))
                meta["total_reported_by_scopus"] = total_reported
                pbar_total = total_reported if max_results is None else min(total_reported, max_results)
                pbar = tqdm(total=pbar_total, desc="[step2] Retrieving", unit="rec", initial=min(retrieved, pbar_total))
                print(f"[step2] Total reported by Scopus: {total_reported:,}")
                if resume_mode:
                    print(f"[step2] RESUME: {retrieved:,} rows already in {out_csv}")

            entries = sr.get("entry", []) or []
            if not entries:
                break

            rows = [_extract_entry_row(e) for e in entries]
            df = pd.DataFrame(rows)

            cols = ["citation_raw", "record_key"] + [c for c in df.columns if c not in ("citation_raw", "record_key")]
            df = df[cols]

            header_needed = (not os.path.exists(out_csv)) or (os.path.getsize(out_csv) == 0)
            df.to_csv(out_csv, mode="a", header=header_needed, index=False)

            got = len(df)
            retrieved += got
            if pbar:
                pbar.update(got)

            if total_reported is not None and retrieved >= total_reported:
                break

            if sleep_s > 0:
                time.sleep(sleep_s)

    if pbar:
        pbar.close()

    meta["ended_utc"] = utc_now_iso()
    meta["retrieved_rows"] = int(retrieved)
    meta["total_reported_by_scopus"] = int(total_reported or 0)
    return int(total_reported or 0), int(retrieved), meta


def _bib_escape(s: str) -> str:
    if s is None:
        return ""
    s = str(s).replace("\n", " ").replace("\r", " ").strip()
    s = s.replace("\\", "\\\\")
    s = s.replace("{", "\\{").replace("}", "\\}")
    return s


def _safe_bib_key(doi_core: str, scopus_id: str, title: str, year: str) -> str:
    if doi_core:
        key = re.sub(r"[^a-z0-9]+", "", doi_core.lower())
        return key[:60] if key else "rec"
    if scopus_id:
        key = re.sub(r"[^a-z0-9]+", "", str(scopus_id).lower())
        return key[:60] if key else "rec"
    base = f"{year}_{clean_text(title or '')}"
    base = re.sub(r"[^a-z0-9_]+", "", base.lower())
    return (base[:60] if base else "rec")


def export_scopus_to_bib(df: pd.DataFrame, out_path: str) -> int:
    n = 0
    with open(out_path, "w", encoding="utf-8") as f:
        for _, r in df.iterrows():
            title = (r.get("title") or "").strip()
            journal = (r.get("publicationName") or "").strip()
            cover = r.get("coverDate")
            year = _year_from_coverdate(cover)
            doi_core = normalize_doi(r.get("doi"))
            url = f"https://doi.org/{doi_core}" if doi_core else ""
            scopus_id = r.get("scopus_id")

            entry_type = "article" if journal else "misc"
            key = _safe_bib_key(doi_core, scopus_id, title, year)

            fields = []
            if title:
                fields.append(f"  title = {{{_bib_escape(title)}}}")
            if journal:
                fields.append(f"  journal = {{{_bib_escape(journal)}}}")
            if year:
                fields.append(f"  year = {{{_bib_escape(year)}}}")
            if doi_core:
                fields.append(f"  doi = {{{_bib_escape(doi_core)}}}")
            if url:
                fields.append(f"  url = {{{_bib_escape(url)}}}")

            if not fields:
                continue

            f.write(f"@{entry_type}{{{key},\n")
            f.write(",\n".join(fields))
            f.write("\n}\n\n")
            n += 1
    return n


def export_scopus_to_ris(df: pd.DataFrame, out_path: str) -> int:
    n = 0
    with open(out_path, "w", encoding="utf-8") as f:
        for _, r in df.iterrows():
            title = (r.get("title") or "").strip()
            journal = (r.get("publicationName") or "").strip()
            cover = r.get("coverDate")
            year = _year_from_coverdate(cover)
            doi_core = normalize_doi(r.get("doi"))
            url = f"https://doi.org/{doi_core}" if doi_core else ""

            ty = "JOUR" if journal else "GEN"
            f.write(f"TY  - {ty}\n")
            if title:
                f.write(f"TI  - {title}\n")
            if journal:
                f.write(f"JO  - {journal}\n")
            if year:
                f.write(f"PY  - {year}\n")
            if doi_core:
                f.write(f"DO  - {doi_core}\n")
            if url:
                f.write(f"UR  - {url}\n")
            f.write("ER  - \n\n")
            n += 1
    return n


def export_scopus_to_word(df: pd.DataFrame, title: str, out_path: str, max_items: Optional[int]) -> int:
    doc = Document()
    doc.add_heading(title, level=1)

    n = 0
    for _, r in df.iterrows():
        if max_items is not None and n >= max_items:
            break
        c = (r.get("citation_raw") or "").strip()
        if not c:
            t = (r.get("title") or "").strip()
            j = (r.get("publicationName") or "").strip()
            y = _year_from_coverdate(r.get("coverDate"))
            doi_core = normalize_doi(r.get("doi"))
            c = _build_scopus_citation_raw("", y, t, j, doi_core)

        if c:
            p = doc.add_paragraph(c, style="List Number")
            p.paragraph_format.space_after = Pt(6)
            n += 1

    doc.save(out_path)
    return n


def step2_retrieve_total(config: dict) -> dict:
    load_dotenv()

    api_key = config.get("scopus_api_key") or os.getenv("SCOPUS_API_KEY")
    inst_token = config.get("scopus_inst_token") or os.getenv("SCOPUS_INST_TOKEN")

    if not api_key:
        raise SystemExit("SCOPUS_API_KEY missing (set it in .env at repo root).")

    auth = ScopusAuth(api_key=api_key, inst_token=inst_token)

    scopus_search_url = config.get("scopus_search_url") or getattr(cfg, "scopus_search_url", None)
    if not scopus_search_url:
        raise SystemExit("SCOPUS_SEARCH_URL missing in config.py (or config dict).")

    search_strings_yml = config.get("search_strings_yml") or getattr(cfg, "SEARCH_STRINGS_YML", None)
    if not search_strings_yml:
        raise SystemExit("SEARCH_STRINGS_YML missing in config.py (or config dict).")

    out_dir = config.get("out_dir") or getattr(cfg, "out_dir", None) or getattr(cfg, "OUT_DIR", None)
    if not out_dir:
        raise SystemExit("out_dir/OUT_DIR missing in config.py (or config dict).")

    step2_dir = os.path.join(out_dir, "step2")
    step2_total_csv = os.path.join(step2_dir, "step2_total_records.csv")
    step2_total_meta_json = os.path.join(step2_dir, "step2_total_records.meta.json")

    use_post = int(config.get("use_post_for_search", getattr(cfg, "USE_POST_FOR_SEARCH", DEFAULT_USE_POST_FOR_SEARCH)))
    view = str(config.get("view", getattr(cfg, "VIEW", DEFAULT_VIEW)))

    count_per_page = int(
        config.get("count_per_page_retrieve", getattr(cfg, "COUNT_PER_PAGE_RETRIEVE", DEFAULT_COUNT_PER_PAGE_RETRIEVE))
    )
    sleep_s = float(config.get("step2_sleep_s", getattr(cfg, "STEP2_SLEEP_S", DEFAULT_STEP2_SLEEP_S)))
    deep_paging_limit = int(config.get("deep_paging_limit", getattr(cfg, "DEEP_PAGING_LIMIT", DEFAULT_DEEP_PAGING_LIMIT)))
    pubyear_min = int(config.get("pubyear_min", getattr(cfg, "PUBYEAR_MIN", DEFAULT_PUBYEAR_MIN)))
    pubyear_max = int(config.get("pubyear_max", getattr(cfg, "PUBYEAR_MAX", DEFAULT_PUBYEAR_MAX)))
    max_results_total = config.get("max_results_total", getattr(cfg, "MAX_RESULTS_TOTAL", DEFAULT_MAX_RESULTS_TOTAL))

    export_word = bool(config.get("export_word", True))
    export_bib = bool(config.get("export_bib", True))
    export_ris = bool(config.get("export_ris", True))
    mw = config.get("max_word_items", getattr(cfg, "MAX_WORD_ITEMS", None))
    max_word_items = int(mw) if mw not in (None, "") else None

    os.makedirs(step2_dir, exist_ok=True)

    already_have_csv = os.path.exists(step2_total_csv) and os.path.getsize(step2_total_csv) > 0

    base_total = 0
    planned_total = 0
    combined_rows = 0
    slice_metas: List[dict] = []

    if already_have_csv:
        print(f"[step2] SKIP retrieval (exists): {step2_total_csv}")
        try:
            df_existing = pd.read_csv(step2_total_csv, engine="python", on_bad_lines="skip")
            combined_rows = int(len(df_existing))
        except Exception:
            combined_rows = 0
    else:
        search_cfg = read_yaml(search_strings_yml)
        _, _, base_q = _build_queries(search_cfg)
        if not base_q:
            raise SystemExit("[step2] TOTAL__ALL query is empty; check search_strings.yml")

        base_total, _ = scopus_count_only(auth, scopus_search_url, base_q, use_post=use_post, view=view, count_per_page=1)
        print(f"[step2] TOTAL__ALL reported by Scopus: {base_total:,}")

        print("[step2] Planning PUBYEAR slices...")
        year_slices = _plan_year_slices(
            auth,
            scopus_search_url,
            base_q,
            pubyear_min,
            pubyear_max,
            deep_paging_limit=deep_paging_limit,
            use_post=use_post,
            view=view,
        )

        # Expand any single-year slices that still exceed deep paging limit
        expanded_slices = _expand_slices(
            auth,
            scopus_search_url,
            base_q,
            year_slices,
            deep_paging_limit=deep_paging_limit,
            use_post=use_post,
            view=view,
        )

        missing_q = _with_pubyear_missing_or_outside(base_q, pubyear_min, pubyear_max)
        missing_total, _ = scopus_count_only(auth, scopus_search_url, missing_q, use_post=use_post, view=view, count_per_page=1)

        planned_total = sum(int(s["expected"]) for s in expanded_slices) + int(missing_total)
        print("[step2] Slice plan:")
        for s in expanded_slices:
            print(f"  - {s['tag']}: {int(s['expected']):,}")
        print(f"  - PUBYEAR_MISSING_OR_OUTSIDE: {missing_total:,}")
        print(f"[step2] Planned total across slices (may exceed base_total due to overlap): {planned_total:,}")
        print(f"[step2] NOTE: overlap can occur in SUBJAREA slices; final output is deduped by record_key.")

        slice_csvs: List[str] = []
        slice_metas = []

        # Retrieve expanded slices
        for s in expanded_slices:
            tag = s["tag"]
            q_slice = s["query"]
            slice_total = int(s["expected"])

            out_csv = os.path.join(step2_dir, f"step2_total_records__{tag}.csv")
            out_meta = os.path.join(step2_dir, f"step2_total_records__{tag}.meta.json")

            existing_rows = 0
            if os.path.exists(out_csv) and os.path.getsize(out_csv) > 0:
                _sanitize_csv_tail(out_csv)
                existing_rows = _count_csv_data_rows(out_csv)

            expected_complete = False
            if os.path.exists(out_meta) and os.path.getsize(out_meta) > 0:
                try:
                    with open(out_meta, "r", encoding="utf-8") as f:
                        prior = json.load(f)
                    prior_total = int(prior.get("total_reported_by_scopus") or 0)
                    if prior_total > 0 and existing_rows >= prior_total:
                        expected_complete = True
                except Exception:
                    pass

            if expected_complete:
                print(f"[step2] SKIP {tag} (complete): {out_csv} ({existing_rows:,} rows)")
                slice_csvs.append(out_csv)
                continue

            if existing_rows > 0:
                print(f"[step2] RESUME {tag}: {out_csv} ({existing_rows:,} rows so far)")

            if slice_total > deep_paging_limit:
                print(f"[step2] WARNING: {tag} still > deep_paging_limit ({slice_total:,} > {deep_paging_limit:,}). "
                      f"This will likely 400 without further slicing.")
            print(f"\n[step2] RETRIEVE {tag} (expected {slice_total:,})")

            total_reported, retrieved, meta = scopus_retrieve_stream_to_csv_start(
                auth,
                scopus_search_url,
                query=q_slice,
                out_csv=out_csv,
                count_per_page=count_per_page,
                sleep_s=sleep_s,
                max_results=max_results_total,
                use_post=use_post,
                view=view,
            )

            meta.update(
                {
                    "slice_tag": tag,
                    "expected_total_from_planner": int(slice_total),
                    "total_reported_by_scopus": int(total_reported),
                    "retrieved_rows": int(retrieved),
                    "timestamp_utc": utc_now_iso(),
                    "slice_kind": s.get("kind"),
                }
            )
            _save_json(out_meta, meta)

            print(f"[step2] Done {tag}: reported={total_reported:,} retrieved={retrieved:,}")
            slice_csvs.append(out_csv)
            slice_metas.append(meta)

        # Missing/outside slice (still year-bounded exclusion)
        if missing_total > 0:
            tag = "PUBYEAR_MISSING_OR_OUTSIDE"
            out_csv = os.path.join(step2_dir, f"step2_total_records__{tag}.csv")
            out_meta = os.path.join(step2_dir, f"step2_total_records__{tag}.meta.json")

            existing_rows = 0
            if os.path.exists(out_csv) and os.path.getsize(out_csv) > 0:
                _sanitize_csv_tail(out_csv)
                existing_rows = _count_csv_data_rows(out_csv)

            expected_complete = False
            if os.path.exists(out_meta) and os.path.getsize(out_meta) > 0:
                try:
                    with open(out_meta, "r", encoding="utf-8") as f:
                        prior = json.load(f)
                    prior_total = int(prior.get("total_reported_by_scopus") or 0)
                    if prior_total > 0 and existing_rows >= prior_total:
                        expected_complete = True
                except Exception:
                    pass

            if expected_complete:
                print(f"[step2] SKIP {tag} (complete): {out_csv} ({existing_rows:,} rows)")
                slice_csvs.append(out_csv)
            else:
                if existing_rows > 0:
                    print(f"[step2] RESUME {tag}: {out_csv} ({existing_rows:,} rows so far)")

                print(f"\n[step2] RETRIEVE {tag} (expected {missing_total:,})")
                total_reported, retrieved, meta = scopus_retrieve_stream_to_csv_start(
                    auth,
                    scopus_search_url,
                    query=missing_q,
                    out_csv=out_csv,
                    count_per_page=count_per_page,
                    sleep_s=sleep_s,
                    max_results=max_results_total,
                    use_post=use_post,
                    view=view,
                )
                meta.update(
                    {
                        "slice_tag": tag,
                        "expected_total_from_planner": int(missing_total),
                        "total_reported_by_scopus": int(total_reported),
                        "retrieved_rows": int(retrieved),
                        "timestamp_utc": utc_now_iso(),
                        "slice_kind": "PUBYEAR_MISSING_OR_OUTSIDE",
                    }
                )
                _save_json(out_meta, meta)
                print(f"[step2] Done {tag}: reported={total_reported:,} retrieved={retrieved:,}")
                slice_csvs.append(out_csv)
                slice_metas.append(meta)

        print(f"\n[step2] Combining slice CSVs -> {step2_total_csv}")
        combined_rows = _concat_csvs(slice_csvs, step2_total_csv)

        print(f"[step2] Deduping combined CSV by record_key -> {step2_total_csv}")
        deduped_rows = _dedupe_csv_in_place(step2_total_csv, key_col="record_key")
        print(f"[step2] Combined rows before dedupe: {combined_rows:,} | after dedupe: {deduped_rows:,}")
        combined_rows = deduped_rows

    meta_out = {
        "base_total_reported_by_scopus": int(base_total),
        "planned_total_across_slices": int(planned_total),
        "combined_rows_written_after_dedupe": int(combined_rows),
        "pubyear_min": int(pubyear_min),
        "pubyear_max": int(pubyear_max),
        "deep_paging_limit": int(deep_paging_limit),
        "count_per_page": int(count_per_page),
        "sleep_s": float(sleep_s),
        "max_results_total": max_results_total,
        "timestamp_utc": utc_now_iso(),
        "slice_metas": slice_metas,
    }

    exports = {}
    df_all = None
    try:
        df_all = pd.read_csv(step2_total_csv, engine="python", on_bad_lines="skip")
    except Exception:
        df_all = None

    if isinstance(df_all, pd.DataFrame) and not df_all.empty:
        if export_bib:
            bib_path = os.path.join(step2_dir, "step2_total_records.bib")
            exports["bib_entries"] = int(export_scopus_to_bib(df_all, bib_path))
            exports["bib_path"] = bib_path

        if export_ris:
            ris_path = os.path.join(step2_dir, "step2_total_records.ris")
            exports["ris_entries"] = int(export_scopus_to_ris(df_all, ris_path))
            exports["ris_path"] = ris_path

        if export_word:
            docx_path = os.path.join(step2_dir, "step2_total_records.docx")
            exports["docx_items"] = int(export_scopus_to_word(df_all, "Scopus Retrieved Records", docx_path, max_word_items))
            exports["docx_path"] = docx_path

    meta_out["exports"] = exports
    _save_json(step2_total_meta_json, meta_out)

    print(f"[step2] Final combined rows (after dedupe): {combined_rows:,}")

    return {
        "status": "ok",
        "step2_total_csv": step2_total_csv,
        "step2_total_meta_json": step2_total_meta_json,
        "base_total_reported_by_scopus": int(base_total),
        "combined_rows_written_after_dedupe": int(combined_rows),
        "exports": exports,
    }


def run_step1b_retrieve_total(config: dict) -> dict:
    return step2_retrieve_total(config)


if __name__ == "__main__":
    step2_retrieve_total({})


def run(config: dict) -> dict:
    return step2_retrieve_total(config)


def run_step2(config: dict) -> dict:
    return step2_retrieve_total(config)


def main(config: dict | None = None) -> dict:
    return step2_retrieve_total(config or {})
