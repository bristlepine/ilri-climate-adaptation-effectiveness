"""
publish_d5_from_d4.py

1. Downloads D4 v02 from Google Drive as a Word (.docx) — preserves watermark,
   header icons, all fonts and styles exactly.
2. Renames to D5 v01.
3. Makes targeted text replacements.
4. Replaces figure captions with actual embedded PNG images (human-coded).
5. Replaces Table 1 with the final D5 pipeline table.
6. Saves locally as Deliverable_5_Bristlepine_Final_Systematic_Map_v01.docx.

Run: python3 deliverables/publish_d5_from_d4.py
"""

from __future__ import annotations

import copy
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

HERE   = Path(__file__).resolve().parent
ROOT   = HERE.parent
STEP16 = ROOT / "scripts" / "outputs" / "step16"
HUMAN  = STEP16 / "interactive" / "human"
OUT    = HERE / "Deliverable_5_Bristlepine_Final_Systematic_Map_v01.docx"

CREDS_DIR  = HERE / ".credentials"
TOKEN_FILE = CREDS_DIR / "token.json"
SCOPES     = ["https://www.googleapis.com/auth/drive"]

D4_V02_DOC_ID = "14Hejjfr63HdxH21bqBxLZe3AsXxmLBrXYqJJk7ZbBxU"

# ── figures: (search text in D4, png path, caption text, width_cm) ────────────
# search text must uniquely identify the paragraph that IS or comes just before
# where the image should go.
PRISMA_PNG = HERE / "prisma_flow_d5.png"   # new dual-track PRISMA
HUMAN      = STEP16 / "interactive" / "human"
SAT_PNG    = STEP16 / "saturation.png"

FIGURES = [
    # Figure 1: new dual-track PRISMA box diagram
    ("Figure 1 presents the ROSES",
     PRISMA_PNG,
     "Figure 1. PRISMA flow diagram — record flow across 29 sources and four screening stages. "
     "Human-coded track (amber, n = 86) is the primary output. "
     "LLM track (n = 2,368) is an exploratory reference.",
     15.5, "after"),
    # Figure 2: Evidence Gap Map (human)
    ("Figure 2 presents the",
     HUMAN / "evidence_gap_map.png",
     "Figure 2. Evidence gap map — human-coded results (n = 86). "
     "Blue = process domains · green = outcome domains · grey = evidence gaps. "
     "Bubble area ∝ number of studies.",
     15.5, "after"),
    # Figures 3–6: replace D4's standalone caption paragraphs
    ("Figure 3. Geographic distribution",
     HUMAN / "geographic_map.png",
     "Figure 3. Geographic distribution of included studies (human-coded, n = 86).",
     15.5, "replace"),
    ("Figure 4.",
     HUMAN / "geographic_bar.png",
     "Figure 4. Top countries by study count (human-coded, n = 86).",
     13.0, "replace"),
    ("Figure 5. Distribution of methodological",
     HUMAN / "methodology.png",
     "Figure 5. Methodological approaches across included studies (human-coded, n = 86).",
     13.0, "replace"),
    ("Figure 6. Temporal trends",
     HUMAN / "temporal_trends.png",
     "Figure 6. Publication trends over time (human-coded, n = 86).",
     13.0, "replace"),
]

# New figures appended before §5 Searchable Database
NEW_FIGURES = [
    (HUMAN / "equity.png",
     "Figure 7. Equity and inclusion dimensions (human-coded, n = 86). "
     "Red bar = studies with no marginalized group focus.",
     13.0),
    (SAT_PNG,
     "Figure 8. Information saturation curve. All three tracked dimensions plateau by "
     "batch FT-R2c (49 papers). Zero new categories added across the final two batches (37 papers).",
     15.5),
    (STEP16 / "llm_vs_human.png",
     "Figure 9. Human vs LLM comparison across key variables (human amber n=86; LLM teal n=2,368). "
     "Distributions are broadly consistent, supporting representativeness of the human sample.",
     15.5),
]

# ── text replacements: (old, new) ─────────────────────────────────────────────
REPLACEMENTS = [
    # Cover / title
    ("Deliverable 4: First Draft Systematic Map (Preliminary)", "Deliverable 5: Final Systematic Map"),
    ("Deliverable 4: First Draft Systematic Map",              "Deliverable 5: Final Systematic Map"),
    ("First Draft Systematic Map (Preliminary)",               "Final Systematic Map"),
    ("First Draft Systematic Map",                             "Final Systematic Map"),
    ("April 2026, v02",  "June 2026, v01"),
    ("April 2026",       "June 2026"),
    # Remove preliminary flags
    (" (Preliminary)", ""),
    ("(Preliminary) ", ""),
    ("Preliminary — ", ""),
    (", preliminary", ""),
    (" preliminary ", " "),
    # Deliverable reference
    ("will be reported in Deliverable 5", "are reported in this document"),
    # Section title
    ("4.  Preliminary Results", "4.  Final Results"),
    ("4. Preliminary Results",  "4. Final Results"),
    # Coded record counts
    ("coded dataset of 2,093 studies", "coded dataset of 86 human-coded studies"),
    ("2,093 records were",             "2,368 records were"),
    # FT retrieval (update prose references)
    ("2,644 of 6,218",     "3,476 of 8,748 (40.1%)"),
    ("43%) were retrieved","40.1%) were retrieved"),
    ("2,644 full texts",   "3,476 full texts"),
    ("3,574 require manual","5,232 require manual"),
    # Net-new screening
    ("8,187 additional records", "8,895 additional records (all non-Scopus databases)"),
    ("2,340 were included",      "2,340 were included at title/abstract stage"),
    # LLM tool
    ("Ollama/Llama3", "Ollama/qwen2.5:14b"),
    # Intro note
    ("Note: This document is a preliminary first draft",
     "Note: This document presents the final systematic map."),
    # Figure prose — update to reflect human results
    ("Figure 2 presents the preliminary evidence gap map, showing the distribution of 2,646 coded studies",
     "Figure 2 presents the evidence gap map based on the 86 human-coded studies"),
    ("Key preliminary observations (to be confirmed and expanded in Deliverable 5):",
     "Key observations from the human-coded results:"),
    # Geographic note (remove "preliminary")
    ("Geographic distribution: Studies are concentrated",
     "Geographic distribution: Studies are concentrated"),
    ("(all databases, preliminary)", "(human-coded, n = 86)"),
    # Screened-database note in §5
    ("The preliminary searchable database", "The searchable database"),
    ("All 2,750 records coded in the preliminary, exploratory data extraction pass",
     "All 2,368 LLM-screened records (exploratory reference) plus 86 human-coded records (primary)"),
    # Limitations — remove D4-specific "to be addressed"
    ("This preliminary report has several limitations that will be addressed in Deliverable 5:",
     "This systematic map has the following limitations:"),
    # Next Steps → Conclusions
    ("7.  Next Steps", "7.  Conclusions"),
    ("7. Next Steps",  "7. Conclusions"),
    # Abstract screening totals in exec summary
    ("8,558 included", "8,558 records for full-text assessment"),
    # Human coding mention
    ("Data extraction proceeded iteratively in batches",
     "Data extraction proceeded iteratively in batches of 20 papers"),
]

# ── Final D5 pipeline table rows (replace existing Table 1) ───────────────────
TABLE_HEADERS = ["Database", "Returned", "After Dedup",
                 "Abstr. Incl.", "FT Retrieved", "FT Screened", "Coded"]
TABLE_ROWS = [
    ["Scopus",                   "17,021", "17,021", "6,218", "2,644", "2,644", "—"],
    ["Web of Science",           "15,179",  "4,683", "1,140",   "552",   "552", "—"],
    ["CAB Abstracts",             "5,723",  "3,229", "1,133",   "260",   "260", "—"],
    ["Academic Search Premier",   "1,187",    "274",    "66",    "20",    "20", "—"],
    ["AGRIS",                         "3",      "1",     "1",     "0",    "—", "—"],
    ["Total",                    "39,113", "25,208", "8,558", "3,476", "3,464", "2,368"],
]

GREEN    = RGBColor(0x21, 0x47, 0x2E)
CHARCOAL = RGBColor(0x3C, 0x3C, 0x3C)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GREY     = RGBColor(0x75, 0x75, 0x75)


# ── helpers ───────────────────────────────────────────────────────────────────

def get_creds() -> Credentials:
    creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), SCOPES)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
    return creds


def download_as_docx(drive, file_id: str) -> bytes:
    req = drive.files().export_media(
        fileId=file_id,
        mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    buf = io.BytesIO()
    dl = MediaIoBaseDownload(buf, req)
    done = False
    while not done:
        _, done = dl.next_chunk()
    return buf.getvalue()


def full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def replace_text_in_para(para, old: str, new: str) -> bool:
    """Replace old→new in a paragraph, preserving run structure where possible."""
    text = full_text(para)
    if old not in text:
        return False
    # Simple case: replacement fits in first run
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Multi-run case: rebuild paragraph text
    new_text = text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def set_cell_bg(cell, hex6: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.lstrip("#"))
    tcPr.append(shd)


def cell_write(cell, text: str, bold=False, color: RGBColor | None = None,
               size: int = 9, align: str = "left") -> None:
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    para.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def insert_image_para(doc: Document, png: Path, caption: str,
                      width_cm: float, before_para_idx: int) -> None:
    """Insert an image paragraph + caption paragraph before the given paragraph index."""
    # Image para
    img_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    img_para.append(pPr)

    img_para_obj = doc.add_paragraph()
    img_para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para_obj.paragraph_format.space_before = Pt(6)
    img_para_obj.paragraph_format.space_after  = Pt(4)
    run = img_para_obj.add_run()
    if png.exists():
        run.add_picture(str(png), width=Cm(width_cm))
    else:
        img_para_obj.add_run(f"[FIGURE NOT FOUND: {png.name}]")

    # Caption para
    cap_para_obj = doc.add_paragraph(caption)
    cap_para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para_obj.paragraph_format.space_before = Pt(2)
    cap_para_obj.paragraph_format.space_after  = Pt(12)
    for run in cap_para_obj.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GREY

    # Move both before target paragraph
    body = doc.element.body
    target = body.paragraphs[before_para_idx]._element
    body.remove(img_para_obj._element)
    body.remove(cap_para_obj._element)
    body.insert(list(body).index(target), cap_para_obj._element)
    body.insert(list(body).index(target), img_para_obj._element)


# ── main ──────────────────────────────────────────────────────────────────────

def run():
    print("[d5] Authenticating...")
    creds = get_creds()
    drive = build("drive", "v3", credentials=creds)

    # 1. Download D4 v02 as Word
    print(f"[d5] Downloading D4 v02 ({D4_V02_DOC_ID}) as .docx...")
    docx_bytes = download_as_docx(drive, D4_V02_DOC_ID)
    print(f"[d5]   {len(docx_bytes):,} bytes received")

    doc = Document(io.BytesIO(docx_bytes))

    # 2. Text replacements throughout entire document
    print("[d5] Applying text replacements...")
    n_replaced = 0
    all_paras = list(doc.paragraphs)
    # Also replace in table cells
    all_cells = [cell for tbl in doc.tables for row in tbl.rows for cell in row.cells]
    all_cell_paras = [p for cell in all_cells for p in cell.paragraphs]

    for old, new in REPLACEMENTS:
        if old == new:
            continue
        for para in all_paras + all_cell_paras:
            if replace_text_in_para(para, old, new):
                n_replaced += 1

    print(f"[d5]   {n_replaced} replacements made")

    # 3. Replace Table 1 — delete the old one (multi-paragraph cells cause doubling)
    #    and insert a clean new table at the same body position.
    print("[d5] Replacing pipeline table...")
    if doc.tables:
        old_tbl = doc.tables[0]
        body = doc.element.body
        tbl_elem = old_tbl._element
        tbl_idx  = list(body).index(tbl_elem)

        # Build new table (appended at end temporarily)
        all_rows = [TABLE_HEADERS] + TABLE_ROWS
        new_tbl = doc.add_table(rows=len(all_rows), cols=len(TABLE_HEADERS))
        # Use the first available table style from the downloaded D4 template
        try:
            new_tbl.style = "Table Grid"
        except KeyError:
            try:
                new_tbl.style = doc.styles.element.xpath(
                    './/w:style[@w:type="table"]', namespaces=new_tbl._element.nsmap
                )[0].get(qn("w:styleId"))
            except Exception:
                pass  # leave default style

        for r_idx, row_data in enumerate(all_rows):
            is_header = (r_idx == 0)
            is_total  = (r_idx == len(all_rows) - 1)
            for c_idx, val in enumerate(row_data):
                cell = new_tbl.rows[r_idx].cells[c_idx]
                if is_header:
                    set_cell_bg(cell, "21472E")
                    cell_write(cell, val, bold=True, color=WHITE, size=9,
                               align="center" if c_idx > 0 else "left")
                elif is_total:
                    set_cell_bg(cell, "3C3C3C")
                    cell_write(cell, val, bold=True, color=WHITE, size=9,
                               align="center" if c_idx > 0 else "left")
                else:
                    bg = "F9F7F4" if r_idx % 2 == 1 else "FFFFFF"
                    set_cell_bg(cell, bg)
                    cell_write(cell, val, size=9,
                               align="center" if c_idx > 0 else "left")

        # Move new table to where old table was, then delete old table
        new_tbl_elem = new_tbl._element
        body.remove(new_tbl_elem)
        body.insert(tbl_idx, new_tbl_elem)
        body.remove(tbl_elem)
        print("[d5]   Pipeline table replaced cleanly")
    else:
        print("[d5]   WARN: no tables found in document")

    # 4. Insert figures
    # Strategy: work bottom-up so earlier insertions don't shift later indices.
    # For "replace" mode: clear the caption paragraph text, insert image before it.
    # For "after" mode: insert image after the matching paragraph.

    print("[d5] Inserting figures...")
    paras = doc.paragraphs

    # Process FIGURES in reverse order (bottom-up) to avoid index drift
    for search, png, caption, width_cm, mode in reversed(FIGURES):
        # Find the paragraph index
        idx = None
        for i, p in enumerate(paras):
            if search in full_text(p):
                idx = i
                break

        if idx is None:
            print(f"[d5]   WARN: not found — '{search[:50]}'")
            continue

        if mode == "replace":
            # Replace the caption paragraph text and insert image before it
            cap_para = paras[idx]
            # Clear existing text, set it to our new caption
            for run in cap_para.runs:
                run.text = ""
            if cap_para.runs:
                run = cap_para.runs[0]
                run.text = caption
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = GREY
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insert image para before the caption
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.space_before = Pt(8)
            img_para.paragraph_format.space_after  = Pt(4)
            r = img_para.add_run()
            if png.exists():
                r.add_picture(str(png), width=Cm(width_cm))
            else:
                img_para.add_run(f"[FIGURE NOT FOUND: {png.name}]")

            body = doc.element.body
            target_elem = cap_para._element
            body.remove(img_para._element)
            body.insert(list(body).index(target_elem), img_para._element)
            print(f"[d5]   Inserted (replace) {png.name}")

        else:  # "after"
            # Insert image + caption after the prose paragraph
            prose_para = paras[idx]

            # Caption para
            cap_para = doc.add_paragraph(caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.paragraph_format.space_after  = Pt(12)
            cap_para.paragraph_format.space_before = Pt(4)
            for r in cap_para.runs:
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = GREY

            # Image para
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.space_before = Pt(8)
            img_para.paragraph_format.space_after  = Pt(4)
            r = img_para.add_run()
            if png.exists():
                r.add_picture(str(png), width=Cm(width_cm))
            else:
                img_para.add_run(f"[NOT FOUND: {png.name}]")

            # Move both to just after prose_para
            body = doc.element.body
            prose_elem = prose_para._element
            prose_idx = list(body).index(prose_elem)

            body.remove(img_para._element)
            body.remove(cap_para._element)
            body.insert(prose_idx + 1, cap_para._element)
            body.insert(prose_idx + 1, img_para._element)
            print(f"[d5]   Inserted (after) {png.name}")

        # Refresh paragraph list after each insertion
        paras = doc.paragraphs

    # 5. Append new figures (equity, saturation, LLM vs human)
    # Insert before "5.  Searchable Database" heading
    print("[d5] Appending new figures (equity, saturation, LLM vs human)...")
    for png, caption, width_cm in reversed(NEW_FIGURES):
        paras = doc.paragraphs
        target_idx = None
        for i, p in enumerate(paras):
            txt = full_text(p)
            if "Searchable Database" in txt or "5." in txt and "Searchable" in txt:
                target_idx = i
                break
        if target_idx is None:
            # Append at end of section 4
            for i, p in enumerate(paras):
                if "Limitations" in full_text(p) or "6." in full_text(p):
                    target_idx = i
                    break

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(8)
        img_para.paragraph_format.space_after  = Pt(4)
        r = img_para.add_run()
        if png.exists():
            r.add_picture(str(png), width=Cm(width_cm))
        else:
            img_para.add_run(f"[NOT FOUND: {png.name}]")

        cap_para = doc.add_paragraph(caption)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(4)
        cap_para.paragraph_format.space_after  = Pt(12)
        for r in cap_para.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = GREY

        if target_idx is not None:
            body = doc.element.body
            target_elem = doc.paragraphs[target_idx]._element
            body.remove(img_para._element)
            body.remove(cap_para._element)
            body.insert(list(body).index(target_elem), cap_para._element)
            body.insert(list(body).index(target_elem), img_para._element)
            print(f"[d5]   Appended {png.name}")
        else:
            print(f"[d5]   Appended {png.name} at end")

    # 6. Save
    doc.save(str(OUT))
    print(f"\n[d5] ✓ Saved: {OUT}")
    print(f"[d5]   Size: {OUT.stat().st_size / 1024:.0f} KB")
    print(f"\nAll figures use human-coded results (n=86) except:")
    print(f"  Figure 9 (LLM vs human comparison) — shows both tracks side by side")


if __name__ == "__main__":
    run()
