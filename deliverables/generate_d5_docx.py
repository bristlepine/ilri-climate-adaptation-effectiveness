"""
generate_d5_docx.py

Generates Deliverable 5 as a full Word document (.docx) with:
  - Cover page
  - Key findings (6 headline stats from systematic-map frontend)
  - Methods overview (Steps 1–7, matching methodology page)
  - ROSES flow diagram
  - Evidence Gap Map — Human (primary output)
  - LLM vs Human comparison figure
  - Domain heatmap
  - Geographic distribution
  - Temporal trends
  - Methodological approaches
  - Equity & inclusion
  - Saturation / convergence
  - Pipeline table (all databases)
  - Data availability note

Run: conda run -n ilri01 python deliverables/generate_d5_docx.py
"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

HERE     = Path(__file__).resolve().parent
ROOT     = HERE.parent
STEP16   = ROOT / "scripts" / "outputs" / "step16"
STEP15C  = ROOT / "scripts" / "outputs" / "step15c"
HUMAN    = STEP16 / "interactive" / "human"
OUT_DOCX = HERE / "Deliverable_5_Bristlepine_Final_Systematic_Map_v01.docx"

# Brand colours
GREEN    = RGBColor(0x21, 0x47, 0x2E)
CHARCOAL = RGBColor(0x3C, 0x3C, 0x3C)
SAND     = RGBColor(0xF5, 0xF2, 0xEC)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
AMBER    = RGBColor(0xC2, 0x70, 0x3A)

# ── PNG helpers ────────────────────────────────────────────────────────────────

def _json_to_png(json_path: Path, out_path: Path, width: int = 1100, height: int = 700) -> bool:
    """Render a Plotly JSON file to PNG using kaleido. Returns True on success."""
    try:
        import plotly.graph_objects as go
        with open(json_path, encoding="utf-8") as f:
            fig_data = json.load(f)
        fig = go.Figure(fig_data)
        fig.write_image(str(out_path), width=width, height=height, scale=2)
        print(f"[png] {out_path.name}")
        return True
    except Exception as e:
        print(f"[png] WARN: {out_path.name} — {e}")
        return False


def ensure_pngs() -> dict[str, Path | None]:
    """Make sure all needed PNGs exist; generate from JSON where missing."""
    pngs: dict[str, Path | None] = {}

    def _reg(key: str, path: Path) -> None:
        pngs[key] = path if path.exists() else None

    # Already-rendered LLM PNGs
    _reg("roses",      STEP16 / "roses_flow.png")
    _reg("egm_llm",   STEP16 / "evidence_gap_map.png")
    _reg("heatmap",   STEP16 / "domain_heatmap.png")
    _reg("geo_map",   STEP16 / "geographic_map.png")
    _reg("geo_bar",   STEP16 / "geographic_bar.png")
    _reg("temporal",  STEP16 / "temporal_trends.png")
    _reg("method",    STEP16 / "methodology_bar.png")
    _reg("equity",    STEP16 / "equity_bar.png")
    _reg("llm_human", STEP16 / "llm_vs_human.png")

    # Human EGM — generate from JSON if PNG missing
    human_egm_png = HUMAN / "evidence_gap_map.png"
    if not human_egm_png.exists():
        human_egm_json = HUMAN / "evidence_gap_map.json"
        if human_egm_json.exists():
            _json_to_png(human_egm_json, human_egm_png, width=1100, height=700)
    _reg("egm_human", human_egm_png)

    # Saturation — generate from JSON if PNG missing
    sat_png = STEP16 / "saturation.png"
    if not sat_png.exists():
        sat_json = STEP15C / "saturation.json"
        if sat_json.exists():
            _json_to_png(sat_json, sat_png, width=1100, height=600)
    _reg("saturation", sat_png)

    # Domain type bar (human)
    human_domtype_png = HUMAN / "domain_type_bar.png"
    if not human_domtype_png.exists():
        human_domtype_json = HUMAN / "domain_type.json"
        if human_domtype_json.exists():
            _json_to_png(human_domtype_json, human_domtype_png, width=900, height=460)
    _reg("domtype_human", human_domtype_png)

    return pngs


# ── docx helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None,
               size: int = 9, align: str = "left") -> None:
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = GREEN
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = GREEN
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def _add_label(doc: Document, text: str) -> None:
    """Small uppercase section label above a heading."""
    p = doc.add_paragraph(text.upper())
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = GREEN
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def _add_figure(doc: Document, png: Path | None, caption: str, width_cm: float = 15.0) -> None:
    if png and png.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png), width=Cm(width_cm))
        # Caption
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        cap.runs[0].italic = True
        cap.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph(f"[Figure not available: {caption}]")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
        p.runs[0].italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)


def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DDDDDD")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── pipeline table ──────────────────────────────────────────────────────────────

PIPELINE_HEADERS = ["Database", "Returned", "After Dedup", "Abstr. Incl.", "FT Retrieved", "FT Screened", "Coded"]
PIPELINE_ROWS = [
    ["Scopus",                   "17,021", "17,021", "6,218", "2,644", "2,644", "—"],
    ["Web of Science",           "15,179",  "4,683", "1,140",   "552",   "552", "—"],
    ["CAB Abstracts",             "5,723",  "3,229", "1,133",   "260",   "260", "—"],
    ["Academic Search Premier",   "1,187",    "274",    "66",    "20",    "20", "—"],
    ["AGRIS",                         "3",      "1",     "1",     "0",    "—", "—"],
    ["Total",                    "39,113", "25,208", "8,558", "3,476", "3,464", "2,368"],
]

CALIBRATION_ROWS = [
    ("R1 — initial criteria",  "205", "0.776", "0.703", "0.436", "0.500", False),
    ("R1a — 1st revision",     "205", "0.761", "0.797", "0.534", "0.500", False),
    ("R1b — 2nd revision",     "205", "0.866", "0.819", "0.645", "0.500", False),
    ("R2a — 3rd revision",     "103", "0.897", "0.905", "0.770", "0.765", False),
    ("R2b — 4th revision ✓",   "103", "0.966", "0.838", "0.720", "0.765", True),
    ("R3a — stability check ✓","107", "0.970", "0.824", "0.721", "0.703", True),
]

KEY_FINDINGS = [
    ("96%",  "A recent but thin evidence base",
     "96% of included studies published since 2015 (median: 2022). Growing fast but long-term outcome evidence is still missing."),
    ("86%",  "Marginalized groups largely invisible",
     "86% of studies do not disaggregate by women, youth, or indigenous peoples — a critical equity gap."),
    ("<6%",  "Non-crop producers underrepresented",
     "Fisheries/aquaculture (5.9%) and agroforestry (4.5%) nearly absent. Evidence skews heavily toward crop systems."),
    ("~80%", "Cost data almost entirely absent",
     "~80% of studies report no cost or efficiency data, making value-for-money assessment nearly impossible."),
    ("28%",  "Three countries dominate",
     "Ethiopia, Ghana, and Kenya account for 28% of all studies. Most LMICs and most of Africa remain uncovered."),
    ("—",    "Process outcomes studied; impacts are not",
     "Knowledge, adoption, and decision-making are well covered. Income, wellbeing, and resilience outcomes remain sparse."),
]


# ── build document ─────────────────────────────────────────────────────────────

def build(pngs: dict[str, Path | None]) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("DELIVERABLE 5")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN

    p = doc.add_paragraph()
    run = p.add_run("Final Systematic Map")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = CHARCOAL

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Measuring what matters: tracking climate adaptation processes\n"
        "and outcomes for smallholder producers in the agriculture sector"
    )
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    for line in [
        "Bristlepine Resilience Consultants  ·  ILRI",
        "June 2026  ·  v01",
        "Zarrar Khan, Jennifer Cisse, Caroline Staub",
    ]:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── Pipeline stats bar ────────────────────────────────────────────────────
    _add_label(doc, "Pipeline overview")
    _add_heading(doc, "Final pipeline summary", 1)

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("40,634", "Records identified\n(29 sources)"),
        ("26,173", "Unique after\ndeduplication"),
        ("8,753",  "Included at\nT&A screening"),
        ("2,368",  "Included at\nfull-text screening (LLM)"),
    ]
    for i, (val, lbl) in enumerate(stats):
        _set_cell_bg(tbl.rows[0].cells[i], "21472E")
        _cell_text(tbl.rows[0].cells[i], val, bold=True, color=WHITE, size=16, align="center")
        _set_cell_bg(tbl.rows[1].cells[i], "F5F2EC")
        _cell_text(tbl.rows[1].cells[i], lbl, color=CHARCOAL, size=8, align="center")
    doc.add_paragraph()

    # ── Key findings ──────────────────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "What the evidence shows")
    _add_heading(doc, "Key findings", 1)
    _add_body(doc,
        "Across 86 human-coded studies, the evidence base is recent, geographically concentrated, "
        "and skewed toward process outcomes and crop systems. Non-crop producers, marginalized groups, "
        "and cost data remain severely underrepresented."
    )

    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (stat, label, text) in enumerate(KEY_FINDINGS):
        row_i, col_i = divmod(idx, 2)
        cell = tbl.rows[row_i].cells[col_i]
        cell.paragraphs[0].clear()
        _set_cell_bg(cell, "F5F2EC")
        p = cell.paragraphs[0]
        r = p.add_run(label + "\n")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = CHARCOAL
        r2 = p.add_run(stat + "\n")
        r2.bold = True
        r2.font.size = Pt(20)
        r2.font.color.rgb = GREEN
        r3 = p.add_run(text)
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    doc.add_paragraph()

    doc.add_page_break()

    # ── ROSES / PRISMA flow ───────────────────────────────────────────────────
    _add_label(doc, "Record flow")
    _add_heading(doc, "ROSES flow diagram", 1)
    _add_body(doc,
        "Record flow across all 29 sources and four screening stages, following ROSES (Reporting "
        "Standards for Systematic Evidence Syntheses) requirements. Records are tracked from initial "
        "database returns through deduplication, title/abstract screening, full-text retrieval, and "
        "full-text screening to the final included set."
    )
    _add_figure(doc, pngs.get("roses"), "Figure 1. ROSES flow diagram — record flow across all databases and screening stages.", width_cm=16)

    # ── Methods overview table ────────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "Methodology")
    _add_heading(doc, "Methods summary", 1)
    _add_body(doc, "The review followed a pre-registered systematic map protocol (Zenodo DOI: 10.5281/zenodo.19811629). Steps 1–7 mirror the protocol and are documented in full on the project website.")

    tbl = doc.add_table(rows=8, cols=3)
    tbl.style = "Table Grid"
    method_rows = [
        ("Step", "Stage", "Key outcomes"),
        ("1", "Database search",         "29 sources · 40,634 records · 6 categories (bibliographic DBs, web search, UN, development agencies, research centres, M&E networks)"),
        ("2", "Deduplication",            "3-pass algorithm (DOI match → exact title+year → fuzzy title). 40,634 → 26,173 unique records (14,461 duplicates removed)"),
        ("3", "Calibration",              "6 rounds with 2 independent human reviewers (EPPI Reviewer). Final: sensitivity = 0.966–0.970, κ = 0.720–0.721. Threshold: sens ≥ 0.95 · κ ≥ 0.60"),
        ("4", "Title/abstract screening", "26,173 screened by LLM (qwen2.5:14b, T = 0.0). 8,753 included (Scopus: 6,218 · multi-DB: 2,535). Uncertain defaults to include"),
        ("5", "Full-text retrieval",      "8,748 FTs sought. 3,505 retrieved (40.1%) via Unpaywall (2,173), Elsevier API (1,100), Semantic Scholar (157), OpenAlex (40), CORE (33), other (2)"),
        ("6", "Full-text screening & coding", "LLM: 3,505 screened → 2,368 included (67.6%). Human: 86 papers across 5 batches (pure random, seeds 42–46). 20-field extraction schema"),
        ("7", "Information saturation",   "All 3 dimensions (domains, methodology, producer type) saturated by batch FT-R2c (49 papers). Zero new canonical categories added across final 2 batches (37 papers)"),
    ]
    for i, (step, stage, detail) in enumerate(method_rows):
        cells = tbl.rows[i].cells
        if i == 0:
            for j, txt in enumerate([step, stage, detail]):
                _set_cell_bg(cells[j], "21472E")
                _cell_text(cells[j], txt, bold=True, color=WHITE, size=9)
        else:
            _set_cell_bg(cells[0], "F5F2EC")
            _cell_text(cells[0], step, bold=True, color=GREEN, size=10, align="center")
            _cell_text(cells[1], stage, bold=True, color=CHARCOAL, size=9)
            _cell_text(cells[2], detail, color=CHARCOAL, size=9)

    # Calibration table
    doc.add_paragraph()
    _add_heading(doc, "Calibration rounds", 2)
    _add_body(doc,
        "Six calibration rounds used two independent human reviewers (Caroline Staub, Jennifer Cisse) "
        "screening the same paper sample in EPPI Reviewer. Criteria were revised between rounds. "
        "Full-corpus screening only commenced once sensitivity ≥ 0.95 AND κ ≥ 0.60 were both achieved."
    )
    tbl = doc.add_table(rows=1 + len(CALIBRATION_ROWS), cols=7)
    tbl.style = "Table Grid"
    for j, h in enumerate(["Round", "n", "Sensitivity", "Specificity", "LLM κ", "Human κ", "Pass"]):
        _set_cell_bg(tbl.rows[0].cells[j], "21472E")
        _cell_text(tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9, align="center")
    for i, (rnd, n, sens, spec, kappa, hkap, passed) in enumerate(CALIBRATION_ROWS):
        row = tbl.rows[i + 1]
        bg = "EDF7ED" if passed else "FFFFFF"
        for j, val in enumerate([rnd, n, sens, spec, kappa, hkap, "✓ Yes" if passed else "—"]):
            _set_cell_bg(row.cells[j], bg)
            bold = passed and j in (0, 2, 4)
            clr = GREEN if (passed and j == 2 and float(sens) >= 0.95) else CHARCOAL
            _cell_text(row.cells[j], val, bold=bold, color=clr, size=9,
                       align="center" if j > 0 else "left")
    p = doc.add_paragraph("R2b 95% CI (Wilson): 0.828–0.994. R3a: 0.847–0.995. Pooled (60/62 true positives): 0.890–0.991.")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    p.runs[0].italic = True

    doc.add_page_break()

    # ── Pipeline table (by database) ──────────────────────────────────────────
    _add_label(doc, "Search results")
    _add_heading(doc, "Table 1. Records by database and screening stage", 1)
    tbl = doc.add_table(rows=1 + len(PIPELINE_ROWS), cols=len(PIPELINE_HEADERS))
    tbl.style = "Table Grid"
    for j, h in enumerate(PIPELINE_HEADERS):
        _set_cell_bg(tbl.rows[0].cells[j], "21472E")
        _cell_text(tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9, align="center")
    for i, row_data in enumerate(PIPELINE_ROWS):
        row = tbl.rows[i + 1]
        is_total = (i == len(PIPELINE_ROWS) - 1)
        bg = "3C3C3C" if is_total else ("F5F2EC" if i % 2 == 0 else "FFFFFF")
        txt_color = WHITE if is_total else CHARCOAL
        for j, val in enumerate(row_data):
            _set_cell_bg(row.cells[j], bg.lstrip("#"))
            _cell_text(row.cells[j], val, bold=is_total, color=txt_color, size=9,
                       align="left" if j == 0 else "center")

    p = doc.add_paragraph(
        "Full-text screening conducted on all automatically retrieved records across all databases "
        "(n = 3,476 retrieved; 3,464 screened). Records requiring manual full-text access (n = 5,243) "
        "are not included in these totals. Coding column shows all LLM-screened included records "
        "combined across all databases."
    )
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    p.runs[0].italic = True

    doc.add_page_break()

    # ── Evidence Gap Map (Human — primary) ────────────────────────────────────
    _add_label(doc, "Primary output")
    _add_heading(doc, "Evidence gap map — human-coded results", 1)
    _add_body(doc,
        "The evidence gap map (EGM) is the primary deliverable. It plots the distribution of included "
        "studies across process/outcome domains (y-axis) and producer types (x-axis). Bubble size is "
        "proportional to the number of studies. Grey markers indicate evidence gaps (cells with no "
        "studies). Results below are based on the 86 human-coded papers, which constitute the "
        "authoritative output for this deliverable."
    )
    _add_figure(doc, pngs.get("egm_human"),
                "Figure 2. Evidence Gap Map — human-coded results (n = 86). "
                "Blue bubbles = process domains; green = outcome domains; grey = evidence gaps. "
                "Bubble area ∝ number of studies.",
                width_cm=16)

    doc.add_page_break()

    # ── LLM vs Human comparison ───────────────────────────────────────────────
    _add_label(doc, "Validation")
    _add_heading(doc, "LLM vs human comparison", 1)
    _add_body(doc,
        "To validate the LLM-extracted reference corpus, we compare the distribution of key variables "
        "between the 86 human-coded records and the full LLM-screened corpus (n = 2,368). Both tracks "
        "show broadly consistent patterns across producer types, domains, and geography, supporting the "
        "use of the LLM corpus as an exploratory reference. The human-coded results remain the "
        "authoritative source for all findings."
    )
    _add_figure(doc, pngs.get("llm_human"),
                "Figure 3. LLM vs human comparison across key categorical variables. "
                "Human (amber) vs LLM (teal). Bars show percentage of studies in each category.",
                width_cm=16)

    doc.add_page_break()

    # ── Domain heatmap ────────────────────────────────────────────────────────
    _add_label(doc, "Results")
    _add_heading(doc, "Domain heatmap — process and outcome dimensions", 1)
    _add_body(doc,
        "The domain heatmap shows the intensity of evidence across all 13 process and outcome "
        "domains, disaggregated by producer type. Darker cells indicate more studies. Process "
        "domains (decision-making, uptake/adoption, knowledge) are more densely covered than outcome "
        "domains (resilience, wellbeing, income)."
    )
    _add_figure(doc, pngs.get("heatmap"),
                "Figure 4. Process/outcome domain heatmap by producer type. "
                "Cell colour intensity ∝ number of studies.",
                width_cm=16)

    # ── Geographic distribution ───────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "Results")
    _add_heading(doc, "Geographic distribution", 2)
    _add_body(doc,
        "Studies are heavily concentrated in Sub-Saharan Africa and South Asia. Ethiopia, Ghana, "
        "and Kenya alone account for 28% of all studies. Large gaps persist across Central America, "
        "Southeast Asia, and MENA regions."
    )
    _add_figure(doc, pngs.get("geo_map"),
                "Figure 5. Geographic distribution of included studies (choropleth). "
                "Countries with no studies shown in grey.",
                width_cm=15)
    _add_figure(doc, pngs.get("geo_bar"),
                "Figure 6. Top countries by study count.",
                width_cm=12)

    doc.add_page_break()

    # ── Temporal trends ───────────────────────────────────────────────────────
    _add_label(doc, "Results")
    _add_heading(doc, "Temporal trends", 1)
    _add_body(doc,
        "Publication volume accelerated sharply after 2015, with 96% of included studies published "
        "in the past decade. The median publication year is 2022. This reflects both growth in the "
        "field and improved evidence infrastructure — but also means the long-term effectiveness "
        "evidence base is very thin."
    )
    _add_figure(doc, pngs.get("temporal"),
                "Figure 7. Publications per year across included studies.",
                width_cm=14)

    # ── Methodological approaches ─────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "Results")
    _add_heading(doc, "Methodological approaches", 2)
    _add_body(doc,
        "Quantitative and mixed-method studies predominate. Experimental designs (RCTs, quasi-experiments) "
        "are rare — less than 10% of the corpus — limiting causal inference about effectiveness. "
        "Qualitative approaches (interviews, focus groups) are common, especially for process outcomes."
    )
    _add_figure(doc, pngs.get("method"),
                "Figure 8. Primary methodological approach across included studies.",
                width_cm=13)

    # ── Equity & inclusion ────────────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "Results")
    _add_heading(doc, "Equity & inclusion", 2)
    _add_body(doc,
        "86% of studies provide no equity disaggregation. Women and gender are the most commonly "
        "reported equity dimension, but still feature in fewer than 10% of studies. Youth, indigenous "
        "peoples, people with disabilities, and pastoralist/migratory groups are nearly absent."
    )
    _add_figure(doc, pngs.get("equity"),
                "Figure 9. Equity and inclusion dimensions across included studies. "
                "Red bar = studies with no marginalized group focus.",
                width_cm=13)

    doc.add_page_break()

    # ── Information saturation ────────────────────────────────────────────────
    _add_label(doc, "Methodological note")
    _add_heading(doc, "Information saturation", 1)
    _add_body(doc,
        "Human coding continued until the distribution of key evidence characteristics stopped "
        "changing — not until all papers were coded. We tracked new canonical category discovery "
        "across three dimensions: process/outcome domains (13 values), methodological approaches (5), "
        "and producer types (5)."
    )
    _add_body(doc,
        "All three dimensions reached saturation by batch FT-R2c (49 papers coded), with zero new "
        "canonical categories added across the final two batches (FT-R2d and FT-R3, covering an "
        "additional 37 papers). This confirms that the extracted evidence space is well-characterised "
        "by the human sample without coding the full 8,748-record corpus."
    )
    _add_figure(doc, pngs.get("saturation"),
                "Figure 10. Information saturation curve. "
                "Top panel: cumulative unique categories as % of final total, by papers coded. "
                "Dashed line = 95% saturation threshold. Bottom panel: new categories per batch. "
                "All dimensions plateau by 49 papers.",
                width_cm=16)

    # ── Data availability ─────────────────────────────────────────────────────
    _add_hr(doc)
    _add_label(doc, "Open data")
    _add_heading(doc, "Data availability", 1)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    for j, h in enumerate(["Resource", "Location"]):
        _set_cell_bg(tbl.rows[0].cells[j], "21472E")
        _cell_text(tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9)
    data_rows = [
        ("Systematic map protocol (D3 + amendment)", "Zenodo: 10.5281/zenodo.19811629"),
        ("First draft systematic map (D4)",           "Zenodo: 10.5281/zenodo.19811622"),
        ("Interactive evidence gap map (HTML)",        "deliverables/evidence_gap_map_d5.html (also at project website)"),
        ("Searchable database (CSV)",                  "Project website: /systematic-map → Download CSV"),
    ]
    for i, (res, loc) in enumerate(data_rows):
        bg = "F5F2EC" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(tbl.rows[i + 1].cells[0], bg)
        _set_cell_bg(tbl.rows[i + 1].cells[1], bg)
        _cell_text(tbl.rows[i + 1].cells[0], res, size=9, color=CHARCOAL)
        _cell_text(tbl.rows[i + 1].cells[1], loc, size=9, color=CHARCOAL)

    doc.add_paragraph()
    _add_body(doc,
        "All pipeline code, criteria YAML, calibration data, and coding templates are publicly "
        "available on GitHub: github.com/bristlepine/ilri-climate-adaptation-effectiveness. "
        "The interactive evidence gap map (HTML) is available as a standalone file and as part "
        "of the project website. All figures are reproducible by re-running "
        "step16_map_visualise.py against the coded dataset."
    )

    # Save
    doc.save(str(OUT_DOCX))
    print(f"\n[docx] Saved: {OUT_DOCX}")
    print(f"[docx] Size:  {OUT_DOCX.stat().st_size / 1024:.0f} KB")


# ── entry point ────────────────────────────────────────────────────────────────

def main():
    print("[d5] Generating missing PNGs...")
    pngs = ensure_pngs()

    found = [k for k, v in pngs.items() if v and v.exists()]
    missing = [k for k, v in pngs.items() if not (v and v.exists())]
    print(f"[d5] Figures ready: {found}")
    if missing:
        print(f"[d5] Missing (will show placeholder): {missing}")

    print("[d5] Building Word document...")
    build(pngs)


if __name__ == "__main__":
    main()
