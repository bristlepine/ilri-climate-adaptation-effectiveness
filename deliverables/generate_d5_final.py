"""
generate_d5_final.py

Creates Deliverable 5 as a Word document (.docx) matching D4's exact format:
  - EB Garamond title (26pt, #4D311A)
  - Lato body text (11pt)
  - H1: 20pt bold #4D311A / H2: 16pt bold #3C3C3C
  - 1-inch margins all sides
  - Header: Bristlepine logo + document title
  - Footer: "Bristlepine Resilience Consultants" + page number
  - Green (#21472E) table headers, charcoal total rows
  - All figures embedded directly as PNGs

Run: conda run -n ilri01 python deliverables/generate_d5_final.py
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_HEADER_FOOTER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import nsmap

HERE   = Path(__file__).resolve().parent
ROOT   = HERE.parent
STEP16 = ROOT / "scripts" / "outputs" / "step16"
HUMAN  = STEP16 / "interactive" / "human"
LOGO   = HERE / "logo_bristlepine.png"
OUT    = HERE / "Deliverable_5_Bristlepine_Final_Systematic_Map_v01.docx"

# ── Brand colours (from D4 GDoc named styles) ─────────────────────────────────
BROWN   = RGBColor(0x4D, 0x31, 0x1A)   # #4D311A — H1, Title
GREY60  = RGBColor(0x3C, 0x3C, 0x3C)   # #3C3C3C — H2, Charcoal
GREY40  = RGBColor(0x66, 0x66, 0x66)   # #666666 — Subtitle
GREEN   = RGBColor(0x21, 0x47, 0x2E)   # #21472E — Table headers
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
SAND    = RGBColor(0xF5, 0xF2, 0xEC)

FONT_TITLE  = "EB Garamond"
FONT_BODY   = "Lato"
FONT_SUB    = "Arial"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _cell_bg(cell, hex6: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.lstrip("#"))
    tcPr.append(shd)


def _cell_write(cell, text: str, bold=False, italic=False,
                color: RGBColor | None = None, size: int = 9,
                align: str = "left") -> None:
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = FONT_BODY
    if color:
        run.font.color.rgb = color
    al = {"center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    para.alignment = al
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)


def _add_run(para, text: str, bold=False, italic=False,
             color: RGBColor | None = None, size: float | None = None,
             font: str | None = None) -> None:
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font or FONT_BODY
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.clear()
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(20)
        run.font.color.rgb = BROWN
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(14)
        run.font.color.rgb = GREY60
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = FONT_BODY
        run.font.size = Pt(12)
        run.font.color.rgb = GREY60


def _body(doc: Document, text: str, space_after: float = 8) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.name = FONT_BODY
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.6)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    _add_run(p, "•  ", bold=True, color=GREEN, size=11)
    _add_run(p, text, size=11)


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.name = FONT_BODY
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = GREY40
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(14)


def _figure(doc: Document, png: Path, caption: str, width_cm: float = 15.5) -> None:
    if png and png.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(str(png), width=Cm(width_cm))
        _caption(doc, caption)
    else:
        p = doc.add_paragraph(f"[Figure: {png.name if png else '?'} — not found]")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D0C8BD")
    pBdr.append(bot)
    pPr.append(pBdr)


def _table_style(doc: Document, rows: list[list[str]],
                 col_widths: list[float] | None = None) -> None:
    """Render a styled table. rows[0] = header, rows[-1] = total if starts with 'Total'."""
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    for r_idx, row_data in enumerate(rows):
        is_header = (r_idx == 0)
        is_total  = (r_idx == len(rows) - 1 and row_data[0].startswith("Total"))
        cells = tbl.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            if is_header:
                _cell_bg(cells[c_idx], "21472E")
                _cell_write(cells[c_idx], val, bold=True, color=WHITE, size=9,
                            align="center" if c_idx > 0 else "left")
            elif is_total:
                _cell_bg(cells[c_idx], "3C3C3C")
                _cell_write(cells[c_idx], val, bold=True, color=WHITE, size=9,
                            align="center" if c_idx > 0 else "left")
            else:
                bg = "F9F7F4" if r_idx % 2 == 0 else "FFFFFF"
                _cell_bg(cells[c_idx], bg)
                _cell_write(cells[c_idx], val, size=9,
                            align="center" if c_idx > 0 else "left")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.name = FONT_BODY
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = GREY40
    p.paragraph_format.space_after = Pt(10)


# ── Header / Footer ───────────────────────────────────────────────────────────

def _setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Main header (all pages except first)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO.exists():
        run_logo = hp.add_run()
        run_logo.add_picture(str(LOGO), height=Pt(20))
        hp.add_run("  ")
    run_title = hp.add_run("Measuring what matters — Systematic Map")
    run_title.font.name = FONT_BODY
    run_title.font.size = Pt(8)
    run_title.font.color.rgb = GREY40
    # Bottom border on header
    hpPr = hp._p.get_or_add_pPr()
    hpBdr = OxmlElement("w:pBdr")
    hbot = OxmlElement("w:bottom")
    hbot.set(qn("w:val"), "single"); hbot.set(qn("w:sz"), "4")
    hbot.set(qn("w:space"), "1"); hbot.set(qn("w:color"), "D0C8BD")
    hpBdr.append(hbot); hpPr.append(hpBdr)

    # First-page header: empty
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    if first_header.paragraphs:
        first_header.paragraphs[0].clear()

    # Footer: Bristlepine + page number (right)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = fp.add_run("Bristlepine Resilience Consultants")
    left_run.font.name = FONT_BODY
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = GREY40
    # Tab stop + page number on right
    tab_run = fp.add_run("\t")
    tab_run.font.size = Pt(8)
    # Add page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    pg_run = fp.add_run()
    pg_run.font.size = Pt(8)
    pg_run.font.color.rgb = GREY40
    pg_run._r.append(fldChar1)
    pg_run._r.append(instrText)
    pg_run._r.append(fldChar2)
    # Right-align tab stop at right margin
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9072")  # 9072 twips = 6.3 inches (width minus margins)
    tabs.append(tab); pPr.append(tabs)


# ── Data ──────────────────────────────────────────────────────────────────────

PIPELINE_TABLE = [
    ["Database", "Returned", "After Dedup", "Abstr. Incl.", "FT Retrieved", "FT Screened", "Coded"],
    ["Scopus",                   "17,021", "17,021", "6,218", "2,644", "2,644", "—"],
    ["Web of Science",           "15,179",  "4,683", "1,140",   "552",   "552", "—"],
    ["CAB Abstracts",             "5,723",  "3,229", "1,133",   "260",   "260", "—"],
    ["Academic Search Premier",   "1,187",    "274",    "66",    "20",    "20", "—"],
    ["AGRIS",                         "3",      "1",     "1",     "0",    "—", "—"],
    ["Total",                    "39,113", "25,208", "8,558", "3,476", "3,464", "2,368"],
]

CAL_TABLE = [
    ["Round", "n", "Sensitivity", "Specificity", "LLM κ", "Human κ", "Pass"],
    ["R1 — initial criteria",    "205", "0.776", "0.703", "0.436", "0.500", "No"],
    ["R1a — 1st revision",       "205", "0.761", "0.797", "0.534", "0.500", "No"],
    ["R1b — 2nd revision",       "205", "0.866", "0.819", "0.645", "0.500", "No"],
    ["R2a — 3rd revision",       "103", "0.897", "0.905", "0.770", "0.765", "No"],
    ["R2b — 4th revision",       "103", "0.966", "0.838", "0.720", "0.765", "YES ✓"],
    ["R3a — stability check",    "107", "0.970", "0.824", "0.721", "0.703", "YES ✓"],
]

KEY_FINDINGS = [
    ("96%",  "A recent but thin evidence base",
     "96% of studies published since 2015 (median 2022). Growing fast — but long-term impact evidence is still missing."),
    ("86%",  "Marginalized groups largely invisible",
     "86% of studies provide no equity disaggregation by gender, age, or ethnicity."),
    ("<6%",  "Non-crop producers severely underrepresented",
     "Fisheries/aquaculture and agroforestry barely present. Evidence skews heavily toward crop farming."),
    ("~80%", "Cost data almost entirely absent",
     "~80% of studies report no cost or efficiency data, making value-for-money assessment near-impossible."),
    ("28%",  "Three countries dominate",
     "Ethiopia, Ghana, and Kenya account for 28% of all studies. Most of Africa and Asia remain very poorly covered."),
    ("—",    "Process outcomes studied; impacts are not",
     "Knowledge, adoption, and decision-making are well covered. Income, wellbeing, and resilience remain sparse."),
]


# ── Build document ─────────────────────────────────────────────────────────────

def build() -> None:
    doc = Document()

    # Margins: 1 inch all sides (matching D4)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    _setup_header_footer(doc)

    # ── Cover page ─────────────────────────────────────────────────────────────
    # Logo centred
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after  = Pt(12)
        p.add_run().add_picture(str(LOGO), height=Pt(48))

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(8)
    _add_run(p, "Measuring what matters: tracking climate adaptation processes and outcomes "
                "for smallholder producers in the agriculture sector",
             bold=False, color=BROWN, size=24, font=FONT_TITLE)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, "Deliverable 5: Final Systematic Map", color=GREY40, size=15, font=FONT_SUB)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    _add_run(p, "June 2026, v01", color=GREY40, size=12, font=FONT_SUB)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    for line in [
        "Jennifer Denno Cissé*, jenn@bristlep.com — Bristlepine Resilience Consultants",
        "Zarrar Khan, zarrar@bristlep.com — Bristlepine Resilience Consultants",
        "Caroline G. Staub, caroline@bristlep.com — Bristlepine Resilience Consultants",
    ]:
        _add_run(p, line + "\n", color=GREY60, size=10, font=FONT_BODY)

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, "Keywords: Climate Adaptation · Smallholder Producers · Systematic Map · "
                "Evidence Synthesis · LMIC · M&E", color=GREY40, size=9.5, font=FONT_BODY)

    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────────────
    _heading(doc, "1.  Executive Summary", 1)
    _body(doc,
        "This document presents the final systematic map of methods used to track and measure "
        "climate adaptation processes and outcomes for smallholder producers across low- and "
        "middle-income countries (LMICs). It fulfils Deliverable 5 of the contract with the "
        "International Livestock Research Institute (ILRI).")
    _body(doc,
        "A total of 39,113 records were identified across five bibliographic databases and 24 "
        "grey literature sources. After deduplication, 25,208 unique records entered title and "
        "abstract screening. The validated LLM screening tool (sensitivity 0.966–0.970; κ = "
        "0.720–0.721) identified 8,558 records for full-text assessment. Of 3,476 full texts "
        "retrieved (40.1%), 3,464 were screened and 2,368 were included.")
    _body(doc,
        "A random sample of 86 studies was selected for human-coded data extraction across "
        "five batches. Information saturation was reached by batch FT-R2c (49 papers): all "
        "three tracked dimensions — process/outcome domains, methodological approaches, and "
        "producer types — showed no new canonical categories across the final two batches. The "
        "human-coded findings are representative of the broader evidence base.")
    _body(doc, "Key findings from the human-coded results (n = 86):", space_after=4)
    for stat, label, text in KEY_FINDINGS:
        _bullet(doc, f"{label} ({stat}): {text}")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc,
        "An interactive evidence gap map and searchable database are available at the project "
        "website. The LLM-extracted corpus (n = 2,368) is provided as an exploratory reference "
        "alongside the human-coded primary results. The systematic map protocol was amended to "
        "document all deviations from the pre-registered protocol (D5.7; Zenodo: "
        "10.5281/zenodo.19811629).")

    # ── 2. Background ──────────────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "2.  Background", 1)
    _body(doc,
        "The background, theory of change, stakeholder engagement process, and full PCCM "
        "eligibility criteria are documented in the protocol published as Deliverable 3 "
        "(Bristlepine Resilience Consultants, January 2026; Zenodo: 10.5281/zenodo.19811629).")
    _body(doc,
        "A growing number of climate adaptation interventions are being deployed across the "
        "agriculture sector to support smallholder producers in LMICs. Despite large investments, "
        "the evidence base on whether and how adaptation processes and outcomes are tracked — and "
        "what methods are used to do so — remains fragmented. This systematic map provides a "
        "structured overview of that evidence base to support better M&E practice and "
        "research prioritisation.")
    _body(doc,
        "The map covers both adaptation processes (knowledge, decision-making, adoption, "
        "behavioural change, participation, governance, access to services) and adaptation "
        "outcomes (yields, income, livelihoods, wellbeing, risk reduction, resilience), as "
        "defined in the PCCM framework.")

    # ── 3. Methods ────────────────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "3.  Methods", 1)
    _body(doc,
        "This systematic map follows the protocol published as Deliverable 3. A summary of "
        "methods is provided here; technical details — calibration data, software stack, and "
        "saturation analysis — are in the Appendix.")

    _heading(doc, "3.1  Search Strategy", 2)
    _body(doc,
        "Searches were conducted in January 2026 using a structured Boolean search string "
        "developed around the PCCM framework. The Scopus search returned 17,021 records. "
        "Additional databases searched in parallel: Web of Science Core Collection (15,179), "
        "CAB Abstracts (5,723), Academic Search Premier (1,187), and AGRIS (3). "
        "Twenty-four grey literature sources were manually searched. Total across all sources: "
        "39,113 records.")

    _heading(doc, "3.2  Deduplication", 2)
    _body(doc,
        "Records from all sources were deduplicated using a three-pass algorithm: exact DOI "
        "match, then exact title-and-year match, then fuzzy title match within the same year "
        "(Jaccard similarity ≥ 0.85). This reduced 39,113 records to 25,208 unique records "
        "(14,905 duplicates removed).")

    _heading(doc, "3.3  Screening", 2)
    _body(doc,
        "Title and abstract screening used a validated LLM-assisted tool (Ollama/qwen2.5:14b), "
        "calibrated through six rounds with two independent human reviewers (Caroline Staub, "
        "Jennifer Cisse). Final performance: sensitivity = 0.966–0.970; κ = 0.720–0.721 — "
        "exceeding the pre-specified thresholds (sensitivity ≥ 0.95; κ ≥ 0.60). A total of "
        "8,558 records were included at title/abstract stage. Full-text screening of 3,464 "
        "retrieved records identified 2,368 included studies.")

    _heading(doc, "3.4  Data Extraction", 2)
    _body(doc,
        "Given the large volume of included studies, data extraction used a random sampling "
        "approach. Papers were drawn in batches of 20 using fixed integer seeds (42, 43, …), "
        "with each batch explicitly excluding previously drawn DOIs. Coders applied the five "
        "PCCM inclusion criteria to confirm inclusion, then completed all 16 extraction fields "
        "(country, scale, producer type, adaptation focus, domain, methodology, equity, "
        "intervention type, etc.).")
    _body(doc,
        "Information saturation was tracked at each batch across three dimensions: "
        "process/outcome domains (13 canonical values), methodological approaches (5), and "
        "producer types (5). Saturation — defined as zero new canonical categories across an "
        "entire batch — was reached by batch FT-R2c (49 papers). Five batches and 86 papers "
        "were coded in total.")

    # ── 4. Final Results ──────────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "4.  Final Results", 1)

    _heading(doc, "4.1  Search and Screening Summary", 2)
    _body(doc,
        "A total of 39,113 records were retrieved across five databases (Table 1). After "
        "deduplication, 25,208 unique records remained. Title and abstract screening identified "
        "8,558 records for full-text assessment (Scopus: 6,218; multi-database net-new: 2,340). "
        "Of 3,476 full texts retrieved (40.1%), 3,464 were screened and 2,368 were included "
        "after full-text screening.")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _table_style(doc, PIPELINE_TABLE,
                 col_widths=[4.5, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8])
    _note(doc,
        "Full-text screening was conducted on all automatically retrieved records across all "
        "databases (3,476 retrieved; 3,464 screened; 12 unprocessable). Records requiring "
        "manual library access (n = 5,243) are not included above. Coded column = all "
        "LLM-screened included records combined.")

    _body(doc,
        "Figure 1 presents the ROSES flow diagram showing the full record flow across all "
        "29 sources and all four screening stages.")
    _figure(doc, STEP16 / "roses_flow.png",
            "Figure 1. ROSES flow diagram — record flow across 29 sources and four screening "
            "stages (n = 39,113 identified; 2,368 included after full-text screening).")

    _heading(doc, "4.2  Evidence Gap Map", 2)
    _body(doc,
        "The evidence gap map (EGM) is the primary output of this deliverable. It plots the "
        "distribution of included human-coded studies (n = 86) across process and outcome "
        "domains (y-axis) and producer types (x-axis). Bubble size is proportional to the "
        "number of studies in each cell; grey markers indicate evidence gaps.")
    _figure(doc, HUMAN / "evidence_gap_map.png",
            "Figure 2. Evidence gap map — human-coded results (n = 86). Blue = process domains; "
            "green = outcome domains; grey = evidence gaps. Bubble area ∝ number of studies.",
            width_cm=15.5)

    _body(doc, "Key observations:")
    _bullet(doc,
        "Process domains (knowledge/awareness, decision-making, uptake/adoption) are more "
        "densely covered than outcome domains (income, wellbeing, resilience). No studies "
        "in the human sample measured resilience or adaptive capacity for livestock, fisheries, "
        "agroforestry, or mixed-system producers.")
    _bullet(doc,
        "Crop farming dominates across all domains. Livestock, fisheries/aquaculture, "
        "agroforestry, and mixed-system producers are consistently underrepresented, "
        "with most cells showing evidence gaps.")
    _bullet(doc,
        "Institutional governance and access to services are nearly absent across all "
        "producer types — a critical gap for systemic adaptation measurement.")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _heading(doc, "4.3  Geographic Distribution", 2)
    _body(doc,
        "Studies are heavily concentrated in Sub-Saharan Africa and South Asia. Ethiopia, Ghana, "
        "and Kenya account for 28% of all human-coded studies. Coverage is very sparse across "
        "Central America, Southeast Asia, the Sahel, and MENA.")
    _figure(doc, STEP16 / "geographic_map.png",
            "Figure 3. Geographic distribution of included studies. Countries by study count; "
            "multi-country studies counted in each country.")
    _figure(doc, STEP16 / "geographic_bar.png",
            "Figure 4. Top countries by study count.", width_cm=13)

    _heading(doc, "4.4  Temporal Trends", 2)
    _body(doc,
        "Publication volume has grown sharply since 2015, with 96% of included studies "
        "published in the past decade and a median year of 2022. This reflects growth in the "
        "field — but also means that long-term impact evidence is very thin, as most studies "
        "were too recently published to capture multi-year outcomes.")
    _figure(doc, STEP16 / "temporal_trends.png",
            "Figure 5. Temporal trends in publication volume, 1990–2025 (all databases, final).",
            width_cm=13)

    _heading(doc, "4.5  Methodological Approaches", 2)
    _body(doc,
        "Survey-based and qualitative approaches dominate the evidence base. Experimental "
        "designs (RCTs, quasi-experiments) represent fewer than 10% of studies, severely "
        "limiting causal inference about adaptation effectiveness. Participatory and "
        "mixed-method approaches are common, particularly for process-level outcomes.")
    _figure(doc, STEP16 / "methodology_bar.png",
            "Figure 6. Distribution of methodological approaches across included studies.",
            width_cm=13)

    _heading(doc, "4.6  Equity and Inclusion", 2)
    _body(doc,
        "86% of included studies provide no equity disaggregation. Where equity is addressed, "
        "gender is the most common focus — but appears in fewer than 15% of studies. Youth, "
        "indigenous peoples, people with disabilities, and pastoralist/migratory groups are "
        "nearly absent. This is a critical gap: adaptation outcomes are rarely equivalent "
        "across social groups, yet the evidence base does not systematically track who benefits.")
    _figure(doc, STEP16 / "equity_bar.png",
            "Figure 7. Equity and inclusion dimensions across included studies. "
            "Red bar = studies with no marginalized group focus.", width_cm=13)

    _heading(doc, "4.7  Information Saturation", 2)
    _body(doc,
        "Figure 8 shows the saturation curve tracking how quickly new evidence categories "
        "were discovered as human coding progressed. All three tracked dimensions plateaued "
        "by 49 papers (batch FT-R2c), with zero new canonical categories added in the final "
        "two batches (37 additional papers). This confirms the human-coded sample is "
        "representative of the broader evidence space.")
    _figure(doc, STEP16 / "saturation.png",
            "Figure 8. Information saturation curve. Top: cumulative unique categories as % "
            "of final total. Bottom: new categories per batch. All dimensions plateau by 49 papers.")

    _heading(doc, "4.8  LLM Corpus — Exploratory Reference", 2)
    _body(doc,
        "In parallel, the LLM-assisted pipeline extracted data from all 2,368 LLM-screened "
        "records. This larger corpus is available on the project website as an exploratory "
        "reference (toggled separately from the human-coded results). The distribution of "
        "key variables is broadly consistent between tracks, supporting the robustness of "
        "the human-coded findings. LLM-extracted data have not been individually validated "
        "and should not be treated as authoritative.")
    _figure(doc, STEP16 / "llm_vs_human.png",
            "Figure 9. LLM vs human comparison across key categorical variables. "
            "Human (amber, n=86) vs LLM (teal, n=2,368). Bars show % of studies in each category.")

    # ── 5. Searchable Database ────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "5.  Searchable Database", 1)
    _body(doc, "The searchable database is available at: "
               "https://main.d3a9jahzuu7xai.amplifyapp.com/systematic-map")
    _body(doc,
        "The database includes all 2,368 LLM-screened records with key metadata (year, title, "
        "country, producer type, domain type, methodological approach). Users can filter by "
        "any combination of fields and download the full CSV. A toggle on the website switches "
        "between the human-coded results (n = 86, authoritative) and the LLM corpus "
        "(n = 2,368, exploratory).")

    # ── 6. Limitations ────────────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "6.  Limitations", 1)
    _body(doc,
        "Full-text retrieval rate. Only 40.1% of abstract-included records (3,476 of 8,748 "
        "full texts sought) were retrieved automatically. Records requiring manual library "
        "access are not included in the full-text screening totals. Absence of a full text "
        "is not grounds for exclusion — these records remain in the database.")
    _body(doc,
        "Human sample size. The 86 human-coded papers are sufficient for saturation across "
        "the three tracked dimensions, but individual cells in the evidence gap map may have "
        "very low counts. Findings about specific domain–producer combinations should be "
        "treated as indicative rather than definitive.")
    _body(doc,
        "LLM corpus (exploratory only). The 2,368 LLM-extracted records have not been "
        "individually validated by human reviewers. LLM extraction errors in individual "
        "fields are likely. The LLM corpus is provided as a reference and labelled "
        "accordingly on the project website.")
    _body(doc,
        "Grey literature. Manual searching of grey literature sources was conducted but is "
        "likely incomplete. Organisational reports not indexed in bibliographic databases "
        "may be underrepresented.")

    # ── 7. Conclusions ────────────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "7.  Conclusions", 1)
    _body(doc,
        "This final systematic map provides a structured characterisation of the evidence base "
        "on methods used to track and measure climate adaptation processes and outcomes for "
        "smallholder producers in LMICs. The evidence base is recent, growing, and "
        "geographically concentrated — but it has significant gaps.")
    _body(doc,
        "The dominant finding is an evidence mismatch: the literature measures early-stage "
        "process outcomes (knowledge, adoption, decision-making) far better than downstream "
        "impact outcomes (income, wellbeing, resilience). For non-crop producers — fishers, "
        "pastoralists, agroforestry practitioners — even process-level evidence is sparse. "
        "Equity disaggregation is almost entirely absent.")
    _body(doc,
        "These gaps have direct implications for M&E practice: current monitoring systems "
        "are not designed to detect adaptation impact, and they systematically overlook "
        "marginalised groups and non-crop livelihood systems. Addressing these gaps will "
        "require deliberate investments in long-term cohort studies, equity-disaggregated "
        "data collection, and expansion beyond crop-focused producer types.")
    _body(doc,
        "The next phase (Deliverable 6) will develop a systematic review and meta-analysis "
        "protocol targeting the most tractable domains identified in this map — focusing on "
        "studies with sufficient outcome measurement to enable quantitative synthesis.")

    # ── 8. Acknowledgements ───────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "8.  Acknowledgements", 1)
    _body(doc,
        "This work was funded by the International Livestock Research Institute (ILRI) under "
        "the CGIAR Initiative on Climate Resilience (ClimBeR). We thank Aditi Krishnapriyan "
        "and Neal Hockley for their guidance and review throughout this project.")

    # ── 9. References ─────────────────────────────────────────────────────────
    _hr(doc)
    _heading(doc, "9.  References", 1)
    for ref in [
        "Bristlepine Resilience Consultants (2026). Deliverable 3: Final Systematic Map "
        "Protocol — Measuring what matters: tracking climate adaptation processes and outcomes "
        "for smallholder producers in the agriculture sector. Zenodo. "
        "https://doi.org/10.5281/zenodo.19811629",

        "Bristlepine Resilience Consultants (2026). Deliverable 4: First Draft Systematic Map "
        "(Preliminary). Zenodo. https://doi.org/10.5281/zenodo.19811622",

        "Collaboration for Environmental Evidence (2018). Guidelines and Standards for Evidence "
        "Synthesis in Environmental Management. Version 5.0.",

        "Haddaway, N.R., Macura, B., Whaley, P. and Pullin, A.S. (2018). ROSES Reporting "
        "standards for Systematic Evidence Syntheses. Environmental Evidence, 7, 7.",

        "O'Mara-Eves, A. et al. (2015). Using text mining for study identification in "
        "systematic reviews. Systematic Reviews, 4, 5.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent       = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after       = Pt(5)
        _add_run(p, ref, size=10)

    # ── Appendix: Technical Methods ───────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Appendix: Technical Methods", 1)
    _body(doc,
        "This appendix provides technical details supplementing Section 3. It is intended for "
        "reviewers assessing methodological rigour.")

    _heading(doc, "A1.  Calibration Rounds", 2)
    _body(doc,
        "Six calibration rounds with two independent human reviewers (Caroline Staub, Jennifer "
        "Cisse) screening in EPPI Reviewer. Criteria revised between rounds. Full-corpus "
        "screening began only once sensitivity ≥ 0.95 AND κ ≥ 0.60 were both achieved.")
    _table_style(doc, CAL_TABLE)
    _note(doc, "R2b 95% CI (Wilson): 0.828–0.994. R3a: 0.847–0.995. "
               "Pooled (60/62 true positives): 0.890–0.991.")

    _heading(doc, "A2.  Full-text Retrieval Sources", 2)
    _table_style(doc, [
        ["Source", "Records", "Format"],
        ["Unpaywall",               "2,173", "PDF"],
        ["Elsevier full-text API",  "1,100", "PDF / HTML"],
        ["Semantic Scholar",          "157", "PDF"],
        ["OpenAlex",                   "40", "PDF"],
        ["CORE",                       "33", "PDF"],
        ["Other",                       "2", "HTML"],
        ["Total",                   "3,505", "—"],
    ], col_widths=[8.0, 3.0, 3.0])

    _heading(doc, "A3.  Software Stack", 2)
    _table_style(doc, [
        ["Component", "Tool", "Purpose"],
        ["LLM inference",        "Ollama (qwen2.5:14b)",                "Local deterministic screening and extraction"],
        ["Scopus API",           "Elsevier REST API",                   "Record retrieval and abstract enrichment"],
        ["DOI enrichment",       "CrossRef, OpenAlex, Semantic Scholar","DOI lookup and abstract retrieval"],
        ["Open access",          "Unpaywall API",                       "Full-text URL discovery"],
        ["PDF parsing",          "pypdf",                               "Full-text extraction from PDFs"],
        ["HTML parsing",         "trafilatura, BeautifulSoup4",         "Full-text extraction from web sources"],
        ["IRR statistics",       "Custom Python (Cohen's κ)",           "Inter-rater reliability analysis"],
        ["Human screening",      "EPPI Reviewer",                       "Human screening and RIS exports"],
        ["Visualisation",        "matplotlib, Plotly + kaleido",        "Static and interactive figures"],
        ["Website",              "Next.js + AWS Amplify",               "Interactive evidence gap map + database"],
        ["Code repository",      "GitHub",                              "github.com/bristlepine/ilri-climate-..."],
    ], col_widths=[4.0, 5.0, 6.5])

    doc.save(str(OUT))
    print(f"\n[d5] ✓ Saved: {OUT}")
    print(f"[d5]   Size: {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
